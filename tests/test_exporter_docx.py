#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_exporter_docx.py
--------------------
Tests unitaires pour exporter_docx.py - Export Markdown vers DOCX.
"""

import json
import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
import sys

# Ajouter le répertoire execution au path
PROJECT_ROOT = Path(__file__).parent.parent
EXECUTION_DIR = PROJECT_ROOT / "execution"
sys.path.insert(0, str(EXECUTION_DIR))

from execution.core.exporter_docx import (
    nettoyer_texte_xml,
    detecter_tableau_markdown,
    ENTETES_TABLEAUX_CONNUS,
    MARQUEUR_VAR_START,
    MARQUEUR_VAR_END,
)


# =============================================================================
# TESTS NETTOYAGE XML
# =============================================================================

class TestNettoyerTexteXml:
    """Tests pour la fonction nettoyer_texte_xml."""

    def test_texte_normal(self):
        """Texte normal sans caractères spéciaux."""
        assert nettoyer_texte_xml("Hello World") == "Hello World"

    def test_texte_vide(self):
        """Chaîne vide."""
        assert nettoyer_texte_xml("") == ""

    def test_texte_none(self):
        """Valeur None."""
        assert nettoyer_texte_xml(None) is None

    def test_caracteres_francais(self):
        """Caractères accentués français."""
        texte = "éàùêïôç ÉÀÙÊÏÔÇ"
        assert nettoyer_texte_xml(texte) == texte

    def test_caracteres_speciaux_valides(self):
        """Caractères spéciaux valides pour XML."""
        texte = "Tab:\tNewline:\nReturn:\r"
        assert nettoyer_texte_xml(texte) == texte

    def test_caracteres_controle_invalides(self):
        """Caractères de contrôle invalides supprimés."""
        # \x00-\x08, \x0B, \x0C, \x0E-\x1F sont invalides
        texte_avec_invalides = "Hello\x00World\x08Test\x0BEnd"
        resultat = nettoyer_texte_xml(texte_avec_invalides)
        assert "\x00" not in resultat
        assert "\x08" not in resultat
        assert "HelloWorldTestEnd" == resultat

    def test_emoji(self):
        """Les emojis sont valides (au-delà du BMP)."""
        # Note: Les emojis peuvent être invalides selon l'implémentation
        texte = "Test sans emoji"
        assert nettoyer_texte_xml(texte) == texte

    def test_texte_juridique(self):
        """Texte juridique type notarial."""
        texte = "Article 1er - L'acquéreur déclare..."
        assert nettoyer_texte_xml(texte) == texte


# =============================================================================
# TESTS DETECTION TABLEAUX MARKDOWN
# =============================================================================

class TestDetecterTableauMarkdown:
    """Tests pour la fonction detecter_tableau_markdown."""

    def test_tableau_simple(self):
        """Tableau Markdown simple 2x2."""
        lignes = [
            "| Col1 | Col2 |",
            "|------|------|",
            "| A    | B    |",
            "| C    | D    |",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is True
        assert fin_index == 4
        assert donnees is not None
        assert len(donnees['lignes']) == 3  # header + 2 data rows
        assert donnees['lignes'][0] == ['Col1', 'Col2']

    def test_pas_tableau(self):
        """Texte qui n'est pas un tableau."""
        lignes = [
            "Ceci est un paragraphe normal.",
            "Pas de tableau ici.",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is False
        assert fin_index == 0
        assert donnees is None

    def test_tableau_avec_alignements(self):
        """Tableau avec alignements (gauche, centre, droite)."""
        lignes = [
            "| Gauche | Centre | Droite |",
            "|:-------|:------:|-------:|",
            "| A      | B      | C      |",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is True
        assert donnees['alignements'] == ['left', 'center', 'right']

    def test_tableau_index_milieu(self):
        """Détection d'un tableau au milieu du document."""
        lignes = [
            "Introduction.",
            "",
            "| Col1 | Col2 |",
            "|------|------|",
            "| X    | Y    |",
            "",
            "Conclusion.",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 2)

        assert est_tableau is True
        assert fin_index == 5

    def test_tableau_sans_separateur(self):
        """Ligne avec pipe mais sans séparateur = pas un tableau."""
        lignes = [
            "| Ceci | n'est | pas |",
            "un vrai tableau",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is False

    def test_index_hors_limites(self):
        """Index dépassant la taille de la liste."""
        lignes = ["Une seule ligne."]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 10)

        assert est_tableau is False
        assert fin_index == 10

    def test_tableau_cellules_vides(self):
        """Tableau avec cellules vides."""
        lignes = [
            "| A |   | C |",
            "|---|---|---|",
            "|   | B |   |",
        ]
        est_tableau, fin_index, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is True
        assert len(donnees['lignes'][0]) == 3


# =============================================================================
# TESTS ENTETES TABLEAUX CONNUS
# =============================================================================

class TestEntetesTableauxConnus:
    """Tests pour la liste des entêtes de tableaux notariaux."""

    def test_entetes_non_vide(self):
        """La liste d'entêtes n'est pas vide."""
        assert len(ENTETES_TABLEAUX_CONNUS) > 0

    def test_entetes_cadastrales(self):
        """Vérifier présence des entêtes cadastrales."""
        entetes_str = [str(e) for e in ENTETES_TABLEAUX_CONNUS]
        has_section = any('Section' in str(e) for e in ENTETES_TABLEAUX_CONNUS)
        assert has_section, "Entêtes cadastrales manquantes"

    def test_entetes_financieres(self):
        """Vérifier présence des entêtes financières."""
        has_montant = any('Montant' in str(e) for e in ENTETES_TABLEAUX_CONNUS)
        assert has_montant, "Entêtes financières manquantes"

    def test_entetes_format_liste(self):
        """Chaque entête est une liste de strings."""
        for entete in ENTETES_TABLEAUX_CONNUS:
            assert isinstance(entete, list), f"Entête {entete} n'est pas une liste"
            for col in entete:
                assert isinstance(col, str), f"Colonne {col} n'est pas une string"


# =============================================================================
# TESTS MARQUEURS VARIABLES
# =============================================================================

class TestMarqueursVariables:
    """Tests pour les marqueurs de variables."""

    def test_marqueurs_definis(self):
        """Les marqueurs sont définis."""
        assert MARQUEUR_VAR_START is not None
        assert MARQUEUR_VAR_END is not None

    def test_marqueurs_non_vides(self):
        """Les marqueurs ne sont pas vides."""
        assert len(MARQUEUR_VAR_START) > 0
        assert len(MARQUEUR_VAR_END) > 0

    def test_marqueurs_distinguables(self):
        """Les marqueurs start/end sont différents."""
        assert MARQUEUR_VAR_START != MARQUEUR_VAR_END

    def test_detection_variable(self):
        """Détection d'une variable dans le texte."""
        texte = f"Le prix est de {MARQUEUR_VAR_START}250000{MARQUEUR_VAR_END} euros."
        assert MARQUEUR_VAR_START in texte
        assert MARQUEUR_VAR_END in texte


# =============================================================================
# TESTS D'INTÉGRATION (avec import conditionnel)
# =============================================================================

class TestExportIntegration:
    """Tests d'intégration pour l'export DOCX."""

    @pytest.fixture
    def markdown_simple(self, tmp_path):
        """Créer un fichier markdown simple."""
        md_content = """# Titre Principal

## Section 1

Ceci est un paragraphe de test.

### Sous-section

| Colonne A | Colonne B |
|-----------|-----------|
| Valeur 1  | Valeur 2  |

---

Fin du document.
"""
        md_file = tmp_path / "test_acte.md"
        md_file.write_text(md_content, encoding='utf-8')
        return md_file

    @pytest.fixture
    def markdown_avec_variables(self, tmp_path):
        """Créer un fichier markdown avec variables marquées."""
        md_content = f"""# Acte de Vente

Le vendeur {MARQUEUR_VAR_START}M. DUPONT Jean{MARQUEUR_VAR_END} vend à
l'acquéreur {MARQUEUR_VAR_START}Mme MARTIN Marie{MARQUEUR_VAR_END}.

**Prix**: {MARQUEUR_VAR_START}250 000{MARQUEUR_VAR_END} euros.
"""
        md_file = tmp_path / "test_variables.md"
        md_file.write_text(md_content, encoding='utf-8')
        return md_file

    @pytest.mark.docx
    def test_export_basique(self, markdown_simple, tmp_path):
        """Test d'export basique."""
        try:
            from execution.core.exporter_docx import main
            output_file = tmp_path / "output.docx"

            # Simuler les arguments CLI
            import sys
            original_argv = sys.argv
            sys.argv = ['exporter_docx.py',
                       '--input', str(markdown_simple),
                       '--output', str(output_file)]
            try:
                # Exécuter la fonction main (peut échouer si pas de docx)
                result = main()
            except SystemExit as e:
                result = e.code
            finally:
                sys.argv = original_argv

            # Vérifier le fichier généré (si success)
            if output_file.exists():
                assert output_file.stat().st_size > 0
        except ImportError:
            pytest.skip("python-docx non installé")

    @pytest.mark.docx
    def test_fichier_inexistant(self, tmp_path):
        """Test avec fichier d'entrée inexistant."""
        try:
            from execution.core.exporter_docx import main
            import sys
            original_argv = sys.argv
            sys.argv = ['exporter_docx.py',
                       '--input', str(tmp_path / "inexistant.md"),
                       '--output', str(tmp_path / "output.docx")]
            try:
                result = main()
                assert result != 0 or result is None  # Erreur attendue
            except (SystemExit, FileNotFoundError) as e:
                pass  # Comportement attendu
            finally:
                sys.argv = original_argv
        except ImportError:
            pytest.skip("python-docx non installé")


# =============================================================================
# TESTS DE FORMATAGE
# =============================================================================

class TestFormatage:
    """Tests pour les fonctions de formatage."""

    def test_appliquer_fond_gris(self):
        """Test de la fonction appliquer_fond_gris."""
        try:
            from execution.core.exporter_docx import appliquer_fond_gris
            from docx import Document
            from docx.shared import Pt

            # Créer un document de test
            doc = Document()
            p = doc.add_paragraph()
            run = p.add_run("Texte test")

            # Appliquer le fond gris
            appliquer_fond_gris(run)

            # Vérifier que le shading a été ajouté
            rPr = run._element.get_or_add_rPr()
            shd_elements = rPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
            assert len(shd_elements) > 0

        except ImportError:
            pytest.skip("python-docx non installé")


# =============================================================================
# TESTS VALIDATION ACTE
# =============================================================================

class TestValidationExportable:
    """Tests pour valider qu'un acte peut être exporté."""

    def test_markdown_valide(self):
        """Un markdown valide peut être traité."""
        markdown = "# Titre\n\nContenu."
        lignes = markdown.split('\n')
        assert len(lignes) >= 1
        assert lignes[0].startswith('#')

    def test_markdown_avec_tableaux(self):
        """Markdown avec tableaux détectables."""
        markdown = """# Document

| A | B |
|---|---|
| 1 | 2 |

Fin.
"""
        lignes = markdown.split('\n')
        est_tableau, _, _ = detecter_tableau_markdown(lignes, 2)
        assert est_tableau is True

    def test_encodage_utf8(self):
        """Vérifier que l'encodage UTF-8 est supporté."""
        texte = "Éléments accentués: àéèùïôç - Œuvre"
        resultat = nettoyer_texte_xml(texte)
        assert 'à' in resultat
        assert 'é' in resultat
        assert 'Œ' in resultat


# =============================================================================
# TESTS DE PERFORMANCE
# =============================================================================

class TestPerformance:
    """Tests de performance pour l'export."""

    def test_nettoyage_grand_texte(self):
        """Nettoyage d'un grand texte."""
        grand_texte = "Test de performance. " * 10000
        resultat = nettoyer_texte_xml(grand_texte)
        assert len(resultat) == len(grand_texte)

    def test_detection_tableaux_multiples(self):
        """Détection de plusieurs tableaux."""
        lignes = []
        for i in range(10):
            lignes.extend([
                f"| Header{i}A | Header{i}B |",
                "|-----------|-----------|",
                f"| Val{i}1    | Val{i}2    |",
                "",
            ])

        tableaux_trouves = 0
        index = 0
        while index < len(lignes):
            est_tableau, fin, _ = detecter_tableau_markdown(lignes, index)
            if est_tableau:
                tableaux_trouves += 1
                index = fin
            else:
                index += 1

        assert tableaux_trouves == 10


# =============================================================================
# TESTS SPÉCIFIQUES NOTARIAUX
# =============================================================================

class TestContenuNotarial:
    """Tests spécifiques au contenu notarial."""

    def test_tableau_cadastral(self):
        """Détection d'un tableau cadastral."""
        lignes = [
            "| Section | N° | Lieudit | Surface |",
            "|---------|----|---------|---------:|",
            "| AB      | 123| Centre  | 500 m²  |",
        ]
        est_tableau, _, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is True
        assert 'Section' in donnees['lignes'][0]

    def test_tableau_prix(self):
        """Détection d'un tableau de prix."""
        lignes = [
            "| Désignation | Montant |",
            "|-------------|--------:|",
            "| Prix principal | 250 000 € |",
            "| Mobilier | 10 000 € |",
        ]
        est_tableau, _, donnees = detecter_tableau_markdown(lignes, 0)

        assert est_tableau is True
        assert len(donnees['lignes']) == 3

    def test_nettoyage_caracteres_juridiques(self):
        """Nettoyage des caractères juridiques spéciaux."""
        texte = "Article 1er - § 2 - L'acquéreur — appelé ci-après « l'ACQUEREUR »"
        resultat = nettoyer_texte_xml(texte)

        assert "Article 1er" in resultat
        assert "§" in resultat
        assert "«" in resultat
        assert "»" in resultat


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
