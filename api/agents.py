# -*- coding: utf-8 -*-
"""
API Endpoints pour les Agents Opus 4.6

Expose les 6 agents spécialisés au frontend:
- workflow-orchestrator (génération parallèle end-to-end)
- cadastre-enricher (enrichissement cadastre API)
- data-collector-qr (collecte Q&R interactive)
- clause-suggester (suggestions clauses contextuelles)
- post-generation-reviewer (QA final)
- schema-validator, template-auditor (existants)

Endpoints:
    POST /agents/orchestrate          - Génération parallèle complète
    POST /agents/{name}/execute       - Exécuter un agent individuel
    GET  /agents/status               - Status et monitoring agents
    GET  /agents                      - Liste agents disponibles

Version: 2.1.0 - Tier 1 Cost Optimizations (-68% costs)
"""

import sys
import time
import asyncio
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Literal
from datetime import datetime

logger = logging.getLogger(__name__)

from fastapi import APIRouter, HTTPException, BackgroundTasks, Request
from pydantic import BaseModel, Field
from sse_starlette.sse import EventSourceResponse

# Ajouter le projet au path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

# Optimisations coûts Tier 1 (v2.1.0)
try:
    from execution.utils.api_cost_config import (
        get_max_tokens,
        get_model,
        get_timeout,
        should_cache,
        get_cachable_system_prompt,
        estimer_cout
    )
    from execution.utils.validation_deterministe import (
        ValidateurDeterministe,
        detecter_type_acte_rapide,
        valider_quotites
    )
    from execution.gestionnaires.orchestrateur import OrchestratorNotaire
    OPTIMIZATIONS_ENABLED = True
    print("✅ Tier 1 optimizations loaded: Smart model selection, deterministic validation, max tokens")
except ImportError as e:
    OPTIMIZATIONS_ENABLED = False
    print(f"⚠️  Tier 1 optimizations not available: {e}")
    # Fallback functions
    def get_max_tokens(agent_name): return 4096
    def get_model(agent_name, fallback_sonnet=False): return "claude-opus-4-6"
    def get_timeout(agent_name): return 30
    def estimer_cout(*args, **kwargs): return 0.0
    ValidateurDeterministe = None
    detecter_type_acte_rapide = None

router = APIRouter(prefix="/agents", tags=["agents"])


# =============================================================================
# Models
# =============================================================================

class AgentExecuteRequest(BaseModel):
    """Requête pour exécuter un agent individuel."""
    agent_name: Literal[
        "workflow-orchestrator",
        "cadastre-enricher",
        "data-collector-qr",
        "clause-suggester",
        "post-generation-reviewer",
        "schema-validator",
        "template-auditor"
    ]
    prompt: str = Field(..., description="Instructions pour l'agent")
    context: Optional[Dict[str, Any]] = Field(default=None, description="Contexte additionnel")
    timeout_seconds: Optional[int] = Field(default=30, ge=1, le=300, description="Timeout max")


class OrchestrateRequest(BaseModel):
    """Requête pour génération parallèle orchestrée."""
    demande: str = Field(..., description="Demande en langage naturel, ex: 'Promesse Martin→Dupont, 67m² Paris, 450k€'")
    strategy: Literal["parallel", "sequential", "auto"] = Field(
        default="parallel",
        description="Stratégie d'exécution: parallel (3-5x rapide), sequential (debug), auto (adaptatif)"
    )
    mode: Literal["auto", "interactive"] = Field(
        default="auto",
        description="Mode collecte données: auto (prefill max), interactive (questions notaire)"
    )
    options: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Options: skip_clauses, skip_qa, verbose, etc."
    )


class AgentStatus(BaseModel):
    """Status d'un agent."""
    name: str
    status: Literal["available", "busy", "error", "unknown"]
    last_execution: Optional[datetime] = None
    avg_duration_ms: Optional[float] = None
    success_rate: Optional[float] = None


class AgentExecuteResponse(BaseModel):
    """Réponse d'exécution d'agent."""
    agent_name: str
    status: Literal["success", "error", "timeout"]
    duration_ms: float
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None

    # Tier 1: Cost tracking (v2.1.0)
    cost_tracking: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Model used, tokens, estimated cost"
    )


class OrchestrateResponse(BaseModel):
    """Réponse d'orchestration complète."""
    workflow_id: str
    status: Literal["success", "error", "blocked"]
    strategy_used: str
    duration_total_ms: float
    speedup_vs_sequential: Optional[float] = None

    agents_executed: List[Dict[str, Any]]
    data_quality: Dict[str, Any]
    output: Optional[Dict[str, Any]] = None

    errors: List[str] = []
    warnings: List[str] = []

    # Tier 1: Cost optimization metrics (v2.1.0)
    cost_summary: Optional[Dict[str, Any]] = Field(
        default=None,
        description="Total cost, model distribution, savings vs baseline"
    )


# =============================================================================
# Agents Registry (mapping nom → fonction d'exécution)
# =============================================================================

# Import conditionnel des agents (certains peuvent être indisponibles)
AGENTS_AVAILABLE = {
    "schema-validator": True,
    "template-auditor": True,
    "security-reviewer": True,
}

# Nouveaux agents Opus 4.6 (nécessitent Claude Code ou implémentation Python)
# Pour Modal, on créera des wrappers Python qui appellent les scripts
AGENTS_OPUS_46 = {
    "cadastre-enricher": "execution.services.cadastre_service",
    "data-collector-qr": "execution.agent_autonome",  # CollecteurInteractif
    "clause-suggester": "execution.utils.suggerer_clauses",  # À créer
    "post-generation-reviewer": "execution.utils.reviewer_qa",  # À créer
    "workflow-orchestrator": "execution.orchestrateur_notaire",  # Existing
}


async def execute_cadastre_enricher(prompt: str, context: dict) -> dict:
    """
    Execute cadastre-enricher agent.

    Input context:
        - adresse: str (e.g., "12 rue de la Paix, 75002 Paris")
        - data: dict (données dossier à enrichir)

    Returns:
        - enriched_data: dict (bien enrichi avec cadastre)
        - source: "api_cadastre_gouv"
        - duration_ms: float
    """
    from execution.services.cadastre_service import CadastreService

    start = time.time()
    service = CadastreService()

    # Extract address from prompt or context
    adresse = context.get("adresse") if context else None
    donnees = context.get("data") if context else {}

    if not adresse and donnees.get("bien", {}).get("adresse"):
        adresse_obj = donnees["bien"]["adresse"]
        adresse = f"{adresse_obj.get('numero', '')} {adresse_obj.get('voie', '')}, {adresse_obj.get('code_postal', '')} {adresse_obj.get('ville', '')}"

    if not adresse:
        return {
            "status": "error",
            "error": "Adresse manquante dans prompt ou context",
            "duration_ms": (time.time() - start) * 1000
        }

    # Enrich data — offload blocking I/O (HTTP calls to BAN/IGN APIs)
    result = await asyncio.to_thread(service.enrichir_cadastre, donnees)

    return {
        "status": "success",
        "enriched": result.get("enriched", False),
        "fields_added": result.get("fields_added", []),
        "enriched_data": result.get("data", donnees),
        "warnings": result.get("warnings", []),
        "duration_ms": (time.time() - start) * 1000
    }


async def execute_data_collector_qr(prompt: str, context: dict) -> dict:
    """
    Execute data-collector-qr agent.

    Input context:
        - type_acte: str
        - donnees_existantes: dict (optional, for prefill)
        - mode: "cli" | "prefill_only"

    Returns:
        - collected: bool
        - data: dict
        - taux_completion: float
        - questions_posees: int
    """
    from execution.agent_autonome import CollecteurInteractif

    start = time.time()

    type_acte = context.get("type_acte", "promesse_vente")
    donnees_existantes = context.get("donnees_existantes", {})
    mode = context.get("mode", "prefill_only")  # Default: auto (pas de questions)

    collecteur = CollecteurInteractif(type_acte=type_acte)

    # En mode API, on utilise toujours prefill_only (pas d'interaction terminal)
    # Le frontend gérera les questions manquantes via son propre workflow
    # Offload blocking I/O (schema loading, file reads)
    resultat = await asyncio.to_thread(
        collecteur.collecter,
        donnees_initiales=donnees_existantes,
        mode="prefill_only"
    )

    return {
        "status": "success",
        "collected": True,
        "data": resultat.get("data", {}),
        "taux_completion": resultat.get("taux_completion", 0),
        "taux_prefill": resultat.get("taux_prefill", 0),
        "questions_posees": 0,  # Mode auto
        "duration_ms": (time.time() - start) * 1000
    }


async def execute_clause_suggester(prompt: str, context: dict) -> dict:
    """
    Execute clause-suggester agent via le vrai moteur suggerer_clauses.py.

    Input context:
        - metadata: dict (type_acte, prix, pret, etc.)
        - donnees: dict (donnees completes de l'acte)

    Returns:
        - suggestions: List[dict]
        - total_suggestions: int
        - critiques: int, recommandees: int, optionnelles: int
    """
    start = time.time()

    metadata = context.get("metadata", {})
    donnees = context.get("donnees", {})
    type_acte = metadata.get("type_acte", "promesse_vente")

    try:
        from execution.utils.suggerer_clauses import suggerer_clauses

        raw_suggestions = suggerer_clauses(donnees, type_acte, verbose=False)

        # Mapper SuggestionClause vers le format API
        priorite_map = {"obligatoire": 1, "recommandee": 2, "optionnelle": 3}
        score_map = {"obligatoire": 95, "recommandee": 65, "optionnelle": 35}

        suggestions = [
            {
                "id": s.clause_id,
                "nom": s.nom,
                "priorite": priorite_map.get(s.priorite, 3),
                "score": score_map.get(s.priorite, 35),
                "justification": s.raison,
                "variables_disponibles": len(s.variables_manquantes) == 0,
                "section_insertion": s.categorie,
            }
            for s in raw_suggestions
        ]
    except Exception as e:
        logger.warning(f"Clause suggester error: {e}")
        suggestions = []
        return {
            "status": "error",
            "error": str(e),
            "suggestions": [],
            "total_suggestions": 0,
            "critiques": 0, "recommandees": 0, "optionnelles": 0,
            "duration_ms": (time.time() - start) * 1000
        }

    return {
        "status": "success",
        "suggestions": suggestions,
        "total_suggestions": len(suggestions),
        "critiques": sum(1 for s in suggestions if s["priorite"] == 1),
        "recommandees": sum(1 for s in suggestions if s["priorite"] == 2),
        "optionnelles": sum(1 for s in suggestions if s["priorite"] == 3),
        "duration_ms": (time.time() - start) * 1000
    }


async def execute_post_generation_reviewer(prompt: str, context: dict) -> dict:
    """
    Execute post-generation-reviewer agent — QA deterministe.

    Verifie le DOCX genere sur 7 dimensions sans necesiter les trames originales.

    Input context:
        - docx_path: str
        - donnees: dict (donnees utilisees pour generation)

    Returns:
        - status: "PASS" | "WARNING" | "BLOCKED"
        - qa_score: int (0-100)
        - issues: List[dict]
    """
    start = time.time()

    docx_path = context.get("docx_path")
    donnees = context.get("donnees", {})

    if not docx_path:
        return {
            "status": "error",
            "error": "docx_path manquant",
            "duration_ms": (time.time() - start) * 1000
        }

    issues = []
    qa_score = 100

    # Dimension 1: Fichier DOCX valide et lisible
    try:
        from docx import Document as DocxDocument
        if Path(docx_path).exists():
            doc = DocxDocument(docx_path)
            paragraphs = [p.text for p in doc.paragraphs]
            full_text = "\n".join(paragraphs)
        else:
            issues.append({"severite": "CRITIQUE", "dimension": "fichier", "message": f"Fichier introuvable: {docx_path}"})
            qa_score -= 30
            doc = None
            full_text = ""
    except Exception as e:
        issues.append({"severite": "CRITIQUE", "dimension": "fichier", "message": f"DOCX invalide: {e}"})
        qa_score -= 30
        doc = None
        full_text = ""

    # Dimension 2: Variables Jinja2 non-resolues
    if full_text:
        import re
        unresolved = re.findall(r'\{\{[^}]+\}\}', full_text)
        if unresolved:
            issues.append({
                "severite": "CRITIQUE",
                "dimension": "variables",
                "message": f"{len(unresolved)} variable(s) Jinja2 non-resolue(s): {', '.join(unresolved[:5])}"
            })
            qa_score -= min(20, len(unresolved) * 5)

    # Dimension 3: Quotites
    for field_name in ["quotites_vendues", "quotites_acquises"]:
        quotites = donnees.get(field_name, [])
        if quotites:
            total = sum(
                q.get("valeur", 0) / max(q.get("base", 1), 1)
                for q in quotites
            )
            if abs(total - 1.0) > 0.001:
                issues.append({
                    "severite": "CRITIQUE",
                    "dimension": "quotites",
                    "message": f"{field_name} = {total*100:.1f}% (attendu: 100%)"
                })
                qa_score -= 15

    # Dimension 4: Prix
    prix = donnees.get("prix", {})
    montant = prix.get("montant", 0)
    if montant <= 0:
        issues.append({
            "severite": "AVERTISSEMENT",
            "dimension": "prix",
            "message": "Prix non renseigne ou <= 0"
        })
        qa_score -= 5

    # Dimension 5: Sections obligatoires (headings dans le DOCX)
    if doc:
        headings = [p.text.strip().upper() for p in doc.paragraphs if p.style and "heading" in (p.style.name or "").lower()]
        sections_attendues = ["COMPARUTION", "DESIGNATION", "PRIX"]
        for section in sections_attendues:
            if not any(section in h for h in headings):
                issues.append({
                    "severite": "AVERTISSEMENT",
                    "dimension": "sections",
                    "message": f"Section '{section}' absente des headings"
                })
                qa_score -= 3

    # Dimension 6: Longueur minimale du document
    if doc and len(doc.paragraphs) < 10:
        issues.append({
            "severite": "AVERTISSEMENT",
            "dimension": "completude",
            "message": f"Document tres court ({len(doc.paragraphs)} paragraphes)"
        })
        qa_score -= 10

    # Dimension 7: Parties presentes
    promettants = donnees.get("promettants", donnees.get("vendeurs", []))
    beneficiaires = donnees.get("beneficiaires", donnees.get("acquereurs", []))
    if not promettants:
        issues.append({"severite": "AVERTISSEMENT", "dimension": "parties", "message": "Aucun vendeur/promettant"})
        qa_score -= 5
    if not beneficiaires:
        issues.append({"severite": "AVERTISSEMENT", "dimension": "parties", "message": "Aucun acquereur/beneficiaire"})
        qa_score -= 5

    qa_score = max(0, qa_score)

    status = "BLOCKED" if any(i["severite"] == "CRITIQUE" for i in issues) else \
             "WARNING" if issues else \
             "PASS"

    return {
        "status": status,
        "qa_score": qa_score,
        "issues": issues,
        "dimensions_checked": [
            "fichier", "variables", "quotites", "prix",
            "sections", "completude", "parties"
        ],
        "duration_ms": (time.time() - start) * 1000
    }


# Agent executors registry
AGENT_EXECUTORS = {
    "cadastre-enricher": execute_cadastre_enricher,
    "data-collector-qr": execute_data_collector_qr,
    "clause-suggester": execute_clause_suggester,
    "post-generation-reviewer": execute_post_generation_reviewer,
}


# =============================================================================
# Cost Tracking Helper (Tier 1 v2.1.0)
# =============================================================================

def track_agent_cost(
    agent_name: str,
    model_used: str,
    tokens_input: int = 0,
    tokens_output: int = 0,
    tokens_cached: int = 0,
    duration_ms: float = 0
) -> Dict[str, Any]:
    """
    Track cost for an agent execution.

    Returns dict with model, tokens, cost, and logs to Supabase if available.
    """
    if not OPTIMIZATIONS_ENABLED:
        return {}

    cost_usd = estimer_cout(
        model=model_used,
        tokens_input=tokens_input,
        tokens_output=tokens_output,
        tokens_cached=tokens_cached
    )

    tracking = {
        "agent_name": agent_name,
        "model_used": model_used,
        "tokens_input": tokens_input,
        "tokens_output": tokens_output,
        "tokens_cached": tokens_cached,
        "cost_usd": round(cost_usd, 6),
        "duration_ms": round(duration_ms, 2)
    }

    # TODO: Log to Supabase api_costs_tracking table (after migration)
    # from execution.database.supabase_client import supabase
    # supabase.table("api_costs_tracking").insert(tracking).execute()

    return tracking


# =============================================================================
# Endpoints
# =============================================================================

@router.get("")
async def list_agents():
    """
    Liste tous les agents disponibles avec leur status.

    v2.1.0: Inclut config Tier 1 (max_tokens, model, timeout)
    """
    agents = []

    for name in AGENTS_AVAILABLE:
        agent_config = {
            "name": name,
            "type": "existant",
            "model": "sonnet" if name == "template-auditor" else "haiku",
            "status": "available"
        }

        # Add Tier 1 config if available
        if OPTIMIZATIONS_ENABLED:
            agent_config["tier1_config"] = {
                "max_tokens": get_max_tokens(name),
                "timeout_seconds": get_timeout(name),
                "caching_enabled": should_cache(name)
            }

        agents.append(agent_config)

    for name, module in AGENTS_OPUS_46.items():
        # Get default model
        default_model = "opus" if "orchestrator" in name or "suggester" in name else "sonnet"

        agent_config = {
            "name": name,
            "type": "opus_4.6",
            "model": default_model,
            "model_fallback": "sonnet" if default_model == "opus" else None,
            "status": "available" if name in AGENT_EXECUTORS else "pending_implementation"
        }

        # Add Tier 1 config if available
        if OPTIMIZATIONS_ENABLED:
            agent_config["tier1_config"] = {
                "max_tokens": get_max_tokens(name),
                "timeout_seconds": get_timeout(name),
                "caching_enabled": should_cache(name),
                "smart_model_selection": name == "workflow-orchestrator"
            }

        agents.append(agent_config)

    response = {
        "agents": agents,
        "total": len(agents),
        "available": sum(1 for a in agents if a["status"] == "available")
    }

    # Add Tier 1 optimization status
    if OPTIMIZATIONS_ENABLED:
        response["tier1_optimizations"] = {
            "enabled": True,
            "features": [
                "smart_model_selection",
                "deterministic_validation",
                "max_tokens_limits",
                "cost_tracking"
            ],
            "expected_savings": "-68% vs baseline"
        }
    else:
        response["tier1_optimizations"] = {
            "enabled": False,
            "message": "Install execution.utils modules to enable"
        }

    return response


@router.post("/{agent_name}/execute", response_model=AgentExecuteResponse)
async def execute_agent(agent_name: str, request: AgentExecuteRequest):
    """
    Exécute un agent individuel.

    Utile pour:
    - Tests unitaires d'agents
    - Workflows personnalisés
    - Debugging

    v2.1.0: Intègre optimisations Tier 1 (max_tokens, cost tracking)
    """
    start = time.time()

    if agent_name not in AGENT_EXECUTORS:
        raise HTTPException(
            status_code=404,
            detail=f"Agent '{agent_name}' non trouvé ou pas encore implémenté. "
                   f"Agents disponibles: {list(AGENT_EXECUTORS.keys())}"
        )

    executor = AGENT_EXECUTORS[agent_name]

    # Tier 1: Get optimal timeout from config
    timeout = request.timeout_seconds or (get_timeout(agent_name) if OPTIMIZATIONS_ENABLED else 30)

    try:
        result = await asyncio.wait_for(
            executor(request.prompt, request.context or {}),
            timeout=timeout
        )

        duration_ms = (time.time() - start) * 1000

        # Tier 1: Track cost if result contains token usage
        cost_tracking = None
        if OPTIMIZATIONS_ENABLED and isinstance(result, dict):
            tokens_input = result.get("tokens_input", 0)
            tokens_output = result.get("tokens_output", 0)
            model_used = result.get("model_used", get_model(agent_name))

            if tokens_input or tokens_output:
                cost_tracking = track_agent_cost(
                    agent_name=agent_name,
                    model_used=model_used,
                    tokens_input=tokens_input,
                    tokens_output=tokens_output,
                    duration_ms=duration_ms
                )

        return AgentExecuteResponse(
            agent_name=agent_name,
            status="success",
            duration_ms=duration_ms,
            result=result,
            cost_tracking=cost_tracking
        )

    except asyncio.TimeoutError:
        return AgentExecuteResponse(
            agent_name=agent_name,
            status="timeout",
            duration_ms=(time.time() - start) * 1000,
            error=f"Agent timeout après {timeout}s"
        )

    except Exception as e:
        return AgentExecuteResponse(
            agent_name=agent_name,
            status="error",
            duration_ms=(time.time() - start) * 1000,
            error=str(e)
        )


@router.post("/orchestrate", response_model=OrchestrateResponse)
async def orchestrate_generation(request: OrchestrateRequest):
    """
    Génération parallèle orchestrée (Opus 4.6).

    Workflow:
    1. Parse demande NL
    2. Décide stratégie (parallel/sequential)
    3. Lance agents (cadastre, collector, auditor en parallèle)
    4. Validation
    5. Assemblage + suggestions clauses
    6. Export DOCX
    7. QA final
    8. Return rapport + fichier

    Speedup attendu: 2.5-3x vs sequential
    """
    from execution.agent_autonome import AgentNotaire, ParseurDemandeNL
    from execution.workflow_rapide import WorkflowRapide

    start_workflow = time.time()
    workflow_id = f"wf-{datetime.now().strftime('%Y%m%d-%H%M%S')}"

    # Phase 1: Parse demande
    parseur = ParseurDemandeNL()
    intent = parseur.analyser(request.demande)

    if intent.intention.value not in ["creer", "generer"]:
        raise HTTPException(
            status_code=400,
            detail=f"Cette API ne gère que la création/génération. Intention détectée: {intent.intention.value}"
        )

    # Phase 2: Stratégie
    strategy = request.strategy
    if strategy == "auto":
        # Décider automatiquement selon complexité
        strategy = "parallel" if intent.confiance > 0.7 else "sequential"

    agents_executed = []
    errors = []
    warnings = []

    try:
        # Phase 3: Exécution agents
        # Build tasks and labels in lockstep to avoid index mismatch
        tasks = []
        task_names = []

        if strategy == "parallel":
            # Cadastre enricher (only if address available)
            if intent.bien:
                tasks.append(execute_cadastre_enricher(
                    prompt=request.demande,
                    context={"data": {"bien": intent.bien}}
                ))
                task_names.append("cadastre-enricher")

            # Data collector (prefill mode)
            tasks.append(execute_data_collector_qr(
                prompt=request.demande,
                context={
                    "type_acte": intent.type_acte.value,
                    "donnees_existantes": {
                        "promettants": [intent.vendeur] if intent.vendeur else [],
                        "beneficiaires": [intent.acquereur] if intent.acquereur else [],
                        "bien": intent.bien or {},
                        "prix": intent.prix or {}
                    },
                    "mode": request.mode
                }
            ))
            task_names.append("data-collector-qr")

            # Execute in parallel
            results = await asyncio.gather(*tasks, return_exceptions=True)

            for agent_name, result in zip(task_names, results):
                if isinstance(result, Exception):
                    errors.append(f"{agent_name}: {str(result)}")
                    agents_executed.append({
                        "name": agent_name,
                        "status": "error",
                        "duration_ms": 0,
                        "error": str(result)
                    })
                else:
                    agents_executed.append({
                        "name": agent_name,
                        "status": "success",
                        "duration_ms": result.get("duration_ms", 0),
                        "result": result
                    })

        else:
            # Sequential fallback — run agents one by one
            if intent.bien:
                try:
                    cadastre_result = await execute_cadastre_enricher(
                        prompt=request.demande,
                        context={"data": {"bien": intent.bien}}
                    )
                    agents_executed.append({"name": "cadastre-enricher", "status": "success", "duration_ms": cadastre_result.get("duration_ms", 0), "result": cadastre_result})
                except Exception as e:
                    errors.append(f"cadastre-enricher: {e}")
                    agents_executed.append({"name": "cadastre-enricher", "status": "error", "duration_ms": 0, "error": str(e)})

            try:
                qr_result = await execute_data_collector_qr(
                    prompt=request.demande,
                    context={"type_acte": intent.type_acte.value, "donnees_existantes": {"promettants": [intent.vendeur] if intent.vendeur else [], "beneficiaires": [intent.acquereur] if intent.acquereur else [], "bien": intent.bien or {}, "prix": intent.prix or {}}, "mode": request.mode}
                )
                agents_executed.append({"name": "data-collector-qr", "status": "success", "duration_ms": qr_result.get("duration_ms", 0), "result": qr_result})
            except Exception as e:
                errors.append(f"data-collector-qr: {e}")
                agents_executed.append({"name": "data-collector-qr", "status": "error", "duration_ms": 0, "error": str(e)})

        # Phase 4-7: Pipeline reel — validation, assemblage, export, QA
        # Merge des donnees collectees par les agents paralleles
        merged_donnees = {
            "promettants": [intent.vendeur] if intent.vendeur else [],
            "beneficiaires": [intent.acquereur] if intent.acquereur else [],
            "bien": intent.bien or {},
            "prix": intent.prix or {},
        }

        # Integrer les resultats des agents paralleles
        for agent_info in agents_executed:
            if agent_info.get("status") != "success" or not agent_info.get("result"):
                continue
            result = agent_info["result"]
            if agent_info["name"] == "cadastre-enricher" and result.get("enriched_data"):
                # Merge cadastre data into bien
                enriched = result["enriched_data"]
                if isinstance(enriched, dict):
                    bien_enrichi = enriched.get("bien", enriched)
                    merged_donnees["bien"].update(bien_enrichi)
            elif agent_info["name"] == "data-collector-qr" and result.get("data"):
                # Merge Q&R collected data (without overwriting existing data)
                for key, val in result["data"].items():
                    if key not in merged_donnees or not merged_donnees[key]:
                        merged_donnees[key] = val

        # Phase 4-5: Generation reelle via le pipeline existant
        output_info = None
        type_acte_str = intent.type_acte.value if hasattr(intent.type_acte, 'value') else str(intent.type_acte)

        try:
            if type_acte_str in ("promesse_vente", "promesse"):
                from execution.gestionnaires.gestionnaire_promesses import GestionnairePromesses
                gestionnaire = GestionnairePromesses()
                gen_result = gestionnaire.generer(merged_donnees, force=True)
                if gen_result.succes and gen_result.fichier_docx:
                    file_path = gen_result.fichier_docx
                    import os
                    output_info = {
                        "file_path": file_path,
                        "file_size_kb": round(os.path.getsize(file_path) / 1024, 1) if os.path.exists(file_path) else 0,
                        "pages": 24,  # Estimation standard promesse
                    }
                    agents_executed.append({
                        "name": "workflow-orchestrator",
                        "status": "success",
                        "duration_ms": gen_result.duree_generation * 1000,
                    })
                else:
                    errors.extend(gen_result.erreurs)
                    agents_executed.append({
                        "name": "workflow-orchestrator",
                        "status": "error",
                        "duration_ms": gen_result.duree_generation * 1000,
                        "error": "; ".join(gen_result.erreurs),
                    })
                warnings.extend(gen_result.warnings)
            else:
                from execution.gestionnaires.orchestrateur import OrchestratorNotaire
                orch = OrchestratorNotaire()
                wf_result = orch.generer_acte_complet(type_acte_str, merged_donnees)
                if wf_result.statut == "succes" and wf_result.fichiers_generes:
                    file_path = wf_result.fichiers_generes[0]
                    import os
                    output_info = {
                        "file_path": file_path,
                        "file_size_kb": round(os.path.getsize(file_path) / 1024, 1) if os.path.exists(file_path) else 0,
                        "pages": 24,
                    }
                    agents_executed.append({
                        "name": "workflow-orchestrator",
                        "status": "success",
                        "duration_ms": wf_result.duree_totale_ms,
                    })
                else:
                    errors.extend(wf_result.erreurs)
                    agents_executed.append({
                        "name": "workflow-orchestrator",
                        "status": "error",
                        "duration_ms": wf_result.duree_totale_ms,
                        "error": "; ".join(wf_result.erreurs),
                    })
                warnings.extend(wf_result.alertes)
        except Exception as gen_err:
            logger.error(f"Erreur generation pipeline: {gen_err}", exc_info=True)
            errors.append(f"generation: {str(gen_err)}")
            agents_executed.append({
                "name": "workflow-orchestrator",
                "status": "error",
                "duration_ms": 0,
                "error": str(gen_err),
            })

        # Phase 6: Suggestions de clauses (reel)
        if output_info and not errors:
            try:
                clause_result = await execute_clause_suggester(
                    prompt=request.demande,
                    context={"metadata": {"type_acte": type_acte_str, "prix": merged_donnees.get("prix", {}), "pret": merged_donnees.get("pret", {})}, "donnees": merged_donnees}
                )
                clause_status = clause_result.get("status", "success")
                agents_executed.append({
                    "name": "clause-suggester",
                    "status": clause_status,
                    "duration_ms": clause_result.get("duration_ms", 0),
                    "result": clause_result,
                })
            except Exception as clause_err:
                logger.warning(f"Clause suggester failed (non-blocking): {clause_err}")

        # Phase 7: QA review (reel)
        if output_info and not errors:
            try:
                qa_result = await execute_post_generation_reviewer(
                    prompt="QA review",
                    context={"docx_path": output_info["file_path"], "donnees": merged_donnees}
                )
                agents_executed.append({
                    "name": "post-generation-reviewer",
                    "status": "success",
                    "duration_ms": qa_result.get("duration_ms", 0),
                    "result": qa_result,
                })
                if qa_result.get("status") == "BLOCKED":
                    errors.append(f"QA blocked: {'; '.join(i.get('message', '') for i in qa_result.get('issues', []))}")
            except Exception as qa_err:
                logger.warning(f"Post-gen reviewer failed (non-blocking): {qa_err}")

        duration_total_ms = (time.time() - start_workflow) * 1000

        return OrchestrateResponse(
            workflow_id=workflow_id,
            status="success" if not errors else "error",
            strategy_used=strategy,
            duration_total_ms=duration_total_ms,
            speedup_vs_sequential=round(max(sum(a.get("duration_ms", 0) for a in agents_executed if a.get("name") in {"cadastre-enricher", "data-collector-qr"}), 1) / max(duration_total_ms, 1), 1) if strategy == "parallel" else None,
            agents_executed=agents_executed,
            data_quality={
                "completion": 100 if output_info else 0,
                "validation_errors": len(errors),
                "warnings": len(warnings)
            },
            output=output_info,
            errors=errors,
            warnings=warnings
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erreur orchestration: {str(e)}")


@router.get("/status")
async def get_agents_status():
    """
    Status et monitoring des agents.

    Returns:
    - Nombre d'agents disponibles
    - Dernières exécutions
    - Temps moyens
    - Taux de succès
    """
    # TODO: Persister stats dans Supabase

    return {
        "agents_available": len(AGENT_EXECUTORS),
        "agents_total": len(AGENTS_AVAILABLE) + len(AGENTS_OPUS_46),
        "status": "operational",
        "last_check": datetime.now().isoformat(),
        "agents": [
            {
                "name": name,
                "status": "available",
                "last_execution": None,
                "avg_duration_ms": None,
                "success_rate": None
            }
            for name in AGENT_EXECUTORS
        ]
    }
