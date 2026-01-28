#!/usr/bin/env python3
"""
Anonymisation PRÉCISE de documents DOCX notariaux.
Version 2 - Détection contextuelle uniquement.

Détecte les noms/prénoms UNIQUEMENT dans des contextes identifiables:
- Après civilités (Monsieur, Madame, M., Mme)
- Dans les blocs d'identité (né(e) le, demeurant)
- Noms en majuscules après "époux/épouse"
- Noms de notaires après "Maître"

Usage:
    python execution/anonymiser_docx_v2.py docs_original/Promesse*.docx --batch
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional
from dataclasses import dataclass, field
from collections import defaultdict

# Encodage UTF-8 pour Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("❌ python-docx requis: pip install python-docx")
    sys.exit(1)


# =============================================================================
# MOTS À NE JAMAIS ANONYMISER (termes juridiques, villes, etc.)
# =============================================================================

MOTS_EXCLUS = {
    # Termes juridiques courants en majuscules
    "PROMETTANT", "PROMETTANTS", "BENEFICIAIRE", "BENEFICIAIRES", "BENFICIAIRE",
    "VENDEUR", "VENDEURS", "ACQUEREUR", "ACQUEREURS", "ACHETEUR", "ACHETEURS",
    "BIEN", "BIENS", "PRIX", "CONDITIONS", "SUSPENSIVES", "VENTE", "PROMESSE",
    "PROPRIETE", "PROPRIÉTÉ", "COPROPRIETE", "COPROPRIÉTÉ", "IMMEUBLE",
    "LOT", "LOTS", "CAVE", "PARKING", "GARAGE", "APPARTEMENT", "MAISON",
    "TERRAIN", "PARCELLE", "SECTION", "CADASTRE", "COMMUNE", "DEPARTEMENT",
    "NOTAIRE", "NOTAIRES", "OFFICE", "ETUDE", "ÉTUDE", "MINUTE", "MINUTES",
    "ACTE", "ACTES", "AUTHENTIQUE", "SIGNATURE", "COMPARANT", "COMPARANTS",
    "PARTIES", "PARTIE", "ARTICLE", "ARTICLES", "CODE", "CIVIL", "LOI",
    "DECRET", "DÉCRET", "ARRETE", "ARRÊTÉ", "TRIBUNAL", "COUR", "JUSTICE",
    "FRANCE", "FRANCAISE", "FRANÇAISE", "REPUBLIQUE", "RÉPUBLIQUE",
    "DESIGNATION", "DÉSIGNATION", "ORIGINE", "URBANISME", "HYPOTHEQUE",
    "HYPOTHÈQUE", "INSCRIPTION", "RADIATION", "MAINLEVEE", "MAINLEVÉE",
    "SERVITUDE", "SERVITUDES", "CHARGES", "ANNEXES", "DIAGNOSTIC",
    "DIAGNOSTICS", "AMIANTE", "PLOMB", "TERMITES", "GAZ", "ELECTRICITE",
    "ÉLECTRICITÉ", "DPE", "CARREZ", "SUPERFICIE", "SURFACE",
    "SYNDIC", "SYNDICAT", "ASSEMBLEE", "ASSEMBLÉE", "GENERALE", "GÉNÉRALE",
    "REGLEMENT", "RÈGLEMENT", "DIVISION", "DESCRIPTIF", "ETAT", "ÉTAT",
    "TANTIEME", "TANTIEMES", "TANTIÈMES", "QUOTE", "PART", "PARTS",
    "INDIVISION", "INDIVISAIRE", "USUFRUIT", "USUFRUITIER", "NUPROPRIETE",
    "DONATION", "SUCCESSION", "HERITAGE", "HÉRITAGE", "TESTAMENT",
    "MARIAGE", "PACS", "DIVORCE", "VEUF", "VEUVE", "CELIBATAIRE",
    "REGIME", "RÉGIME", "MATRIMONIAL", "COMMUNAUTE", "COMMUNAUTÉ",
    "SEPARATION", "SÉPARATION", "ACQUETS", "ACQUÊTS", "UNIVERSEL",
    "BANQUE", "PRET", "PRÊT", "CREDIT", "CRÉDIT", "EMPRUNT", "FINANCEMENT",
    "CAUTION", "GARANTIE", "HYPOTHECAIRE", "HYPOTHÉCAIRE", "PRIVILEGIE",
    "EUR", "EUROS", "EURO", "TVA", "TTC", "HT",
    "SCI", "SARL", "SAS", "SA", "SASU", "EURL", "SNC", "GFA",
    "RCS", "SIRET", "SIREN", "APE", "NAF",
    "NORD", "SUD", "EST", "OUEST", "RUE", "AVENUE", "BOULEVARD", "PLACE",
    "IMPASSE", "ALLEE", "ALLÉE", "CHEMIN", "ROUTE", "VOIE", "PASSAGE",
    # Mois
    "JANVIER", "FEVRIER", "FÉVRIER", "MARS", "AVRIL", "MAI", "JUIN",
    "JUILLET", "AOUT", "AOÛT", "SEPTEMBRE", "OCTOBRE", "NOVEMBRE", "DECEMBRE", "DÉCEMBRE",
    # Jours
    "LUNDI", "MARDI", "MERCREDI", "JEUDI", "VENDREDI", "SAMEDI", "DIMANCHE",
    # Chiffres en lettres
    "UN", "DEUX", "TROIS", "QUATRE", "CINQ", "SIX", "SEPT", "HUIT", "NEUF", "DIX",
    "ONZE", "DOUZE", "TREIZE", "QUATORZE", "QUINZE", "SEIZE", "VINGT", "TRENTE",
    "QUARANTE", "CINQUANTE", "SOIXANTE", "CENT", "CENTS", "MILLE", "MILLION", "MILLIONS",
    # Autres termes juridiques
    "EXPOSE", "EXPOSÉ", "CONVENTION", "STIPULATION", "STIPULATIONS",
    "CLAUSE", "CLAUSES", "CONDITION", "DELAI", "DÉLAI", "DATE", "JOUR", "JOURS",
    "MOIS", "AN", "ANS", "ANNEE", "ANNÉE", "ANNEES", "ANNÉES",
    "OBLIGATION", "OBLIGATIONS", "DROIT", "DROITS", "POUVOIR", "POUVOIRS",
    "PROCURATION", "MANDAT", "MANDATAIRE", "REPRESENTANT", "REPRÉSENTANT",
    "PRESENT", "PRÉSENT", "PRESENTS", "PRÉSENTS", "SOUSSIGNE", "SOUSSIGNÉ",
    "DECLARE", "DÉCLARE", "RECONNAIT", "RECONNAÎT", "CERTIFIE", "ATTESTE",
    "DONT", "SOIT", "SAVOIR", "SUIVANT", "SUIVANTE", "SUIVANTS", "SUIVANTES",
    # Termes techniques
    "BÂTIMENT", "BATIMENT", "ETAGE", "ÉTAGE", "REZ", "CHAUSSEE", "CHAUSSÉE",
    "ESCALIER", "ASCENSEUR", "ENTREE", "ENTRÉE", "SORTIE", "COULOIR",
    "PIECE", "PIÈCE", "PIECES", "PIÈCES", "CHAMBRE", "SALON", "CUISINE",
    "SALLE", "BAIN", "DOUCHE", "WC", "TOILETTES", "BALCON", "TERRASSE",
    "JARDIN", "COUR", "LOCAL", "LOCAUX", "COMMERCIAL", "COMMERCIAUX",
    # Abréviations
    "M", "MM", "MME", "MMES", "MR", "MLLE", "DR", "ME", "STE", "CIE",
    # Villes principales
    "PARIS", "LYON", "MARSEILLE", "TOULOUSE", "NICE", "NANTES", "STRASBOURG",
    "MONTPELLIER", "BORDEAUX", "LILLE", "RENNES", "REIMS", "TOULON",
    "ASNIERES", "ASNIÈRES", "SEINE", "HAUTS",
    # =========================================================================
    # MOTS FRANÇAIS COURANTS (faux positifs détectés)
    # =========================================================================
    # Articles, prépositions, conjonctions
    "LE", "LA", "LES", "UN", "UNE", "DES", "DU", "DE", "AU", "AUX",
    "AVEC", "SANS", "POUR", "PAR", "SUR", "SOUS", "DANS", "ENTRE",
    "ET", "OU", "NI", "MAIS", "DONC", "OR", "CAR",
    "CE", "CET", "CETTE", "CES", "MON", "TON", "SON", "MA", "TA", "SA",
    "QUI", "QUE", "QUOI", "DONT", "LEQUEL", "LAQUELLE", "LESQUELS", "LESQUELLES",
    # Verbes et formes verbales courants
    "EST", "SONT", "SERA", "SERONT", "ETAIT", "ÉTAIT", "ETAIENT", "ÉTAIENT",
    "A", "ONT", "AURA", "AURONT", "AVAIT", "AVAIENT",
    "FAIT", "FAIRE", "FONT", "FERA", "FERONT",
    "PEUT", "PEUVENT", "POURRA", "POURRONT", "POUVAIT", "POUVAIENT",
    "DOIT", "DOIVENT", "DEVRA", "DEVRONT", "DEVAIT", "DEVAIENT",
    "VU", "VUS", "VUE", "VUES", "VOIR", "VOIT", "VOIENT",
    "DIT", "DITE", "DITS", "DITES", "DIRE",
    "PRIS", "PRISE", "PRISES", "PRENDRE",
    "MIS", "MISE", "MISES", "METTRE",
    # Adverbes et locutions
    "AINSI", "AUSSI", "ALORS", "APRÈS", "AVANT", "BIEN", "COMME", "COMMENT",
    "DEPUIS", "ENCORE", "ENFIN", "ICI", "JAMAIS", "LORS", "MAINTENANT",
    "MOINS", "PLUS", "POURQUOI", "QUAND", "RIEN", "SEULEMENT", "TOUJOURS",
    "TOUT", "TOUTE", "TOUS", "TOUTES", "TRÈS", "TROP",
    # Adjectifs courants
    "AUTRE", "AUTRES", "CERTAIN", "CERTAINE", "CERTAINS", "CERTAINES",
    "CHAQUE", "DERNIER", "DERNIÈRE", "DERNIERS", "DERNIÈRES",
    "GRAND", "GRANDE", "GRANDS", "GRANDES", "PETIT", "PETITE", "PETITS", "PETITES",
    "PREMIER", "PREMIÈRE", "PREMIERS", "PREMIÈRES",
    "SEUL", "SEULE", "SEULS", "SEULES", "TEL", "TELLE", "TELS", "TELLES",
    # Noms communs qui pourraient être confondus
    "ADJOINT", "ADJOINTE", "ADJOINTS", "ADJOINTES",
    "AFFAIRE", "AFFAIRES", "AIDE", "AIDES",
    "ANALYSE", "ANALYSES", "ANNONCE", "ANNONCES",
    "ATTESTATION", "ATTESTATIONS", "AUDIT", "AUDITS",
    "AVANTAGE", "AVANTAGES", "AVOIR", "AVOIRS",
    "BA", "BQ", "BP",  # Abréviations
    "CARTE", "CARTES", "CAS", "CAUSE", "CAUSES",
    "COMPTE", "COMPTES", "COMPTABLE", "COMPTABLES",
    "CONSIDÉRANT", "CONSIDERANT",
    "COURRIEL", "COURRIELS", "COURRIER", "COURRIERS",
    "DECISION", "DÉCISION", "DECISIONS", "DÉCISIONS",
    "DOCUMENT", "DOCUMENTS", "DOSSIER", "DOSSIERS",
    "EFFET", "EFFETS", "EXEMPLE", "EXEMPLES",
    "FICHE", "FICHES", "FICHIER", "FICHIERS",
    "FRAIS", "GENERAL", "GÉNÉRALE", "GENERAUX", "GÉNÉRAUX",
    "INFORMATION", "INFORMATIONS", "INTERET", "INTÉRÊT", "INTERETS", "INTÉRÊTS",
    "LETTRE", "LETTRES", "LISTE", "LISTES",
    "MENTION", "MENTIONS", "MESSAGE", "MESSAGES",
    "MONSIEUR", "MESSIEURS", "MADAME", "MESDAMES", "MADEMOISELLE",
    "NOTE", "NOTES", "NUMERO", "NUMÉRO", "NUMEROS", "NUMÉROS",
    "OBJET", "OBJETS", "OBSERVATION", "OBSERVATIONS",
    "OPERATION", "OPÉRATION", "OPERATIONS", "OPÉRATIONS",
    "ORDRE", "ORDRES", "PAGE", "PAGES",
    "PASSEPORT", "PASSEPORTS", "PERMIS",
    "POINT", "POINTS", "PROJET", "PROJETS",
    "QUESTION", "QUESTIONS", "RAPPORT", "RAPPORTS",
    "REFERENCE", "RÉFÉRENCE", "REFERENCES", "RÉFÉRENCES",
    "REGISTRE", "REGISTRES", "REMARQUE", "REMARQUES",
    "SERVICE", "SERVICES", "SITUATION", "SITUATIONS",
    "SUITE", "SUITES", "SUJET", "SUJETS",
    "TELEPHONE", "TÉLÉPHONE", "TELEPHONES", "TÉLÉPHONES",
    "TITRE", "TITRES", "TYPE", "TYPES",
    "USAGE", "USAGES", "VALEUR", "VALEURS",
    # Verbes à l'infinitif courants
    "ACCEPTER", "ACCORDER", "AJOUTER", "ALLER", "APPELER", "APPORTER",
    "ARRIVER", "ASSURER", "ATTENDRE", "AVOIR", "CHERCHER", "COMMENCER",
    "COMPRENDRE", "CONNAÎTRE", "CONNAITRE", "CONSIDERER", "CONSIDÉRER",
    "CONTINUER", "CROIRE", "DEMANDER", "DEVENIR", "DONNER", "ÉCRIRE", "ECRIRE",
    "ENTENDRE", "ENTRER", "ENVOYER", "ÊTRE", "ETRE", "EXISTER",
    "EXPLIQUER", "FAIRE", "FALLOIR", "FINIR", "GARDER", "INDIQUER",
    "LAISSER", "LIRE", "METTRE", "MONTRER", "OBTENIR", "OUVRIR",
    "PARLER", "PARTIR", "PASSER", "PAYER", "PENSER", "PERDRE", "PERMETTRE",
    "PORTER", "POSER", "POUVOIR", "PRENDRE", "PRESENTER", "PRÉSENTER",
    "PRODUIRE", "PROPOSER", "RAPPELER", "RECEVOIR", "RECONNAITRE", "RECONNAÎTRE",
    "REMETTRE", "RENDRE", "REPONDRE", "RÉPONDRE", "RESTER", "RETENIR",
    "RETROUVER", "REVENIR", "SAVOIR", "SEMBLER", "SENTIR", "SERVIR",
    "SORTIR", "SUIVRE", "TENIR", "TERMINER", "TOMBER", "TOUCHER",
    "TOURNER", "TROUVER", "UTILISER", "VENIR", "VIVRE", "VOIR", "VOULOIR",
    # Termes de documents
    "ACCEPTATION", "AFFIRMATION", "AMPLIATIONS", "AVANCES",
    "DÉCLARATION", "DECLARATION",
}

# Ajouter les versions minuscules et title case
MOTS_EXCLUS_LOWER = {m.lower() for m in MOTS_EXCLUS}
MOTS_EXCLUS_TITLE = {m.title() for m in MOTS_EXCLUS}
MOTS_EXCLUS_ALL = MOTS_EXCLUS | MOTS_EXCLUS_LOWER | MOTS_EXCLUS_TITLE


@dataclass
class PersonneDetectee:
    """Une personne détectée dans le document."""
    nom: str  # Nom de famille
    prenoms: List[str] = field(default_factory=list)  # Liste des prénoms
    contexte: str = ""  # Contexte de détection
    position: int = 0  # Position dans le texte


@dataclass
class ResultatAnonymisation:
    """Résultat de l'anonymisation."""
    fichier_source: Path
    fichier_sortie: Path
    personnes_detectees: List[PersonneDetectee] = field(default_factory=list)
    mapping: Dict[str, str] = field(default_factory=dict)
    nb_remplacements: int = 0
    succes: bool = True
    erreur: str = ""


class AnonymiseurPrecis:
    """Anonymise les documents avec détection contextuelle précise."""

    # Patterns de détection contextuelle
    PATTERNS = {
        # Monsieur/Madame NOM Prénom ou Prénom NOM
        'civilite_nom': re.compile(
            r'\b(Monsieur|Madame|Mademoiselle|M\.|Mme|Mlle|Mr\.?)\s+'
            r'([A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+(?:\s+[A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+){0,3})',
            re.UNICODE
        ),

        # NOM EN MAJUSCULES (contexte parties)
        'nom_majuscules': re.compile(
            r'\b([A-ZÀ-Ÿ]{2,}(?:\s+[A-ZÀ-Ÿ]{2,})?)\s*,\s*'
            r'([A-ZÀ-Ÿ][a-zà-ÿ]+(?:[\s\-][A-ZÀ-Ÿ][a-zà-ÿ]+)*)',
            re.UNICODE
        ),

        # né(e) NOM ou née NOM
        'nee': re.compile(
            r'\bné(?:e)?\s+([A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+)',
            re.UNICODE
        ),

        # époux/épouse NOM
        'epoux': re.compile(
            r'\b(?:époux|épouse|veuf|veuve)\s+([A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+)',
            re.UNICODE
        ),

        # Maître NOM (notaires)
        'maitre': re.compile(
            r'\bMa[îi]tre\s+([A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+(?:\s+[A-ZÀ-Ÿ][A-ZÀ-Ÿa-zà-ÿ\-]+)?)',
            re.UNICODE
        ),

        # Prénom NOM après virgule dans contexte identité
        'prenom_nom_identite': re.compile(
            r',\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:[\s\-][A-ZÀ-Ÿ][a-zà-ÿ]+)*)\s*,',
            re.UNICODE
        ),
    }

    # Noms de remplacement
    NOMS_ANONYMES = [
        "DUPONT", "MARTIN", "DURAND", "BERNARD", "MOREAU", "PETIT",
        "SIMON", "LAURENT", "MICHEL", "LEROY", "GARCIA", "ROUX"
    ]

    PRENOMS_ANONYMES = [
        "Jean", "Pierre", "Marie", "Paul", "Anne", "Louis", "Claire",
        "Michel", "Sophie", "André", "Julie", "Henri", "Lucie", "Marc"
    ]

    def __init__(self):
        self.mapping_noms: Dict[str, str] = {}
        self.mapping_prenoms: Dict[str, str] = {}
        self.compteur_noms = 0
        self.compteur_prenoms = 0

    def _est_mot_exclu(self, mot: str) -> bool:
        """Vérifie si le mot est exclu de l'anonymisation."""
        return mot in MOTS_EXCLUS_ALL or mot.upper() in MOTS_EXCLUS

    def _obtenir_nom_anonyme(self, nom_original: str) -> str:
        """Retourne un nom anonyme cohérent pour un nom donné."""
        # Normaliser pour le mapping
        cle = nom_original.upper()

        if cle in self.mapping_noms:
            nom_anonyme = self.mapping_noms[cle]
        else:
            # Nouveau nom
            idx = self.compteur_noms % len(self.NOMS_ANONYMES)
            nom_anonyme = self.NOMS_ANONYMES[idx]
            self.mapping_noms[cle] = nom_anonyme
            self.compteur_noms += 1

        # Adapter la casse
        if nom_original.isupper():
            return nom_anonyme.upper()
        elif nom_original.istitle():
            return nom_anonyme.title()
        return nom_anonyme

    def _obtenir_prenom_anonyme(self, prenom_original: str) -> str:
        """Retourne un prénom anonyme cohérent."""
        cle = prenom_original.lower()

        if cle in self.mapping_prenoms:
            prenom_anonyme = self.mapping_prenoms[cle]
        else:
            idx = self.compteur_prenoms % len(self.PRENOMS_ANONYMES)
            prenom_anonyme = self.PRENOMS_ANONYMES[idx]
            self.mapping_prenoms[cle] = prenom_anonyme
            self.compteur_prenoms += 1

        # Les prénoms sont généralement en Title Case
        return prenom_anonyme.title() if prenom_original.istitle() else prenom_anonyme

    def _detecter_noms_contextuels(self, texte: str) -> Dict[str, str]:
        """Détecte les noms dans leur contexte et retourne le mapping."""
        noms_trouves = {}

        # Pattern 1: Civilité + Nom/Prénom
        for match in self.PATTERNS['civilite_nom'].finditer(texte):
            noms_bruts = match.group(2).split()
            for nom in noms_bruts:
                if not self._est_mot_exclu(nom) and len(nom) >= 2:
                    if nom.isupper() or nom[0].isupper():
                        # C'est probablement un nom ou prénom
                        if nom.isupper():
                            noms_trouves[nom] = self._obtenir_nom_anonyme(nom)
                        else:
                            noms_trouves[nom] = self._obtenir_prenom_anonyme(nom)

        # Pattern 2: NOM, Prénom
        for match in self.PATTERNS['nom_majuscules'].finditer(texte):
            nom = match.group(1).strip()
            prenoms = match.group(2).strip()

            if not self._est_mot_exclu(nom):
                noms_trouves[nom] = self._obtenir_nom_anonyme(nom)

            for prenom in prenoms.split():
                if not self._est_mot_exclu(prenom) and len(prenom) >= 2:
                    noms_trouves[prenom] = self._obtenir_prenom_anonyme(prenom)

        # Pattern 3: née NOM
        for match in self.PATTERNS['nee'].finditer(texte):
            nom = match.group(1)
            if not self._est_mot_exclu(nom):
                noms_trouves[nom] = self._obtenir_nom_anonyme(nom)

        # Pattern 4: époux/épouse NOM
        for match in self.PATTERNS['epoux'].finditer(texte):
            nom = match.group(1)
            if not self._est_mot_exclu(nom):
                noms_trouves[nom] = self._obtenir_nom_anonyme(nom)

        # Pattern 5: Maître NOM
        for match in self.PATTERNS['maitre'].finditer(texte):
            noms_bruts = match.group(1).split()
            for nom in noms_bruts:
                if not self._est_mot_exclu(nom):
                    noms_trouves[nom] = self._obtenir_nom_anonyme(nom)

        return noms_trouves

    def _extraire_texte_complet(self, doc: Document) -> str:
        """Extrait tout le texte pour analyse."""
        textes = []

        for para in doc.paragraphs:
            textes.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    textes.append(cell.text)

        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    textes.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    textes.append(para.text)

        return '\n'.join(textes)

    def _remplacer_dans_run(self, run, mapping: Dict[str, str]) -> int:
        """Remplace dans un run en préservant le formatage."""
        if not run.text:
            return 0

        texte = run.text
        nb_remplacements = 0

        # Trier par longueur décroissante
        for original, anonyme in sorted(mapping.items(), key=lambda x: -len(x[0])):
            if original in texte:
                texte = texte.replace(original, anonyme)
                nb_remplacements += 1

        if nb_remplacements > 0:
            run.text = texte

        return nb_remplacements

    def anonymiser(self, fichier_source: Path, fichier_sortie: Optional[Path] = None) -> ResultatAnonymisation:
        """Anonymise un document DOCX."""
        fichier_source = Path(fichier_source)

        if fichier_sortie is None:
            fichier_sortie = fichier_source.parent / f"{fichier_source.stem}_anonyme{fichier_source.suffix}"

        resultat = ResultatAnonymisation(
            fichier_source=fichier_source,
            fichier_sortie=Path(fichier_sortie)
        )

        try:
            print(f"\n{'='*60}")
            print(f"📄 Fichier: {fichier_source.name}")
            print('='*60)

            # Charger
            doc = Document(str(fichier_source))

            # Extraire et analyser le texte
            texte_complet = self._extraire_texte_complet(doc)

            # Détecter les noms contextuellement
            print("🔍 Détection contextuelle des noms...")
            mapping = self._detecter_noms_contextuels(texte_complet)

            print(f"   → {len(mapping)} noms/prénoms détectés")

            if mapping:
                print("\n📋 Mapping:")
                for orig, anon in sorted(mapping.items()):
                    print(f"   {orig} → {anon}")

            resultat.mapping = mapping.copy()

            # Appliquer les remplacements
            print("\n✏️  Application des remplacements...")

            # Paragraphes
            for para in doc.paragraphs:
                for run in para.runs:
                    resultat.nb_remplacements += self._remplacer_dans_run(run, mapping)

            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                resultat.nb_remplacements += self._remplacer_dans_run(run, mapping)

            # Headers/Footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        for run in para.runs:
                            resultat.nb_remplacements += self._remplacer_dans_run(run, mapping)
                if section.footer:
                    for para in section.footer.paragraphs:
                        for run in para.runs:
                            resultat.nb_remplacements += self._remplacer_dans_run(run, mapping)

            # Sauvegarder
            doc.save(str(fichier_sortie))

            print(f"\n💾 Sauvegardé: {fichier_sortie.name}")
            print(f"✅ {resultat.nb_remplacements} remplacements effectués")

            resultat.succes = True

        except Exception as e:
            resultat.succes = False
            resultat.erreur = str(e)
            print(f"❌ Erreur: {e}")
            import traceback
            traceback.print_exc()

        return resultat


def main():
    parser = argparse.ArgumentParser(
        description="Anonymise les documents DOCX avec détection contextuelle précise"
    )
    parser.add_argument("fichiers", nargs="+", help="Fichiers DOCX")
    parser.add_argument("-o", "--output", help="Dossier de sortie")
    parser.add_argument("--batch", action="store_true", help="Mode batch")
    parser.add_argument("--mapping", help="Exporter le mapping en JSON")

    args = parser.parse_args()

    # Résoudre les fichiers
    fichiers = []
    for f in args.fichiers:
        p = Path(f)
        if p.exists():
            fichiers.append(p)
        else:
            fichiers.extend(Path('.').glob(f))

    if not fichiers:
        print("❌ Aucun fichier trouvé")
        sys.exit(1)

    print(f"📁 {len(fichiers)} fichier(s) à traiter")

    anonymiseur = AnonymiseurPrecis()
    resultats = []

    for fichier in fichiers:
        if args.output:
            dossier = Path(args.output)
            dossier.mkdir(exist_ok=True)
            sortie = dossier / f"{fichier.stem}_anonyme{fichier.suffix}"
        else:
            sortie = None

        resultat = anonymiseur.anonymiser(fichier, sortie)
        resultats.append(resultat)

    # Exporter mapping global si demandé
    if args.mapping:
        mapping_global = {
            "noms": anonymiseur.mapping_noms,
            "prenoms": anonymiseur.mapping_prenoms,
            "fichiers": {
                r.fichier_source.name: {
                    "sortie": r.fichier_sortie.name,
                    "mapping": r.mapping,
                    "remplacements": r.nb_remplacements,
                    "succes": r.succes
                }
                for r in resultats
            }
        }

        with open(args.mapping, 'w', encoding='utf-8') as f:
            json.dump(mapping_global, f, ensure_ascii=False, indent=2)

        print(f"\n📄 Mapping exporté: {args.mapping}")

    # Résumé
    print("\n" + "="*60)
    print("📊 RÉSUMÉ FINAL")
    print("="*60)

    succes = sum(1 for r in resultats if r.succes)
    total_remp = sum(r.nb_remplacements for r in resultats)

    print(f"   Fichiers: {len(resultats)}")
    print(f"   Succès: {succes}/{len(resultats)}")
    print(f"   Total remplacements: {total_remp}")

    print("\n📋 Mapping global:")
    print("   Noms:")
    for orig, anon in sorted(anonymiseur.mapping_noms.items()):
        print(f"      {orig} → {anon}")
    print("   Prénoms:")
    for orig, anon in sorted(anonymiseur.mapping_prenoms.items()):
        print(f"      {orig} → {anon}")


if __name__ == "__main__":
    main()
