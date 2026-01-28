#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_statuts_final.py
=========================

Copie EXACTE du document Statuts (1).docx avec uniquement les remplacements de texte.
AUCUNE modification de formatage.
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from pathlib import Path
import comtypes.client
import pythoncom


def generer_statuts_final():
    """Génère le document Statuts Officiel - copie exacte avec variables."""

    print('=' * 80)
    print('GENERATION STATUTS OFFICIEL - COPIE EXACTE')
    print('=' * 80)

    # Charger le document original
    print('\n1. Chargement document original...')
    doc = Document('docs_original/Statuts (1).docx')
    print(f'   OK - {len(doc.paragraphs)} paragraphes')

    # Dictionnaire de remplacements - ORDRE IMPORTANT (plus spécifique en premier)
    remplacements = {
        # Organismes (avant les simples mots pour éviter remplacement partiel)
        'FIDUCIAL SOFIRAL': 'CABINET_EXEMPLE',
        'Fiducial Sofiral': 'Cabinet_Exemple',
        'FIDUCIAL BANQUE': 'BANQUE_EXEMPLE',
        'Fiducial Banque': 'Banque_Exemple',
        'FIDUCIAL': 'CABINET',
        'Fiducial': 'Cabinet',

        # Adresses complètes (avant codes postaux)
        '29 Chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 Chemin du Panorama': '00 rue Exemple',
        '29 chemin du Panorama': '00 rue Exemple',

        # Villes avec apostrophe (avant autres villes)
        "TAIN L'HERMITAGE": 'VILLE_NAISSANCE',
        "Tain l'Hermitage": 'Ville_Naissance',
        'CHANTEMERLE-LES-BLES': 'VILLE_MARIAGE',
        'Chantemerle-les-Bles': 'Ville_Mariage',

        # Société
        '985 331 354 R.C.S. LYON': '000 000 000 R.C.S. METROPOLE',
        '985 331 354': '000 000 000',
        'BLOUGE': 'NOM_SOCIETE',
        'Blouge': 'Nom_Societe',

        # Personnes - Noms
        'AUVRAY': 'NOM_FAMILLE',
        'Auvray': 'Nom_Famille',
        'ROBIN': 'NOM_NAISSANCE',
        'Robin': 'Nom_Naissance',
        'COHENDET': 'NOM_NOTAIRE',
        'Cohendet': 'Nom_Notaire',
        'GIEN': 'NOM_AVOCAT',
        'Gien': 'Nom_Avocat',

        # Prénoms
        'Dominique': 'Prenom1',
        'Christian': 'Prenom3',
        'Viviane': 'Prenom2',
        'Blanche': 'Prenom6',
        'Célie': 'Prenom7',
        'Pierre': 'Prenom_Notaire',

        # Villes
        'DARDILLY': 'VILLE',
        'Dardilly': 'Ville',
        'PARIS': 'METROPOLE',
        'Paris': 'Metropole',
        'LYON': 'METROPOLE',
        'Lyon': 'Metropole',
        'Charolles': 'Ville_Notaire',

        # Codes postaux
        '69570': '00000',
        '75012': '00001',
        '26600': '00002',
        '71120': '00003',

        # Dates
        '20 février 1961': 'JJ mois AAAA',
        '16 avril 1961': 'JJ mois AAAA',
        '09 septembre 1989': 'JJ mois AAAA',
        '24 juin 1989': 'JJ mois AAAA',

        # Emails
        'lyon.avocat@fiducial.fr': 'email@exemple.fr',
    }

    # Remplacer dans les paragraphes - AUCUNE MODIFICATION DE FORMATAGE
    print('\n2. Remplacement dans paragraphes...')
    count_para = 0
    for para in doc.paragraphs:
        for run in para.runs:
            texte_orig = run.text
            # Appliquer tous les remplacements
            for ancien, nouveau in remplacements.items():
                run.text = run.text.replace(ancien, nouveau)
            if run.text != texte_orig:
                count_para += 1

    print(f'   {count_para} modifications')

    # Remplacer dans les tableaux - AUCUNE MODIFICATION DE FORMATAGE
    print('\n3. Remplacement dans tableaux...')
    count_table = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        texte_orig = run.text
                        # Appliquer tous les remplacements
                        for ancien, nouveau in remplacements.items():
                            run.text = run.text.replace(ancien, nouveau)
                        if run.text != texte_orig:
                            count_table += 1

    print(f'   {count_table} modifications')

    # Sauvegarder
    print('\n4. Sauvegarde...')
    output_dir = Path('outputs')
    output_dir.mkdir(exist_ok=True)

    docx_path = output_dir / 'officiel statuts.docx'
    doc.save(str(docx_path))
    print(f'   OK - {docx_path}')

    # Convertir en PDF
    print('\n5. Conversion PDF...')
    pdf_path = output_dir / 'officiel statuts.pdf'

    pythoncom.CoInitialize()
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        docx_abs = str(docx_path.resolve())
        pdf_abs = str(pdf_path.resolve())

        doc_word = word.Documents.Open(docx_abs)
        doc_word.SaveAs(pdf_abs, FileFormat=17)
        doc_word.Close()
        word.Quit()

        print(f'   OK - {pdf_path}')
    finally:
        pythoncom.CoUninitialize()

    # Vérification
    print('\n' + '=' * 80)
    print('VERIFICATION')
    print('=' * 80)

    doc_verif = Document(str(docx_path))
    import PyPDF2
    with open(str(pdf_path), 'rb') as f:
        nb_pages = len(PyPDF2.PdfReader(f).pages)

    print(f'\nStructure:')
    print(f'  Paragraphes: {len(doc_verif.paragraphs)}')
    print(f'  Tableaux: {len(doc_verif.tables)}')
    print(f'  Sections: {len(doc_verif.sections)}')
    print(f'  Pages PDF: {nb_pages}')

    # Vérifier variables
    texte = ' '.join(p.text for p in doc_verif.paragraphs)
    interdits = ['BLOUGE', 'AUVRAY', 'Dominique', 'DARDILLY', '69570']
    trouves = [x for x in interdits if x in texte]

    if trouves:
        print(f'\n! Valeurs specifiques restantes: {trouves}')
    else:
        print(f'\nOK - Toutes variables generiques')

    print('\n' + '=' * 80)
    print('TERMINE')
    print('=' * 80)


if __name__ == '__main__':
    generer_statuts_final()
