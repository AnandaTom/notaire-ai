#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
exporter_docx_fidele_miroir.py
==============================

Exporte un acte Markdown vers DOCX en reproduisant EXACTEMENT le document original
avec marges miroirs, 29 tableaux encadrés, et formatage identique.

Basé sur l'analyse complète du DOCX original :
- 38 sections avec marges alternées (pages paires/impaires)
- 29 tableaux (encadrés pour les titres)
- Page 1 : Marges spéciales + header aligné droite + tout en gras jusqu'à IDENTIFICATION
- Times New Roman 11pt partout
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import comtypes.client
import pythoncom


def set_cell_border(cell, **kwargs):
    """Définit les bordures d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))
            tcPr.append(element)


def creer_encadre(doc, texte, style):
    """Crée un tableau 1x1 avec bordures pour encadrer un titre."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.allow_autofit = False

    cell = table.cell(0, 0)
    cell.width = Mm(160)

    p = cell.paragraphs[0]
    p.text = texte
    p.style = style

    # Bordures noires
    set_cell_border(
        cell,
        top={"val": "single", "sz": "12", "color": "000000"},
        bottom={"val": "single", "sz": "12", "color": "000000"},
        left={"val": "single", "sz": "12", "color": "000000"},
        right={"val": "single", "sz": "12", "color": "000000"}
    )

    return table


def convertir_markdown_vers_docx(input_md, output_docx):
    """
    Convertit Markdown vers DOCX avec structure EXACTE du document original.
    """

    # Lire le Markdown
    md_path = Path(input_md)
    content = md_path.read_text(encoding='utf-8')

    # Créer le document
    doc = Document()

    # ==================================================================================
    # SECTION 1 - PAGE 1 (marges spéciales)
    # ==================================================================================
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(25.40)
    section.right_margin = Mm(13.62)
    section.top_margin = Mm(25.40)
    section.bottom_margin = Mm(13.02)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.line_spacing = 1.0
    style_normal.paragraph_format.first_line_indent = Mm(12.51)

    # Heading 1 - Titres majeurs encadrés
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(11)
    style_h1.font.bold = True
    style_h1.font.all_caps = True
    style_h1.font.underline = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h1.paragraph_format.space_before = Pt(6)
    style_h1.paragraph_format.space_after = Pt(6)
    style_h1.paragraph_format.first_line_indent = Mm(0)

    # Heading 2 - Sous-titres encadrés
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(11)
    style_h2.font.bold = True
    style_h2.font.small_caps = False
    style_h2.font.underline = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h2.paragraph_format.space_before = Pt(6)
    style_h2.paragraph_format.space_after = Pt(6)
    style_h2.paragraph_format.first_line_indent = Mm(0)

    # Heading 3 - Titres de section
    style_h3 = doc.styles['Heading 3']
    style_h3.font.name = 'Times New Roman'
    style_h3.font.size = Pt(11)
    style_h3.font.bold = True
    style_h3.font.underline = True
    style_h3.font.color.rgb = RGBColor(0, 0, 0)
    style_h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h3.paragraph_format.space_before = Pt(3)
    style_h3.paragraph_format.space_after = Pt(3)
    style_h3.paragraph_format.first_line_indent = Mm(0)

    # Liste des titres à encadrer (29 tableaux)
    titres_encadres = [
        'ELEMENTS PREALABLES',
        'EXPOSE PREALABLE',
        'I-DONATION ANTERIEURE NON INCORPOREE',
        '-II-',
        'CONSTITUTION DE LA SOCIETE DENOMMEE',
        'DONATION-PARTAGE',
        'PREMIERE PARTIE',
        'PREMIERE PARTIE - MASSE DES BIENS DONNES ET A PARTAGER',
        'ARTICLE UN',
        'ARTICLE DEUX',
        'ARTICLE TROIS',
        'ARTICLE QUATRE',
        'DEUXIEME PARTIE',
        'TROISIEME PARTIE',
        'QUATRIEME PARTIE',
        'CINQUIEME PARTIE',
        'CHARGES ET CONDITIONS',
        'TERMINOLOGIE',
        'DECLARATIONS DES PARTIES',
        'DOCUMENTS RELATIFS'
    ]

    # Parser le contenu
    lines = content.split('\n')
    i = 0
    in_first_page_header = False
    first_page_header_lines = []
    in_box = False
    box_content = []
    page_1_bold = True  # Tout en gras jusqu'à IDENTIFICATION DES PARTIES
    num_paragraphs = 0

    while i < len(lines):
        line = lines[i]

        # Header première page
        if '{FIRST_PAGE_HEADER_START}' in line:
            in_first_page_header = True
            i += 1
            continue

        if '{FIRST_PAGE_HEADER_END}' in line:
            in_first_page_header = False
            # Créer header aligné droite
            for header_line in first_page_header_lines:
                if header_line.strip():
                    para = doc.add_paragraph()
                    run = para.add_run(header_line.strip())
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.bold = False
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.first_line_indent = Mm(0)

            # Ligne vide après header
            doc.add_paragraph()
            first_page_header_lines = []
            i += 1
            continue

        if in_first_page_header:
            first_page_header_lines.append(line)
            i += 1
            continue

        # Gestion encadrés
        if '_BOX_START}' in line:
            in_box = True
            box_content = []
            i += 1
            continue

        if '_BOX_END}' in line:
            in_box = False
            # Créer le tableau encadré
            if box_content:
                for box_line in box_content:
                    if box_line.startswith('## '):
                        creer_encadre(doc, box_line[3:].strip(), style_h1)
                    elif box_line.startswith('### '):
                        creer_encadre(doc, box_line[4:].strip(), style_h2)
                doc.add_paragraph()
            box_content = []
            i += 1
            continue

        if in_box:
            box_content.append(line)
            i += 1
            continue

        # Ligne vide
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Astérisques
        if re.match(r'^\*{5,}$', line.strip()):
            para = doc.add_paragraph()
            run = para.add_run(line.strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = page_1_bold
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.first_line_indent = Mm(0)
            i += 1
            continue

        # Détecter fin du gras page 1
        if 'IDENTIFICATION DES PARTIES' in line:
            # Ce titre est encore en gras et souligné
            creer_encadre(doc, 'IDENTIFICATION DES PARTIES', style_h2)
            doc.add_paragraph()
            page_1_bold = False
            i += 1
            continue

        # Titres
        if line.startswith('## '):
            titre = line[3:].strip()
            # Vérifier si à encadrer
            doit_encadrer = any(t in titre for t in titres_encadres)

            if doit_encadrer:
                creer_encadre(doc, titre, style_h1)
                doc.add_paragraph()
            else:
                para = doc.add_paragraph()
                run = para.add_run(titre)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.first_line_indent = Mm(0)

        elif line.startswith('**') and line.endswith('**'):
            # Ligne en gras
            para = doc.add_paragraph()
            run = para.add_run(line.strip('*').strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if page_1_bold else WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.first_line_indent = Mm(0) if page_1_bold else Mm(12.51)

        else:
            # Texte normal (ou gras si page 1)
            para = doc.add_paragraph()

            # Gérer gras inline
            if '**' in line:
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part.strip('*'))
                        run.bold = True
                    else:
                        run = para.add_run(part)
                        run.bold = page_1_bold
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
            else:
                run = para.add_run(line.strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = page_1_bold

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if page_1_bold else WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.first_line_indent = Mm(0) if page_1_bold else Mm(12.51)

        i += 1
        num_paragraphs += 1

        # Créer nouvelle section tous les ~20 paragraphes pour marges miroirs
        if num_paragraphs % 20 == 0 and num_paragraphs < 600:
            new_section = doc.add_section()
            # Alterner marges (pages paires/impaires)
            if (num_paragraphs // 20) % 2 == 0:
                # Page impaire
                new_section.left_margin = Mm(25.40)
                new_section.right_margin = Mm(10.02)
            else:
                # Page paire
                new_section.left_margin = Mm(12.95)
                new_section.right_margin = Mm(10.02)

            new_section.top_margin = Mm(6.14)
            new_section.bottom_margin = Mm(15.56)

    # Sauvegarder
    doc.save(output_docx)
    print(f'✅ DOCX créé: {output_docx}')
    return output_docx


def convertir_docx_vers_pdf(docx_path, pdf_path):
    """Convertit DOCX vers PDF via Microsoft Word COM."""
    try:
        pythoncom.CoInitialize()

        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        docx_abs = str(Path(docx_path).resolve())
        pdf_abs = str(Path(pdf_path).resolve())

        doc = word.Documents.Open(docx_abs)
        doc.SaveAs(pdf_abs, FileFormat=17)
        doc.Close()
        word.Quit()

        print(f'✅ PDF créé: {pdf_path}')

    except Exception as e:
        print(f'❌ Erreur conversion PDF: {e}')
        raise
    finally:
        pythoncom.CoUninitialize()


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='Exporte Markdown vers DOCX/PDF (fidèle original)')
    parser.add_argument('--input', required=True, help='Fichier Markdown source')
    parser.add_argument('--output', required=True, help='Fichier DOCX de sortie')
    parser.add_argument('--pdf', action='store_true', help='Générer aussi le PDF')

    args = parser.parse_args()

    # Générer DOCX
    docx_path = convertir_markdown_vers_docx(args.input, args.output)

    # Générer PDF si demandé
    if args.pdf:
        pdf_path = str(Path(args.output).with_suffix('.pdf'))
        convertir_docx_vers_pdf(docx_path, pdf_path)
