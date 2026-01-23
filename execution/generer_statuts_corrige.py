#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_statuts_corrige.py
===========================

Génère les statuts avec corrections:
1. Ajoute fond bleu au tableau STATUTS (perdu lors conversion PDF→DOCX)
2. Supprime paragraphes vides qui créent page blanche entre page 1 et 2
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import comtypes.client
import pythoncom


def set_cell_background(cell, color):
    """Définit la couleur de fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Supprimer ancien shd si existe
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)

    # Créer nouveau shd
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def generer_statuts_corrige():
    """Génère le document Statuts avec corrections."""

    print('=' * 80)
    print('GENERATION STATUTS - VERSION CORRIGEE')
    print('=' * 80)

    # Charger le document original
    print('\n1. Chargement document original...')
    doc = Document('docs_originels/Statuts (1).docx')
    print(f'   OK - {len(doc.paragraphs)} paragraphes, {len(doc.tables)} tableaux')

    # Dictionnaire de remplacements - ORDRE IMPORTANT
    remplacements = {
        # Organismes (avant les simples mots)
        'FIDUCIAL SOFIRAL': 'CABINET_EXEMPLE',
        'Fiducial Sofiral': 'Cabinet_Exemple',
        'FIDUCIAL BANQUE': 'BANQUE_EXEMPLE',
        'Fiducial Banque': 'Banque_Exemple',
        'FIDUCIAL': 'CABINET',
        'Fiducial': 'Cabinet',

        # Adresses complètes
        '29 Chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 Chemin du Panorama': '00 rue Exemple',
        '29 chemin du Panorama': '00 rue Exemple',

        # Villes avec apostrophe
        "TAIN L'HERMITAGE": 'VILLE_NAISSANCE',
        "Tain l'Hermitage": 'Ville_Naissance',
        'CHANTEMERLE-LES-BLES': 'VILLE_MARIAGE',
        'Chantemerle-les-Bles': 'Ville_Mariage',

        # Société
        '985 331 354 R.C.S. LYON': '000 000 000 R.C.S. METROPOLE',
        '985 331 354': '000 000 000',
        'BLOUGE': 'NOM_SOCIETE',
        'Blouge': 'Nom_Societe',

        # Personnes
        'AUVRAY': 'NOM_FAMILLE',
        'Auvray': 'Nom_Famille',
        'ROBIN': 'NOM_NAISSANCE',
        'Robin': 'Nom_Naissance',
        'COHENDET': 'NOM_NOTAIRE',
        'Cohendet': 'Nom_Notaire',
        'GIEN': 'NOM_AVOCAT',
        'Gien': 'Nom_Avocat',

        # Prénoms
        'Dominique': 'Prenom1',
        'Christian': 'Prenom3',
        'Viviane': 'Prenom2',
        'Blanche': 'Prenom6',
        'Célie': 'Prenom7',
        'Pierre': 'Prenom_Notaire',

        # Villes
        'DARDILLY': 'VILLE',
        'Dardilly': 'Ville',
        'PARIS': 'METROPOLE',
        'Paris': 'Metropole',
        'LYON': 'METROPOLE',
        'Lyon': 'Metropole',
        'Charolles': 'Ville_Notaire',

        # Codes postaux
        '69570': '00000',
        '75012': '00001',
        '26600': '00002',
        '71120': '00003',

        # Dates
        '20 février 1961': 'JJ mois AAAA',
        '16 avril 1961': 'JJ mois AAAA',
        '09 septembre 1989': 'JJ mois AAAA',
        '24 juin 1989': 'JJ mois AAAA',

        # Emails
        'lyon.avocat@fiducial.fr': 'email@exemple.fr',
    }

    # Remplacer dans les paragraphes
    print('\n2. Remplacement dans paragraphes...')
    count_para = 0
    for para in doc.paragraphs:
        for run in para.runs:
            texte_orig = run.text
            for ancien, nouveau in remplacements.items():
                run.text = run.text.replace(ancien, nouveau)
            if run.text != texte_orig:
                count_para += 1

    print(f'   {count_para} modifications')

    # Remplacer dans les tableaux
    print('\n3. Remplacement dans tableaux...')
    count_table = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        texte_orig = run.text
                        for ancien, nouveau in remplacements.items():
                            run.text = run.text.replace(ancien, nouveau)
                        if run.text != texte_orig:
                            count_table += 1

    print(f'   {count_table} modifications')

    # CORRECTION 1: Ajouter fond bleu au tableau "STATUTS"
    print('\n4. Correction formatage tableau STATUTS...')
    table_statuts = doc.tables[1]  # Table "STATUTS"
    cell_statuts = table_statuts.cell(0, 0)
    set_cell_background(cell_statuts, '001f5f')
    print('   ✅ Fond bleu 001f5f ajouté au tableau STATUTS')

    # CORRECTION 2: Supprimer paragraphes vides et forcer saut de page
    print('\n5. Correction mise en page...')

    # Structure originale (dans le body XML):
    # - Para 0: vide
    # - Para 1: Header
    # - Table 0: STATUTS CONSTITUTIFS
    # - Para 2: vide
    # - Table 1: STATUTS
    # - Para 3-6: vides (4 paragraphes à supprimer)
    # - Table 2: LES SOUSSIGNÉS ← Doit commencer page 2!

    # Supprimer les paragraphes 3, 4, 5, 6 (de la fin vers le début)
    for i in range(6, 2, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    # Après suppression, la structure est:
    # - Para 0: vide
    # - Para 1: Header
    # - Table 0: STATUTS CONSTITUTIFS
    # - Para 2: vide
    # - Table 1: STATUTS
    # - Table 2: LES SOUSSIGNÉS

    # Forcer saut de page sur le premier paragraphe qui vient APRÈS le tableau STATUTS
    # Après suppression, la structure dans doc.paragraphs est:
    # Para 0: vide
    # Para 1: Header
    # Para 2: vide (entre STATUTS CONSTITUTIFS et STATUTS)
    # Para 3 et suivants: contenu après les tableaux

    # On cherche le premier paragraphe après le tableau LES SOUSSIGNÉS
    # Ce sera le paragraphe "1. Monsieur..."
    # Il faut trouver son index

    # La structure body contient tableaux ET paragraphes mélangés
    # Après suppression para 3-6, il reste:
    # Body: Para0, Para1, Table0(STATUTS CONST), Para2, Table1(STATUTS), Table2(LES SOUSS), Para3(1.Monsieur)...

    # On va simplement mettre le saut de page sur Para 2 (celui entre les 2 tableaux de page 1)
    # NON! Ça va mettre Para 2 en page 2

    # Solution: insérer un paragraphe vide JUSTE AVANT Table2, avec page_break
    body = doc.element.body
    table_soussignes_elem = doc.tables[2]._element

    # Créer paragraphe avec saut de page
    from docx.oxml import OxmlElement
    new_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # Le paragraphe lui-même sera vide mais causera un saut de page
    pageBreak = OxmlElement('w:pageBreakBefore')
    pPr.append(pageBreak)
    new_para.append(pPr)

    # Insérer JUSTE AVANT le tableau LES SOUSSIGNÉS
    table_souss_index = body.index(table_soussignes_elem)
    body.insert(table_souss_index, new_para)

    print('   ✅ Paragraphes vides 3-6 supprimés')
    print('   ✅ Saut de page inséré avant LES SOUSSIGNÉS')

    # Sauvegarder
    print('\n6. Sauvegarde...')
    output_dir = Path('outputs')
    output_dir.mkdir(exist_ok=True)

    docx_path = output_dir / 'officiel statuts.docx'
    doc.save(str(docx_path))
    print(f'   OK - {docx_path}')

    # Convertir en PDF
    print('\n7. Conversion PDF...')
    pdf_path = output_dir / 'officiel statuts.pdf'

    pythoncom.CoInitialize()
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        docx_abs = str(docx_path.resolve())
        pdf_abs = str(pdf_path.resolve())

        doc_word = word.Documents.Open(docx_abs)
        doc_word.SaveAs(pdf_abs, FileFormat=17)
        doc_word.Close()
        word.Quit()

        print(f'   OK - {pdf_path}')
    finally:
        pythoncom.CoUninitialize()

    # Vérification
    print('\n' + '=' * 80)
    print('VERIFICATION')
    print('=' * 80)

    doc_verif = Document(str(docx_path))
    import PyPDF2
    with open(str(pdf_path), 'rb') as f:
        nb_pages = len(PyPDF2.PdfReader(f).pages)

    print(f'\nStructure:')
    print(f'  Paragraphes: {len(doc_verif.paragraphs)}')
    print(f'  Tableaux: {len(doc_verif.tables)}')
    print(f'  Pages PDF: {nb_pages}')

    # Vérifier variables
    texte = ' '.join(p.text for p in doc_verif.paragraphs)
    interdits = ['BLOUGE', 'AUVRAY', 'Dominique', 'DARDILLY', '69570']
    trouves = [x for x in interdits if x in texte]

    if trouves:
        print(f'\n❌ Valeurs spécifiques restantes: {trouves}')
    else:
        print(f'\n✅ Toutes variables génériques')

    # Vérifier le fond du tableau STATUTS
    print('\nVérification tableau STATUTS:')
    table_verif = doc_verif.tables[1]
    cell_verif = table_verif.cell(0, 0)
    tc = cell_verif._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))

    if shd is not None:
        fill = shd.get(qn('w:fill'))
        if fill == '001f5f':
            print(f'  ✅ Fond bleu 001f5f confirmé')
        else:
            print(f'  ❌ Fond: {fill}')
    else:
        print(f'  ❌ Pas de fond')

    # Vérifier pages 1 et 2 du PDF
    print('\nVérification PDF:')
    with open(str(pdf_path), 'rb') as f:
        pdf = PyPDF2.PdfReader(f)

        page1 = pdf.pages[0].extract_text()
        page2 = pdf.pages[1].extract_text()

        if 'STATUTS' in page1:
            print('  ✅ Page 1: "STATUTS" présent')
        else:
            print('  ❌ Page 1: "STATUTS" manquant')

        if 'SOUSSIGN' in page2:
            print('  ✅ Page 2: "LES SOUSSIGNÉS" présent')
        else:
            print('  ❌ Page 2: "LES SOUSSIGNÉS" manquant')

    print('\n' + '=' * 80)
    print('TERMINE')
    print('=' * 80)


if __name__ == '__main__':
    generer_statuts_corrige()
