#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_depuis_template_docx.py
===============================

Génère un nouvel acte notarial en utilisant un DOCX original comme template.
Copie la structure EXACTE (styles, formatage, tableaux) et substitue les variables.

AVANTAGE: Format 100% identique à l'original car on copie le document, pas qu'on le recrée.

Usage:
    python generer_depuis_template_docx.py \\
        --template "docs_originels/Trame vente lots de copropriété.docx" \\
        --variables donnees.json \\
        --output "acte_genere.docx"

Format du fichier de variables (JSON):
{
    "vendeur": {
        "nom": "DUPONT",
        "prenoms": "Jean Marie",
        "profession": "Ingénieur",
        "adresse": "12 rue de la Paix, 75001 Paris",
        "date_naissance": "15/03/1975",
        "lieu_naissance": "Lyon",
        "nationalite": "française",
        "situation_matrimoniale": "Célibataire"
    },
    "acquereur": { ... },
    "bien": {
        "adresse": "45 avenue des Champs, 69001 Lyon",
        "lot": "12",
        "tantiemes": "150/10000"
    },
    "prix": {
        "montant": 250000,
        "lettres": "deux cent cinquante mille"
    },
    "date_acte": "15/01/2025"
}
"""

import argparse
import json
import re
import copy
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =============================================================================
# PATTERNS DE SUBSTITUTION
# =============================================================================

# Patterns de variables dans les documents notariaux originaux
PATTERNS_SUBSTITUTION = [
    # Points de suspension simples
    (r'…+', 'VAR'),
    # Points multiples
    (r'\.{3,}', 'VAR'),
    # Underscores multiples
    (r'_{3,}', 'VAR'),
    # Texte entre parenthèses qui ressemble à une instruction
    (r'\(tous les prénoms et nom\)', 'NOM_COMPLET'),
    (r'\(profession\)', 'PROFESSION'),
    (r'\(adresse\)', 'ADRESSE'),
    (r'\(date\)', 'DATE'),
]


# =============================================================================
# SUBSTITUTION INTELLIGENTE
# =============================================================================

class SubstituteurActe:
    """
    Substitue les variables dans un document Word en préservant le formatage.
    """

    def __init__(self, variables: dict):
        self.variables = self._aplatir_variables(variables)
        self.substitutions_effectuees = []

    def _aplatir_variables(self, d: dict, parent_key: str = '') -> dict:
        """Aplatit un dictionnaire imbriqué en clés avec underscores."""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}_{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._aplatir_variables(v, new_key).items())
            else:
                items.append((new_key.lower(), v))
        return dict(items)

    def substituer_texte(self, texte: str, contexte: str = "") -> str:
        """
        Substitue les variables dans un texte.
        Utilise le contexte pour déterminer quelle variable utiliser.

        Patterns détectés dans les originaux:
        - … (ellipsis unicode U+2026) pour les valeurs à remplir
        - (tous les prénoms et nom), (profession), (adresse) pour les instructions
        """
        original = texte

        # Caractère ellipsis unicode
        ELLIPSIS = '\u2026'  # …

        # Si pas de pattern de variable, retourner tel quel
        if ELLIPSIS not in texte and '(tous' not in texte.lower() and '(profession)' not in texte.lower() and '(adresse)' not in texte.lower():
            return texte

        # Substitutions contextuelles basées sur le texte environnant
        contexte_lower = contexte.lower()

        # Déterminer la partie concernée par le contexte
        partie = None
        if 'vendeur' in contexte_lower or 'vend' in texte.lower():
            partie = 'vendeur'
        elif 'acquereur' in contexte_lower or 'acquéreur' in contexte_lower or 'acquiert' in texte.lower():
            partie = 'acquereur'

        # 1. Instructions entre parenthèses
        texte = self._substituer_instructions(texte, partie)

        # 2. Patterns avec ellipsis
        texte = self._substituer_ellipsis(texte, partie, contexte_lower)

        # 3. Prix
        if 'prix' in contexte_lower or 'euros' in texte.lower() or '€' in texte:
            texte = self._substituer_prix(texte)

        # 4. Date
        if 'le ' + ELLIPSIS in texte.lower() or f'le{ELLIPSIS}' in texte.lower():
            texte = self._substituer_date(texte)

        if texte != original:
            self.substitutions_effectuees.append((original[:50], texte[:50]))

        return texte

    def _substituer_instructions(self, texte: str, partie: str = None) -> str:
        """Substitue les instructions entre parenthèses."""
        prefix = partie or 'vendeur'

        # (tous les prénoms et nom)
        if '(tous les prénoms et nom)' in texte:
            nom = self.variables.get(f'{prefix}_nom', '...')
            prenoms = self.variables.get(f'{prefix}_prenoms', '...')
            nom_complet = f"{prenoms} {nom}".strip()
            texte = texte.replace('(tous les prénoms et nom)', nom_complet)

        # (profession)
        if '(profession)' in texte:
            profession = self.variables.get(f'{prefix}_profession', '...')
            texte = texte.replace('(profession)', profession)

        # (adresse)
        if '(adresse)' in texte:
            adresse = self.variables.get(f'{prefix}_adresse', '...')
            texte = texte.replace('(adresse)', adresse)

        return texte

    def _substituer_ellipsis(self, texte: str, partie: str = None, contexte: str = "") -> str:
        """Substitue les ellipsis (…) selon le contexte."""
        ELLIPSIS = '\u2026'
        prefix = partie or 'vendeur'

        # Madame … ou Monsieur …
        if 'Madame ' + ELLIPSIS in texte or f'Madame{ELLIPSIS}' in texte:
            nom = self.variables.get(f'{prefix}_nom', '...')
            prenoms = self.variables.get(f'{prefix}_prenoms', '...')
            texte = re.sub(rf'Madame\s*{ELLIPSIS}+', f'Madame {prenoms} {nom}', texte)

        if 'Monsieur ' + ELLIPSIS in texte or f'Monsieur{ELLIPSIS}' in texte:
            nom = self.variables.get(f'{prefix}_nom', '...')
            prenoms = self.variables.get(f'{prefix}_prenoms', '...')
            texte = re.sub(rf'Monsieur\s*{ELLIPSIS}+', f'Monsieur {prenoms} {nom}', texte)

        # Maître …
        if 'Maître ' + ELLIPSIS in texte or f'Maître{ELLIPSIS}' in texte:
            notaire_nom = self.variables.get('notaire_nom', '...')
            texte = re.sub(rf'Maître\s*{ELLIPSIS}+\.?,?', f'Maître {notaire_nom}', texte)

        # Le … (date)
        if re.search(rf'[Ll]e\s*{ELLIPSIS}', texte):
            date = self.variables.get('date_acte', '...')
            texte = re.sub(rf'[Ll]e\s*{ELLIPSIS}+\.?', f'Le {date}', texte)

        # Né(e) à … le …
        if 'à ' + ELLIPSIS in texte and 'le' + ELLIPSIS in texte:
            lieu = self.variables.get(f'{prefix}_lieu_naissance', '...')
            date_naissance = self.variables.get(f'{prefix}_date_naissance', '...')
            texte = re.sub(rf'à\s*{ELLIPSIS}+\s+le\s*{ELLIPSIS}+\.?', f'à {lieu} le {date_naissance}', texte)

        return texte

    def _substituer_partie(self, texte: str, partie: str) -> str:
        """Substitue les variables d'une partie (vendeur/acquéreur)."""
        prefix = partie

        # Nom complet
        if self.variables.get(f'{prefix}_nom'):
            nom = self.variables.get(f'{prefix}_nom', '')
            prenoms = self.variables.get(f'{prefix}_prenoms', '')
            nom_complet = f"{prenoms} {nom}".strip()

            texte = re.sub(r'\(tous les prénoms et nom\)', nom_complet, texte, flags=re.IGNORECASE)
            texte = re.sub(r'Madame\s*…+', f'Madame {nom_complet}', texte)
            texte = re.sub(r'Monsieur\s*…+', f'Monsieur {nom_complet}', texte)
            texte = re.sub(r'M\.\s*…+', f'M. {nom_complet}', texte)
            texte = re.sub(r'Mme\s*…+', f'Mme {nom_complet}', texte)

        # Profession
        if self.variables.get(f'{prefix}_profession'):
            texte = re.sub(r'\(profession\)', self.variables[f'{prefix}_profession'], texte, flags=re.IGNORECASE)

        # Adresse
        if self.variables.get(f'{prefix}_adresse'):
            texte = re.sub(r'\(adresse\)', self.variables[f'{prefix}_adresse'], texte, flags=re.IGNORECASE)

        # Date de naissance
        if self.variables.get(f'{prefix}_date_naissance'):
            texte = re.sub(r'Née?\s+à\s*…+\s+le\s*…+',
                          f"Né(e) à {self.variables.get(f'{prefix}_lieu_naissance', '...')} le {self.variables[f'{prefix}_date_naissance']}",
                          texte, flags=re.IGNORECASE)

        return texte

    def _substituer_prix(self, texte: str) -> str:
        """Substitue les variables de prix."""
        if self.variables.get('prix_montant'):
            montant = self.variables['prix_montant']
            if isinstance(montant, (int, float)):
                montant_str = f"{montant:,.0f}".replace(',', ' ')
            else:
                montant_str = str(montant)

            texte = re.sub(r'…+\s*(euros?|€)', f'{montant_str} euros', texte, flags=re.IGNORECASE)

        if self.variables.get('prix_lettres'):
            texte = re.sub(r'\(en lettres\)', self.variables['prix_lettres'], texte, flags=re.IGNORECASE)

        return texte

    def _substituer_date(self, texte: str) -> str:
        """Substitue les variables de date."""
        if self.variables.get('date_acte'):
            date = self.variables['date_acte']
            texte = re.sub(r'Le\s+…+', f'Le {date}', texte)

        return texte

    def _substituer_bien(self, texte: str) -> str:
        """Substitue les variables du bien."""
        if self.variables.get('bien_adresse'):
            texte = re.sub(r'\(adresse du bien\)', self.variables['bien_adresse'], texte, flags=re.IGNORECASE)

        if self.variables.get('bien_lot'):
            texte = re.sub(r'lot\s+n°?\s*…+', f"lot n° {self.variables['bien_lot']}", texte, flags=re.IGNORECASE)

        return texte

    def _substituer_generique(self, texte: str) -> str:
        """Substitution générique des patterns restants."""
        # Ne pas substituer si on n'a pas assez de contexte
        return texte


# =============================================================================
# COPIE PROFONDE DOCUMENT
# =============================================================================

def copier_document(doc_source: Document) -> Document:
    """
    Crée une copie profonde d'un document Word.
    Préserve tous les styles, formatages et structures.
    """
    # python-docx ne permet pas facilement de copier un document
    # On va donc travailler directement sur le document source
    # et sauvegarder sous un nouveau nom
    return doc_source


def substituer_dans_paragraphe(para, substitueur: SubstituteurActe, contexte: str = ""):
    """
    Substitue les variables dans un paragraphe en préservant le formatage de chaque run.
    """
    for run in para.runs:
        if run.text:
            nouveau_texte = substitueur.substituer_texte(run.text, contexte)
            if nouveau_texte != run.text:
                run.text = nouveau_texte


def substituer_dans_tableau(table, substitueur: SubstituteurActe):
    """
    Substitue les variables dans un tableau.
    """
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                substituer_dans_paragraphe(para, substitueur, "tableau")


def substituer_dans_document(doc: Document, substitueur: SubstituteurActe):
    """
    Parcourt tout le document et effectue les substitutions.
    """
    # Contexte courant (section, partie du document)
    contexte_courant = ""

    # Paragraphes
    for para in doc.paragraphs:
        # Mettre à jour le contexte basé sur les titres
        style_name = para.style.name if para.style else ""
        if 'Heading' in style_name:
            contexte_courant = para.text.lower()

        substituer_dans_paragraphe(para, substitueur, contexte_courant)

    # Tableaux
    for table in doc.tables:
        substituer_dans_tableau(table, substitueur)

    # En-têtes et pieds de page
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                substituer_dans_paragraphe(para, substitueur, "en-tete")
        if section.footer:
            for para in section.footer.paragraphs:
                substituer_dans_paragraphe(para, substitueur, "pied-page")


# =============================================================================
# GENERATION PRINCIPALE
# =============================================================================

def generer_acte(
    chemin_template: Path,
    variables: dict,
    chemin_sortie: Path,
    verbose: bool = False
) -> bool:
    """
    Génère un acte en copiant le template et en substituant les variables.

    Args:
        chemin_template: Chemin vers le DOCX template
        variables: Dictionnaire des variables à substituer
        chemin_sortie: Chemin de sortie pour l'acte généré
        verbose: Afficher les substitutions effectuées

    Returns:
        True si succès
    """
    print(f"[INFO] Chargement du template: {chemin_template}")

    # Charger le document template
    doc = Document(str(chemin_template))

    # Créer le substitueur
    substitueur = SubstituteurActe(variables)

    print(f"[INFO] Variables chargées: {len(substitueur.variables)}")

    # Effectuer les substitutions
    print("[INFO] Substitution des variables...")
    substituer_dans_document(doc, substitueur)

    # Sauvegarder
    chemin_sortie.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(chemin_sortie))

    print(f"[OK] Acte généré: {chemin_sortie}")
    print(f"     Substitutions: {len(substitueur.substitutions_effectuees)}")

    if verbose:
        print("\n[DETAIL] Substitutions effectuees:")
        for orig, nouveau in substitueur.substitutions_effectuees[:20]:
            # Nettoyer pour affichage console
            orig_clean = orig.encode('ascii', 'replace').decode('ascii')
            nouveau_clean = nouveau.encode('ascii', 'replace').decode('ascii')
            print(f"  '{orig_clean}' -> '{nouveau_clean}'")
        if len(substitueur.substitutions_effectuees) > 20:
            print(f"  ... et {len(substitueur.substitutions_effectuees) - 20} autres")

    return True


# =============================================================================
# VALIDATION VARIABLES
# =============================================================================

def valider_variables(variables: dict) -> list:
    """
    Valide que les variables essentielles sont présentes.
    Retourne la liste des erreurs.
    """
    erreurs = []

    # Variables obligatoires
    obligatoires = [
        ('vendeur', ['nom', 'prenoms']),
        ('acquereur', ['nom', 'prenoms']),
        ('prix', ['montant']),
    ]

    for section, champs in obligatoires:
        if section not in variables:
            erreurs.append(f"Section manquante: {section}")
        else:
            for champ in champs:
                if champ not in variables[section]:
                    erreurs.append(f"Champ manquant: {section}.{champ}")

    return erreurs


# =============================================================================
# EXEMPLE DE VARIABLES
# =============================================================================

EXEMPLE_VARIABLES = {
    "vendeur": {
        "nom": "MARTIN",
        "prenoms": "Jean Pierre",
        "profession": "Cadre commercial",
        "adresse": "25 rue Victor Hugo, 69002 Lyon",
        "date_naissance": "15/06/1970",
        "lieu_naissance": "Paris 12ème",
        "nationalite": "française",
        "situation_matrimoniale": "Célibataire"
    },
    "acquereur": {
        "nom": "DUBOIS",
        "prenoms": "Marie Claire",
        "profession": "Médecin",
        "adresse": "12 avenue Jean Jaurès, 69003 Lyon",
        "date_naissance": "22/09/1985",
        "lieu_naissance": "Lyon 6ème",
        "nationalite": "française",
        "situation_matrimoniale": "Célibataire"
    },
    "bien": {
        "adresse": "45 rue de la République, 69001 Lyon",
        "lot": "15",
        "batiment": "A",
        "etage": "3ème",
        "tantiemes": "250/10000",
        "surface": "75 m²"
    },
    "prix": {
        "montant": 320000,
        "lettres": "trois cent vingt mille"
    },
    "date_acte": "25 janvier 2025",
    "notaire": {
        "nom": "DIAZ",
        "prenom": "Charlotte",
        "etude": "Office Notarial des Cèdres",
        "adresse": "93 avenue du 11 Novembre 1918, Tassin-la-Demi-Lune"
    }
}


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Génère un acte notarial depuis un template DOCX"
    )
    parser.add_argument(
        "--template", "-t",
        type=Path,
        required=True,
        help="Chemin vers le DOCX template"
    )
    parser.add_argument(
        "--variables", "-v",
        type=Path,
        help="Fichier JSON contenant les variables"
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        required=True,
        help="Chemin de sortie pour l'acte généré"
    )
    parser.add_argument(
        "--exemple",
        action="store_true",
        help="Utiliser les variables d'exemple"
    )
    parser.add_argument(
        "--generer-exemple",
        type=Path,
        help="Générer un fichier JSON d'exemple de variables"
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Afficher les détails des substitutions"
    )

    args = parser.parse_args()

    # Générer fichier exemple
    if args.generer_exemple:
        with open(args.generer_exemple, 'w', encoding='utf-8') as f:
            json.dump(EXEMPLE_VARIABLES, f, indent=2, ensure_ascii=False)
        print(f"[OK] Fichier exemple généré: {args.generer_exemple}")
        return 0

    # Vérifier template
    if not args.template.exists():
        print(f"[ERREUR] Template introuvable: {args.template}")
        return 1

    # Charger variables
    if args.exemple:
        variables = EXEMPLE_VARIABLES
        print("[INFO] Utilisation des variables d'exemple")
    elif args.variables:
        if not args.variables.exists():
            print(f"[ERREUR] Fichier variables introuvable: {args.variables}")
            return 1
        with open(args.variables, 'r', encoding='utf-8') as f:
            variables = json.load(f)
    else:
        print("[ERREUR] Spécifiez --variables ou --exemple")
        return 1

    # Valider
    erreurs = valider_variables(variables)
    if erreurs:
        print("[ATTENTION] Variables manquantes:")
        for e in erreurs:
            print(f"  - {e}")

    # Générer
    try:
        generer_acte(args.template, variables, args.output, args.verbose)
        return 0
    except Exception as e:
        print(f"[ERREUR] {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
