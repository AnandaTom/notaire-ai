#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
comparer_documents_v2.py
------------------------
Version améliorée de la comparaison de documents DOCX avec scoring intelligent.

Améliorations v2:
- Scoring pondéré par importance fonctionnelle (H1=3pts, H2=2pts, H3/H4=1pt, Tableaux=5pts)
- Ignore les différences de hiérarchie si le titre existe (H1 vs H4 même titre = OK)
- Métrique "sections critiques présentes" pour valider les éléments MUST-HAVE
- Longueur du document compte pour 10% seulement (focus sur structure, pas volume)
- Conformité fonctionnelle plutôt que similarité structurelle brute

Usage:
    python comparer_documents_v2.py --original docs_original/trame.docx --genere outputs/acte.docx
    python comparer_documents_v2.py --original docs_original/trame.docx --genere outputs/acte.docx --rapport .tmp/rapport.json
"""

import argparse
import json
import re
import sys
from collections import Counter
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Set

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Erreur: python-docx n'est pas installé. Exécutez: pip install python-docx")
    sys.exit(1)

try:
    from rich.console import Console
    from rich.table import Table
    from rich.panel import Panel
    from rich.progress import track
except ImportError:
    print("Erreur: rich n'est pas installé. Exécutez: pip install rich")
    sys.exit(1)

# Fix encoding pour Windows
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# Configuration
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent

console = Console()

# Sections critiques MUST-HAVE pour les actes notariaux
SECTIONS_CRITIQUES = [
    "comparution",
    "désignation",
    "prix",
    "paiement",
    "origine de propriété",
    "charges et conditions",
    "déclarations",
    "état descriptif de division",
    "copropriété",
    "servitudes",
]

# Pondération des éléments pour le scoring
POIDS_H1 = 3
POIDS_H2 = 2
POIDS_H3 = 1
POIDS_H4 = 1
POIDS_TABLEAU = 5
POIDS_LONGUEUR = 0.1  # 10% seulement


@dataclass
class StructureDocument:
    """Structure extraite d'un document DOCX."""
    fichier: str
    nb_paragraphes: int
    nb_tableaux: int
    nb_sections: int
    titres: List[Dict[str, Any]]
    styles_utilises: Dict[str, int]
    longueur_totale: int
    tables_info: List[Dict[str, Any]]
    sections_detectees: List[str]
    sections_critiques_presentes: List[str]
    titres_par_niveau: Dict[int, int]


@dataclass
class ResultatComparaison:
    """Résultat de la comparaison entre deux documents."""
    conformite_globale: float  # Score de 0 à 100
    conformite_structure: float
    conformite_titres: float
    conformite_tableaux: float
    conformite_sections_critiques: float
    score_pondere: float
    differences: List[Dict[str, Any]]
    avertissements: List[str]
    recommandations: List[str]
    sections_critiques_manquantes: List[str]
    sections_critiques_presentes: List[str]


def extraire_structure(chemin_docx: Path) -> StructureDocument:
    """
    Extrait la structure d'un document DOCX.
    """
    doc = Document(str(chemin_docx))

    titres = []
    styles_utilises = Counter()
    longueur_totale = 0
    sections_detectees = []
    sections_critiques_presentes = []
    titres_par_niveau = {1: 0, 2: 0, 3: 0, 4: 0}

    # Analyse des paragraphes
    for i, para in enumerate(doc.paragraphs):
        texte = para.text.strip()
        style_name = para.style.name if para.style else "Normal"

        styles_utilises[style_name] += 1
        longueur_totale += len(texte)

        # Détection des titres
        niveau = None
        if style_name.startswith("Heading") or style_name.startswith("Titre"):
            if "1" in style_name or style_name in ["Heading 1", "Titre 1"]:
                niveau = 1
            elif "2" in style_name or style_name in ["Heading 2", "Titre 2"]:
                niveau = 2
            elif "3" in style_name or style_name in ["Heading 3", "Titre 3"]:
                niveau = 3
            elif "4" in style_name or style_name in ["Heading 4", "Titre 4"]:
                niveau = 4
            else:
                niveau = 3  # Par défaut

        if niveau:
            titres.append({
                "index": i,
                "niveau": niveau,
                "texte": texte[:100],  # Limite à 100 caractères
                "style": style_name
            })
            titres_par_niveau[niveau] += 1

        # Détection des sections par mots-clés
        texte_lower = texte.lower()

        # Détection des sections critiques
        for section_critique in SECTIONS_CRITIQUES:
            if section_critique in texte_lower:
                if section_critique not in sections_critiques_presentes:
                    sections_critiques_presentes.append(section_critique)

        if any(kw in texte_lower for kw in ["chapitre", "partie", "section", "titre"]):
            sections_detectees.append(texte[:80])

        # Détection par formatage (majuscules, gras)
        elif texte and texte.isupper() and len(texte) > 5:
            sections_detectees.append(texte[:80])

    # Analyse des tableaux
    tables_info = []
    for i, table in enumerate(doc.tables):
        nb_lignes = len(table.rows)
        nb_colonnes = len(table.columns)

        # Extrait les en-têtes si possible
        headers = []
        if nb_lignes > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip()[:30])

        tables_info.append({
            "index": i,
            "lignes": nb_lignes,
            "colonnes": nb_colonnes,
            "headers": headers
        })

    # Calcul du nombre de sections (basé sur les titres Heading 1)
    nb_sections = titres_par_niveau[1]

    return StructureDocument(
        fichier=str(chemin_docx.name),
        nb_paragraphes=len(doc.paragraphs),
        nb_tableaux=len(doc.tables),
        nb_sections=nb_sections,
        titres=titres,
        styles_utilises=dict(styles_utilises),
        longueur_totale=longueur_totale,
        tables_info=tables_info,
        sections_detectees=sections_detectees,
        sections_critiques_presentes=sections_critiques_presentes,
        titres_par_niveau=titres_par_niveau
    )


def normaliser_texte(texte: str) -> str:
    """Normalise un texte pour comparaison."""
    # Supprime les espaces multiples, la ponctuation, met en minuscules
    texte = texte.lower()
    texte = re.sub(r'\s+', ' ', texte)
    texte = re.sub(r'[^\w\s]', '', texte)
    return texte.strip()


def calculer_similarite_titres_v2(titres1: List[Dict], titres2: List[Dict]) -> Tuple[float, List[Dict], Dict]:
    """
    Compare les titres avec scoring pondéré intelligent.

    Améliorations:
    - H1 = 3 points, H2 = 2 points, H3/H4 = 1 point
    - Ignore les différences de hiérarchie si le titre existe
    - Retourne aussi les statistiques de matching
    """
    if not titres1 and not titres2:
        return 100.0, [], {}

    if not titres1 or not titres2:
        return 0.0, [{"type": "missing_titles", "message": "Un des documents n'a pas de titres détectés"}], {}

    differences = []
    points_gagnes = 0.0
    points_totaux = 0.0

    stats = {
        "h1_matches": 0,
        "h2_matches": 0,
        "h3_matches": 0,
        "h4_matches": 0,
        "total_matches": 0,
        "hierarchy_differences_ignored": 0
    }

    # Normalise les titres
    titres1_norm = [(normaliser_texte(t["texte"]), t["niveau"]) for t in titres1]
    titres2_norm = [(normaliser_texte(t["texte"]), t["niveau"]) for t in titres2]

    # Calcul des points totaux basé sur l'original
    for texte_norm, niveau in titres1_norm:
        if niveau == 1:
            points_totaux += POIDS_H1
        elif niveau == 2:
            points_totaux += POIDS_H2
        elif niveau == 3:
            points_totaux += POIDS_H3
        else:
            points_totaux += POIDS_H4

    # Cherche les correspondances
    matched_indices = set()
    for i, (t1_norm, niveau1) in enumerate(titres1_norm):
        matched = False
        niveau_different = False

        for j, (t2_norm, niveau2) in enumerate(titres2_norm):
            if j in matched_indices:
                continue

            # Similarité basique: texte identique ou contenu l'un dans l'autre
            if t1_norm == t2_norm or (len(t1_norm) > 3 and (t1_norm in t2_norm or t2_norm in t1_norm)):
                # Titre trouvé - ignore la différence de hiérarchie
                if niveau1 != niveau2:
                    niveau_different = True
                    stats["hierarchy_differences_ignored"] += 1

                matched = True
                matched_indices.add(j)

                # Attribution des points selon le niveau ORIGINAL
                if niveau1 == 1:
                    points_gagnes += POIDS_H1
                    stats["h1_matches"] += 1
                elif niveau1 == 2:
                    points_gagnes += POIDS_H2
                    stats["h2_matches"] += 1
                elif niveau1 == 3:
                    points_gagnes += POIDS_H3
                    stats["h3_matches"] += 1
                else:
                    points_gagnes += POIDS_H4
                    stats["h4_matches"] += 1

                stats["total_matches"] += 1

                if niveau_different:
                    differences.append({
                        "type": "titre_niveau_different",
                        "titre": titres1[i]["texte"][:60],
                        "niveau_original": niveau1,
                        "niveau_genere": niveau2,
                        "severite": "info"  # Info seulement, pas pénalisé
                    })

                break

            # Similarité par mots communs (pour titres longs)
            mots1 = set(t1_norm.split())
            mots2 = set(t2_norm.split())
            if mots1 and mots2 and len(mots1) > 2:
                intersection = len(mots1 & mots2)
                union = len(mots1 | mots2)
                if union > 0 and intersection / union > 0.6:  # 60% de similarité
                    matched = True
                    matched_indices.add(j)

                    # Match partiel = 80% des points
                    if niveau1 == 1:
                        points_gagnes += POIDS_H1 * 0.8
                        stats["h1_matches"] += 0.8
                    elif niveau1 == 2:
                        points_gagnes += POIDS_H2 * 0.8
                        stats["h2_matches"] += 0.8
                    elif niveau1 == 3:
                        points_gagnes += POIDS_H3 * 0.8
                        stats["h3_matches"] += 0.8
                    else:
                        points_gagnes += POIDS_H4 * 0.8
                        stats["h4_matches"] += 0.8

                    stats["total_matches"] += 0.8

                    differences.append({
                        "type": "titre_partiellement_different",
                        "titre_original": titres1[i]["texte"][:60],
                        "titre_genere": titres2[j]["texte"][:60],
                        "similarite": f"{intersection/union*100:.0f}%",
                        "severite": "warning"
                    })
                    break

        # Titre manquant
        if not matched and len(titres1[i]["texte"]) > 3:
            differences.append({
                "type": "titre_manquant_genere",
                "titre_original": titres1[i]["texte"][:60],
                "niveau": niveau1,
                "severite": "error"
            })

    # Titres en trop dans le document généré (moins grave)
    for j, (t2_norm, niveau2) in enumerate(titres2_norm):
        if j not in matched_indices and len(titres2[j]["texte"]) > 3:
            differences.append({
                "type": "titre_supplementaire",
                "titre_genere": titres2[j]["texte"][:60],
                "niveau": niveau2,
                "severite": "info"
            })

    # Score basé sur les points pondérés
    score = (points_gagnes / points_totaux * 100) if points_totaux > 0 else 100.0
    return min(score, 100.0), differences, stats


def calculer_similarite_tableaux_v2(tables1: List[Dict], tables2: List[Dict]) -> Tuple[float, List[Dict]]:
    """
    Compare les tableaux avec pondération forte (5 points par tableau).
    """
    if not tables1 and not tables2:
        return 100.0, []

    differences = []
    nb1 = len(tables1)
    nb2 = len(tables2)

    if nb1 != nb2:
        differences.append({
            "type": "nombre_tableaux_different",
            "original": nb1,
            "genere": nb2,
            "severite": "error" if abs(nb1 - nb2) > 2 else "warning"
        })

    # Points pondérés
    points_gagnes = 0.0
    points_totaux = nb1 * POIDS_TABLEAU

    # Compare les tableaux existants
    for i in range(min(nb1, nb2)):
        t1 = tables1[i]
        t2 = tables2[i]

        points_tableau = POIDS_TABLEAU
        penalite = 0

        if t1["colonnes"] != t2["colonnes"]:
            penalite += 0.3
            differences.append({
                "type": "colonnes_differentes",
                "tableau": i + 1,
                "original": t1["colonnes"],
                "genere": t2["colonnes"],
                "severite": "warning"
            })

        # Compare les headers
        h1_norm = [normaliser_texte(h) for h in t1["headers"]]
        h2_norm = [normaliser_texte(h) for h in t2["headers"]]

        if h1_norm != h2_norm:
            # Vérifier si au moins 50% des headers correspondent
            matches = sum(1 for h in h1_norm if h in h2_norm)
            total = max(len(h1_norm), 1)
            if matches / total < 0.5:
                penalite += 0.3
                differences.append({
                    "type": "headers_tres_differents",
                    "tableau": i + 1,
                    "original": t1["headers"],
                    "genere": t2["headers"],
                    "severite": "warning"
                })
            else:
                penalite += 0.1
                differences.append({
                    "type": "headers_legerement_differents",
                    "tableau": i + 1,
                    "severite": "info"
                })

        # Attribution des points avec pénalité
        points_gagnes += points_tableau * (1 - min(penalite, 1.0))

    # Score basé sur les points pondérés
    score = (points_gagnes / points_totaux * 100) if points_totaux > 0 else 100.0
    return min(score, 100.0), differences


def verifier_sections_critiques(struct_original: StructureDocument, struct_genere: StructureDocument) -> Tuple[float, List[str], List[str]]:
    """
    Vérifie la présence des sections critiques MUST-HAVE.

    Returns:
        Tuple (score, sections_manquantes, sections_presentes)
    """
    sections_orig = set([s.lower() for s in struct_original.sections_critiques_presentes])
    sections_gen = set([s.lower() for s in struct_genere.sections_critiques_presentes])

    sections_critiques_lower = [s.lower() for s in SECTIONS_CRITIQUES]

    # Sections critiques présentes dans l'original
    sections_orig_critiques = sections_orig.intersection(sections_critiques_lower)

    # Sections critiques présentes dans le généré
    sections_gen_critiques = sections_gen.intersection(sections_critiques_lower)

    # Calcul du score
    if len(sections_orig_critiques) == 0:
        # Si l'original n'a pas de sections critiques détectées, on ne peut pas évaluer
        score = 100.0
        manquantes = []
        presentes = list(sections_gen_critiques)
    else:
        # Score = % de sections critiques de l'original présentes dans le généré
        manquantes = list(sections_orig_critiques - sections_gen_critiques)
        presentes = list(sections_gen_critiques)
        score = (len(sections_gen_critiques) / len(sections_orig_critiques)) * 100

    return score, manquantes, presentes


def comparer_documents_v2(original: Path, genere: Path) -> ResultatComparaison:
    """
    Compare deux documents DOCX avec scoring intelligent v2.
    """
    console.print(f"[cyan]Extraction de la structure de l'original...[/cyan]")
    struct_original = extraire_structure(original)

    console.print(f"[cyan]Extraction de la structure du document généré...[/cyan]")
    struct_genere = extraire_structure(genere)

    differences = []
    avertissements = []
    recommandations = []

    # Comparaison de la longueur (10% seulement)
    ratio_longueur = min(struct_genere.longueur_totale, struct_original.longueur_totale) / \
                     max(struct_genere.longueur_totale, struct_original.longueur_totale) \
                     if max(struct_genere.longueur_totale, struct_original.longueur_totale) > 0 else 1.0

    score_longueur = ratio_longueur * 100

    if ratio_longueur < 0.5:
        avertissements.append(
            f"Le document généré est significativement {'plus court' if struct_genere.longueur_totale < struct_original.longueur_totale else 'plus long'} "
            f"({struct_genere.longueur_totale} vs {struct_original.longueur_totale} caractères)"
        )

    # Comparaison des titres avec scoring pondéré v2
    score_titres, diff_titres, stats_titres = calculer_similarite_titres_v2(
        struct_original.titres,
        struct_genere.titres
    )
    differences.extend(diff_titres)

    # Comparaison des tableaux avec pondération forte
    score_tableaux, diff_tableaux = calculer_similarite_tableaux_v2(
        struct_original.tables_info,
        struct_genere.tables_info
    )
    differences.extend(diff_tableaux)

    # Vérification des sections critiques
    score_sections_critiques, sections_manquantes, sections_presentes = verifier_sections_critiques(
        struct_original,
        struct_genere
    )

    if sections_manquantes:
        avertissements.append(
            f"Sections critiques manquantes: {', '.join(sections_manquantes)}"
        )

    # Score pondéré global
    # Priorité: Titres (40%) > Tableaux (30%) > Sections critiques (20%) > Longueur (10%)
    score_pondere = (
        score_titres * 0.40 +
        score_tableaux * 0.30 +
        score_sections_critiques * 0.20 +
        score_longueur * 0.10
    )

    # Score de structure (basé sur la présence des éléments clés)
    score_structure = (
        (struct_genere.nb_tableaux / max(struct_original.nb_tableaux, 1) * 50 if struct_original.nb_tableaux > 0 else 50) +
        (len(struct_genere.titres) / max(len(struct_original.titres), 1) * 50 if struct_original.titres else 50)
    )
    score_structure = min(score_structure, 100)

    # Génération des recommandations
    if score_titres < 80:
        recommandations.append("Vérifier que tous les titres de section sont présents dans le template")

        # Recommandations spécifiques par niveau
        for niveau in [1, 2, 3, 4]:
            if stats_titres.get(f"h{niveau}_matches", 0) < struct_original.titres_par_niveau.get(niveau, 0) * 0.8:
                recommandations.append(f"Manque de titres H{niveau} (seulement {stats_titres.get(f'h{niveau}_matches', 0):.0f} sur {struct_original.titres_par_niveau.get(niveau, 0)})")

    if score_tableaux < 80:
        recommandations.append("Vérifier la structure des tableaux (colonnes, en-têtes)")

    if score_sections_critiques < 80:
        recommandations.append("Vérifier la présence des sections critiques obligatoires")
        for section in sections_manquantes:
            recommandations.append(f"  → Ajouter la section: {section}")

    if ratio_longueur < 0.7:
        recommandations.append("Le document généré semble incomplet. Vérifier que toutes les sections sont remplies.")

    # Avertissement si différences de hiérarchie
    if stats_titres.get("hierarchy_differences_ignored", 0) > 0:
        avertissements.append(
            f"{stats_titres['hierarchy_differences_ignored']} titre(s) ont une hiérarchie différente (H1 vs H2, etc.) "
            f"mais le contenu correspond. Pas de pénalité appliquée."
        )

    return ResultatComparaison(
        conformite_globale=round(score_pondere, 1),
        conformite_structure=round(score_structure, 1),
        conformite_titres=round(score_titres, 1),
        conformite_tableaux=round(score_tableaux, 1),
        conformite_sections_critiques=round(score_sections_critiques, 1),
        score_pondere=round(score_pondere, 1),
        differences=differences,
        avertissements=avertissements,
        recommandations=recommandations,
        sections_critiques_manquantes=sections_manquantes,
        sections_critiques_presentes=sections_presentes
    )


def afficher_rapport_v2(resultat: ResultatComparaison, struct_original: StructureDocument, struct_genere: StructureDocument) -> None:
    """Affiche un rapport de comparaison formaté v2."""

    # Score global
    couleur_score = "green" if resultat.conformite_globale >= 80 else "yellow" if resultat.conformite_globale >= 60 else "red"

    console.print()
    console.print(Panel(
        f"[bold {couleur_score}]Score de conformité fonctionnelle: {resultat.conformite_globale}%[/bold {couleur_score}]\n"
        f"[dim](Scoring pondéré: Titres 40% | Tableaux 30% | Sections critiques 20% | Longueur 10%)[/dim]",
        title="RÉSULTAT v2 - Scoring Intelligent",
        border_style=couleur_score
    ))

    # Détail des scores
    table = Table(title="Détail des scores pondérés", show_header=True)
    table.add_column("Critère", style="cyan")
    table.add_column("Score", justify="right")
    table.add_column("Poids", justify="right")
    table.add_column("Statut", justify="center")

    for critere, score, poids in [
        ("Titres (H1=3pts, H2=2pts, H3/H4=1pt)", resultat.conformite_titres, "40%"),
        ("Tableaux (5pts chacun)", resultat.conformite_tableaux, "30%"),
        ("Sections critiques MUST-HAVE", resultat.conformite_sections_critiques, "20%"),
        ("Longueur document", round((resultat.score_pondere - resultat.conformite_titres*0.4 - resultat.conformite_tableaux*0.3 - resultat.conformite_sections_critiques*0.2)/0.1, 1), "10%")
    ]:
        if score >= 80:
            statut = "[green]✓[/green]"
        elif score >= 60:
            statut = "[yellow]~[/yellow]"
        else:
            statut = "[red]✗[/red]"
        table.add_row(critere, f"{score}%", poids, statut)

    console.print(table)

    # Sections critiques
    console.print()
    console.print("[bold]Sections critiques MUST-HAVE:[/bold]")

    sections_table = Table(show_header=True)
    sections_table.add_column("Section", style="cyan")
    sections_table.add_column("Statut", justify="center")

    for section in SECTIONS_CRITIQUES:
        if section in resultat.sections_critiques_presentes:
            sections_table.add_row(section.title(), "[green]✓ Présente[/green]")
        elif section in resultat.sections_critiques_manquantes:
            sections_table.add_row(section.title(), "[red]✗ Manquante[/red]")
        else:
            sections_table.add_row(section.title(), "[dim]- Non détectée[/dim]")

    console.print(sections_table)

    # Statistiques comparatives
    console.print()
    console.print("[bold]Statistiques comparatives:[/bold]")

    stats_table = Table(show_header=True)
    stats_table.add_column("Métrique", style="cyan")
    stats_table.add_column("Original", justify="right")
    stats_table.add_column("Généré", justify="right")
    stats_table.add_column("Écart", justify="right")

    stats_table.add_row(
        "Titres H1 (3pts chacun)",
        str(struct_original.titres_par_niveau.get(1, 0)),
        str(struct_genere.titres_par_niveau.get(1, 0)),
        str(struct_genere.titres_par_niveau.get(1, 0) - struct_original.titres_par_niveau.get(1, 0))
    )
    stats_table.add_row(
        "Titres H2 (2pts chacun)",
        str(struct_original.titres_par_niveau.get(2, 0)),
        str(struct_genere.titres_par_niveau.get(2, 0)),
        str(struct_genere.titres_par_niveau.get(2, 0) - struct_original.titres_par_niveau.get(2, 0))
    )
    stats_table.add_row(
        "Titres H3/H4 (1pt chacun)",
        str(struct_original.titres_par_niveau.get(3, 0) + struct_original.titres_par_niveau.get(4, 0)),
        str(struct_genere.titres_par_niveau.get(3, 0) + struct_genere.titres_par_niveau.get(4, 0)),
        str((struct_genere.titres_par_niveau.get(3, 0) + struct_genere.titres_par_niveau.get(4, 0)) -
            (struct_original.titres_par_niveau.get(3, 0) + struct_original.titres_par_niveau.get(4, 0)))
    )
    stats_table.add_row(
        "Tableaux (5pts chacun)",
        str(struct_original.nb_tableaux),
        str(struct_genere.nb_tableaux),
        str(struct_genere.nb_tableaux - struct_original.nb_tableaux)
    )
    stats_table.add_row(
        "Longueur (10% poids)",
        f"{struct_original.longueur_totale:,} car.",
        f"{struct_genere.longueur_totale:,} car.",
        f"{struct_genere.longueur_totale - struct_original.longueur_totale:,}"
    )

    console.print(stats_table)

    # Différences par sévérité
    if resultat.differences:
        console.print()

        errors = [d for d in resultat.differences if d.get("severite") == "error"]
        warnings = [d for d in resultat.differences if d.get("severite") == "warning"]
        infos = [d for d in resultat.differences if d.get("severite") == "info"]

        if errors:
            console.print(f"[bold red]Erreurs critiques ({len(errors)}):[/bold red]")
            for diff in errors[:10]:
                if diff["type"] == "titre_manquant_genere":
                    console.print(f"  [red]✗ Titre manquant:[/red] {diff['titre_original']}")
                elif diff["type"] == "nombre_tableaux_different":
                    console.print(f"  [red]✗ Tableaux:[/red] {diff['original']} attendus, {diff['genere']} trouvés")
                else:
                    console.print(f"  [red]✗[/red] {diff}")

        if warnings:
            console.print(f"\n[bold yellow]Avertissements ({len(warnings)}):[/bold yellow]")
            for diff in warnings[:10]:
                if diff["type"] == "titre_partiellement_different":
                    console.print(f"  [yellow]⚠ Titre similaire à {diff.get('similarite', '?')}:[/yellow] {diff.get('titre_original', '')[:50]}")
                elif diff["type"] == "colonnes_differentes":
                    console.print(f"  [yellow]⚠ Tableau {diff['tableau']}:[/yellow] {diff['original']} vs {diff['genere']} colonnes")
                else:
                    console.print(f"  [yellow]⚠[/yellow] {diff.get('type', 'unknown')}")

        if infos:
            console.print(f"\n[bold blue]Informations ({len(infos)}):[/bold blue]")
            for diff in infos[:5]:
                if diff["type"] == "titre_niveau_different":
                    console.print(f"  [blue]ℹ Hiérarchie différente (ignoré):[/blue] H{diff['niveau_original']} → H{diff['niveau_genere']}")
                elif diff["type"] == "titre_supplementaire":
                    console.print(f"  [blue]+ Titre supplémentaire:[/blue] {diff.get('titre_genere', '')[:50]}")
                else:
                    console.print(f"  [blue]ℹ[/blue] {diff.get('type', 'unknown')}")

            if len(infos) > 5:
                console.print(f"  [dim]... et {len(infos) - 5} autres informations[/dim]")

    # Avertissements
    if resultat.avertissements:
        console.print()
        console.print("[bold yellow]Avertissements:[/bold yellow]")
        for avert in resultat.avertissements:
            console.print(f"  [yellow]⚠[/yellow] {avert}")

    # Recommandations
    if resultat.recommandations:
        console.print()
        console.print("[bold blue]Recommandations:[/bold blue]")
        for reco in resultat.recommandations:
            console.print(f"  [blue]→[/blue] {reco}")


def main():
    """Point d'entrée principal."""
    parser = argparse.ArgumentParser(
        description="Compare la structure d'un document généré avec un original (v2 - Scoring intelligent)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples:
  python comparer_documents_v2.py --original docs_original/trame.docx --genere outputs/acte.docx
  python comparer_documents_v2.py -o docs_original/trame.docx -g outputs/acte.docx --rapport .tmp/rapport.json

Améliorations v2:
  - Scoring pondéré: H1=3pts, H2=2pts, H3/H4=1pt, Tableaux=5pts
  - Ignore différences de hiérarchie si titre existe
  - Métrique "sections critiques présentes"
  - Longueur compte pour 10% seulement
  - Focus sur conformité fonctionnelle vs similarité structurelle
        """
    )

    parser.add_argument(
        "--original", "-o",
        required=True,
        type=str,
        help="Chemin vers le document original (trame)"
    )

    parser.add_argument(
        "--genere", "-g",
        required=True,
        type=str,
        help="Chemin vers le document généré"
    )

    parser.add_argument(
        "--rapport", "-r",
        type=str,
        help="Fichier de sortie pour le rapport JSON (optionnel)"
    )

    parser.add_argument(
        "--seuil",
        type=float,
        default=80.0,
        help="Seuil de conformité minimum (défaut: 80%%)"
    )

    args = parser.parse_args()

    # Vérifie les fichiers
    original_path = Path(args.original)
    genere_path = Path(args.genere)

    if not original_path.exists():
        console.print(f"[red]Fichier original non trouvé: {original_path}[/red]")
        sys.exit(1)

    if not genere_path.exists():
        console.print(f"[red]Fichier généré non trouvé: {genere_path}[/red]")
        sys.exit(1)

    # Affiche l'en-tête
    console.print()
    console.print(Panel(
        f"[bold white]COMPARAISON DE DOCUMENTS v2 - Scoring Intelligent[/bold white]\n\n"
        f"[cyan]Original:[/cyan] {original_path.name}\n"
        f"[cyan]Généré:[/cyan] {genere_path.name}\n\n"
        f"[dim]Pondération: Titres 40% | Tableaux 30% | Sections critiques 20% | Longueur 10%[/dim]",
        border_style="blue"
    ))

    # Extraction des structures
    struct_original = extraire_structure(original_path)
    struct_genere = extraire_structure(genere_path)

    # Comparaison
    resultat = comparer_documents_v2(original_path, genere_path)

    # Affichage du rapport
    afficher_rapport_v2(resultat, struct_original, struct_genere)

    # Export JSON si demandé
    if args.rapport:
        rapport_path = Path(args.rapport)
        rapport_path.parent.mkdir(parents=True, exist_ok=True)

        rapport_data = {
            "version": "2.0",
            "scoring": {
                "titres_poids": "40%",
                "tableaux_poids": "30%",
                "sections_critiques_poids": "20%",
                "longueur_poids": "10%",
                "ponderations": {
                    "h1": POIDS_H1,
                    "h2": POIDS_H2,
                    "h3": POIDS_H3,
                    "h4": POIDS_H4,
                    "tableau": POIDS_TABLEAU
                }
            },
            "original": asdict(struct_original),
            "genere": asdict(struct_genere),
            "resultat": asdict(resultat)
        }

        with open(rapport_path, "w", encoding="utf-8") as f:
            json.dump(rapport_data, f, ensure_ascii=False, indent=2)

        console.print(f"\n[green]Rapport exporté: {rapport_path}[/green]")

    # Code de retour basé sur le seuil
    if resultat.conformite_globale >= args.seuil:
        console.print(f"\n[green]✓ Conformité fonctionnelle atteinte ({resultat.conformite_globale}% >= {args.seuil}%)[/green]")
        sys.exit(0)
    else:
        console.print(f"\n[red]✗ Conformité fonctionnelle insuffisante ({resultat.conformite_globale}% < {args.seuil}%)[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main()
