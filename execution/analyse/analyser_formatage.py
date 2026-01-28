#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analyser_formatage_complet.py
=============================

Analyse ULTRA-DÉTAILLÉE d'un document Word pour reproduction fidèle.
Extrait CHAQUE détail de formatage: paragraphes, runs, tableaux, styles, etc.

Usage:
    python analyser_formatage_complet.py "fichier.docx" -o rapport.json
    python analyser_formatage_complet.py "fichier.docx" --text  # Affiche le texte structuré
"""

import argparse
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# CONVERSIONS
# =============================================================================

def emu_to_cm(emu):
    """Convertit EMU (English Metric Units) en cm."""
    if emu is None:
        return None
    return round(emu / 914400 * 2.54, 3)

def emu_to_pt(emu):
    """Convertit EMU en points."""
    if emu is None:
        return None
    return round(emu / 12700, 2)

def twips_to_pt(twips):
    """Convertit twips en points (20 twips = 1 pt)."""
    if twips is None:
        return None
    return round(twips / 20, 2)

def safe_value(val):
    """Extrait une valeur de manière sécurisée."""
    if val is None:
        return None
    if hasattr(val, 'pt'):
        return round(val.pt, 2)
    if hasattr(val, 'cm'):
        return round(val.cm, 3)
    if hasattr(val, 'emu'):
        return val.emu
    return val


# =============================================================================
# ANALYSE SECTIONS
# =============================================================================

def analyser_section(section, index):
    """Analyse complète d'une section."""
    info = {
        "numero": index + 1,
        "page": {
            "largeur_cm": round(section.page_width.cm, 2) if section.page_width else None,
            "hauteur_cm": round(section.page_height.cm, 2) if section.page_height else None,
            "orientation": "portrait" if section.page_width and section.page_height and section.page_width < section.page_height else "paysage",
        },
        "marges": {
            "gauche_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "droite_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
            "haut_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bas_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "reliure_cm": round(section.gutter.cm, 2) if section.gutter else None,
        },
        "en_tete_pied": {
            "distance_en_tete_cm": round(section.header_distance.cm, 2) if section.header_distance else None,
            "distance_pied_cm": round(section.footer_distance.cm, 2) if section.footer_distance else None,
            "en_tete_different_premiere_page": section.different_first_page_header_footer,
        },
    }

    # Vérifier marges miroir via XML
    sectPr = section._sectPr
    if sectPr is not None:
        mirror = sectPr.find(qn('w:mirrorMargins'))
        info["marges"]["miroir"] = mirror is not None

        # Colonnes
        cols = sectPr.find(qn('w:cols'))
        if cols is not None:
            info["colonnes"] = {
                "nombre": cols.get(qn('w:num')),
                "espacement": cols.get(qn('w:space')),
            }

    return info


# =============================================================================
# ANALYSE STYLES
# =============================================================================

def analyser_style_complet(style):
    """Analyse complète d'un style."""
    info = {
        "nom": style.name,
        "id": style.style_id,
        "type": str(style.type).split('.')[-1] if style.type else None,
        "base_sur": style.base_style.name if style.base_style else None,
        "style_suivant": style.next_paragraph_style.name if hasattr(style, 'next_paragraph_style') and style.next_paragraph_style else None,
        "integre": style.builtin,
        "masque": style.hidden,
    }

    # Police
    if style.font:
        f = style.font
        info["police"] = {
            "nom": f.name,
            "taille_pt": safe_value(f.size),
            "gras": f.bold,
            "italique": f.italic,
            "souligne": bool(f.underline) if f.underline is not None else None,
            "type_souligne": str(f.underline) if f.underline else None,
            "barre": f.strike,
            "double_barre": f.double_strike,
            "majuscules": f.all_caps,
            "petites_majuscules": f.small_caps,
            "exposant": f.superscript,
            "indice": f.subscript,
            "masque": f.hidden,
            "couleur_rgb": str(f.color.rgb) if f.color and f.color.rgb else None,
            "couleur_theme": str(f.color.theme_color) if f.color and f.color.theme_color else None,
            "surlignage": str(f.highlight_color) if f.highlight_color else None,
        }

    # Paragraphe (seulement pour styles paragraphe)
    if style.type == WD_STYLE_TYPE.PARAGRAPH and style.paragraph_format:
        pf = style.paragraph_format
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "gauche",
            WD_ALIGN_PARAGRAPH.CENTER: "centre",
            WD_ALIGN_PARAGRAPH.RIGHT: "droite",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justifie",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribue",
        }

        line_spacing_map = {
            WD_LINE_SPACING.SINGLE: "simple",
            WD_LINE_SPACING.ONE_POINT_FIVE: "1.5",
            WD_LINE_SPACING.DOUBLE: "double",
            WD_LINE_SPACING.AT_LEAST: "au_moins",
            WD_LINE_SPACING.EXACTLY: "exactement",
            WD_LINE_SPACING.MULTIPLE: "multiple",
        }

        info["paragraphe"] = {
            "alignement": align_map.get(pf.alignment, str(pf.alignment) if pf.alignment else None),
            "retrait_gauche_cm": round(pf.left_indent.cm, 3) if pf.left_indent else None,
            "retrait_droite_cm": round(pf.right_indent.cm, 3) if pf.right_indent else None,
            "retrait_premiere_ligne_cm": round(pf.first_line_indent.cm, 3) if pf.first_line_indent else None,
            "espace_avant_pt": safe_value(pf.space_before),
            "espace_apres_pt": safe_value(pf.space_after),
            "interligne_regle": line_spacing_map.get(pf.line_spacing_rule, str(pf.line_spacing_rule) if pf.line_spacing_rule else None),
            "interligne_valeur": pf.line_spacing,
            "conserver_avec_suivant": pf.keep_with_next,
            "conserver_ensemble": pf.keep_together,
            "saut_page_avant": pf.page_break_before,
            "veuve_orphelin": pf.widow_control,
        }

        # Tabulations
        if pf.tab_stops:
            tabs = []
            for tab in pf.tab_stops:
                tabs.append({
                    "position_cm": round(tab.position.cm, 3) if tab.position else None,
                    "alignement": str(tab.alignment) if tab.alignment else None,
                    "caractere_remplissage": str(tab.leader) if tab.leader else None,
                })
            if tabs:
                info["paragraphe"]["tabulations"] = tabs

    return info


# =============================================================================
# ANALYSE PARAGRAPHES
# =============================================================================

def analyser_run(run, index):
    """Analyse complète d'un run (segment de texte)."""
    info = {
        "index": index,
        "texte": run.text,
        "longueur": len(run.text),
    }

    # Formatage direct
    f = run.font
    formatage = {}

    if f.name:
        formatage["police"] = f.name
    if f.size:
        formatage["taille_pt"] = safe_value(f.size)
    if f.bold is not None:
        formatage["gras"] = f.bold
    if f.italic is not None:
        formatage["italique"] = f.italic
    if f.underline is not None:
        formatage["souligne"] = bool(f.underline)
    if f.all_caps is not None:
        formatage["majuscules"] = f.all_caps
    if f.small_caps is not None:
        formatage["petites_majuscules"] = f.small_caps
    if f.strike is not None:
        formatage["barre"] = f.strike
    if f.subscript is not None:
        formatage["indice"] = f.subscript
    if f.superscript is not None:
        formatage["exposant"] = f.superscript
    if f.color and f.color.rgb:
        formatage["couleur"] = str(f.color.rgb)

    if formatage:
        info["formatage"] = formatage

    return info


def analyser_paragraphe(para, index):
    """Analyse complète d'un paragraphe."""
    info = {
        "index": index,
        "texte": para.text,
        "longueur": len(para.text),
        "style": para.style.name if para.style else None,
    }

    # Formatage direct du paragraphe
    pf = para.paragraph_format
    if pf:
        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "gauche",
            WD_ALIGN_PARAGRAPH.CENTER: "centre",
            WD_ALIGN_PARAGRAPH.RIGHT: "droite",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justifie",
        }

        format_para = {}
        if pf.alignment is not None:
            format_para["alignement"] = align_map.get(pf.alignment, str(pf.alignment))
        if pf.left_indent:
            format_para["retrait_gauche_cm"] = round(pf.left_indent.cm, 3)
        if pf.right_indent:
            format_para["retrait_droite_cm"] = round(pf.right_indent.cm, 3)
        if pf.first_line_indent:
            format_para["retrait_premiere_ligne_cm"] = round(pf.first_line_indent.cm, 3)
        if pf.space_before:
            format_para["espace_avant_pt"] = safe_value(pf.space_before)
        if pf.space_after:
            format_para["espace_apres_pt"] = safe_value(pf.space_after)
        if pf.line_spacing:
            format_para["interligne"] = pf.line_spacing

        if format_para:
            info["format_direct"] = format_para

    # Analyser chaque run
    runs = []
    for i, run in enumerate(para.runs):
        if run.text:  # Ignorer les runs vides
            runs.append(analyser_run(run, i))

    if runs:
        info["runs"] = runs

    return info


# =============================================================================
# ANALYSE TABLEAUX
# =============================================================================

def analyser_cellule(cell, row_idx, col_idx):
    """Analyse complète d'une cellule de tableau."""
    info = {
        "position": f"[{row_idx},{col_idx}]",
        "texte": cell.text[:100] + "..." if len(cell.text) > 100 else cell.text,
    }

    # Dimensions
    if cell.width:
        info["largeur_cm"] = round(cell.width.cm, 2)

    # Alignement vertical
    valign_map = {
        WD_CELL_VERTICAL_ALIGNMENT.TOP: "haut",
        WD_CELL_VERTICAL_ALIGNMENT.CENTER: "centre",
        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bas",
    }
    if cell.vertical_alignment:
        info["alignement_vertical"] = valign_map.get(cell.vertical_alignment, str(cell.vertical_alignment))

    # Paragraphes dans la cellule
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.style:
            info["style_paragraphe"] = para.style.name
        if para.paragraph_format.alignment:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "gauche",
                WD_ALIGN_PARAGRAPH.CENTER: "centre",
                WD_ALIGN_PARAGRAPH.RIGHT: "droite",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justifie",
            }
            info["alignement_texte"] = align_map.get(para.paragraph_format.alignment)

    return info


def analyser_tableau(table, index):
    """Analyse complète d'un tableau."""
    info = {
        "index": index,
        "lignes": len(table.rows),
        "colonnes": len(table.columns),
        "style": table.style.name if table.style else None,
    }

    # Alignement du tableau
    align_map = {
        WD_TABLE_ALIGNMENT.LEFT: "gauche",
        WD_TABLE_ALIGNMENT.CENTER: "centre",
        WD_TABLE_ALIGNMENT.RIGHT: "droite",
    }
    if table.alignment:
        info["alignement"] = align_map.get(table.alignment, str(table.alignment))

    # Largeurs des colonnes
    largeurs = []
    for col in table.columns:
        if col.width:
            largeurs.append(round(col.width.cm, 2))
    if largeurs:
        info["largeurs_colonnes_cm"] = largeurs

    # Contenu détaillé
    contenu = []
    for i, row in enumerate(table.rows):
        ligne = []
        for j, cell in enumerate(row.cells):
            ligne.append(analyser_cellule(cell, i, j))
        contenu.append(ligne)
    info["contenu"] = contenu

    return info


# =============================================================================
# ANALYSE EN-TÊTES ET PIEDS DE PAGE
# =============================================================================

def analyser_header_footer(hf, type_name):
    """Analyse un en-tête ou pied de page."""
    if not hf or not hf.paragraphs:
        return None

    info = {
        "type": type_name,
        "paragraphes": [],
    }

    for i, para in enumerate(hf.paragraphs):
        if para.text.strip():
            p_info = {
                "texte": para.text,
                "style": para.style.name if para.style else None,
            }
            if para.paragraph_format.alignment:
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "gauche",
                    WD_ALIGN_PARAGRAPH.CENTER: "centre",
                    WD_ALIGN_PARAGRAPH.RIGHT: "droite",
                }
                p_info["alignement"] = align_map.get(para.paragraph_format.alignment)
            info["paragraphes"].append(p_info)

    return info if info["paragraphes"] else None


# =============================================================================
# EXPORT TEXTE STRUCTURÉ
# =============================================================================

def exporter_texte_structure(doc):
    """Exporte le document en texte structuré avec marqueurs de formatage."""
    output = []

    output.append("=" * 80)
    output.append("DOCUMENT WORD - ANALYSE STRUCTURÉE")
    output.append("=" * 80)
    output.append("")

    # Sections
    for i, section in enumerate(doc.sections):
        s = analyser_section(section, i)
        output.append(f"[SECTION {i+1}]")
        output.append(f"  Page: {s['page']['largeur_cm']} x {s['page']['hauteur_cm']} cm ({s['page']['orientation']})")
        output.append(f"  Marges: G={s['marges']['gauche_cm']}cm D={s['marges']['droite_cm']}cm H={s['marges']['haut_cm']}cm B={s['marges']['bas_cm']}cm")
        if s['marges'].get('miroir'):
            output.append("  Marges miroir: OUI")
        output.append("")

    # Paragraphes
    output.append("-" * 80)
    output.append("CONTENU")
    output.append("-" * 80)
    output.append("")

    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        # Style et formatage
        style_info = f"[{para.style.name}]" if para.style else "[?]"

        # Alignement
        if para.paragraph_format.alignment:
            align_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "←",
                WD_ALIGN_PARAGRAPH.CENTER: "↔",
                WD_ALIGN_PARAGRAPH.RIGHT: "→",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "⇔",
            }
            align = align_map.get(para.paragraph_format.alignment, "?")
        else:
            align = " "

        # Retraits
        indent_info = ""
        if para.paragraph_format.first_line_indent:
            indent_cm = round(para.paragraph_format.first_line_indent.cm, 2)
            indent_info = f" [indent:{indent_cm}cm]"

        # Formatage des runs
        formatted_text = ""
        for run in para.runs:
            text = run.text
            if not text:
                continue

            # Marqueurs de formatage
            prefix = ""
            suffix = ""
            if run.bold:
                prefix += "**"
                suffix = "**" + suffix
            if run.italic:
                prefix += "*"
                suffix = "*" + suffix
            if run.underline:
                prefix += "__"
                suffix = "__" + suffix
            if run.font.all_caps:
                text = text.upper()
                prefix += "[CAPS]"
            if run.font.small_caps:
                prefix += "[sc]"
                suffix = "[/sc]" + suffix

            formatted_text += prefix + text + suffix

        # Ligne de sortie
        output.append(f"P{i:03d} {align} {style_info}{indent_info}")
        output.append(f"     {formatted_text[:200]}{'...' if len(formatted_text) > 200 else ''}")
        output.append("")

    # Tableaux
    if doc.tables:
        output.append("-" * 80)
        output.append("TABLEAUX")
        output.append("-" * 80)
        output.append("")

        for i, table in enumerate(doc.tables):
            output.append(f"[TABLEAU {i+1}] {len(table.rows)} lignes x {len(table.columns)} colonnes")
            for row_idx, row in enumerate(table.rows):
                cells_text = [cell.text[:30].replace('\n', ' ') for cell in row.cells]
                output.append(f"  L{row_idx}: | " + " | ".join(cells_text) + " |")
            output.append("")

    return "\n".join(output)


# =============================================================================
# ANALYSE PRINCIPALE
# =============================================================================

def analyser_document_complet(chemin_docx):
    """Analyse COMPLÈTE d'un document Word."""
    doc = Document(chemin_docx)

    analyse = {
        "fichier": str(chemin_docx),
        "meta": {
            "nb_paragraphes": len(doc.paragraphs),
            "nb_tableaux": len(doc.tables),
            "nb_sections": len(doc.sections),
        },
        "sections": [],
        "styles": {
            "paragraphe": [],
            "caractere": [],
        },
        "paragraphes": [],
        "tableaux": [],
        "en_tetes_pieds": [],
    }

    # Sections
    for i, section in enumerate(doc.sections):
        analyse["sections"].append(analyser_section(section, i))

        # En-têtes et pieds de page
        header = analyser_header_footer(section.header, f"en_tete_section_{i+1}")
        if header:
            analyse["en_tetes_pieds"].append(header)
        footer = analyser_header_footer(section.footer, f"pied_section_{i+1}")
        if footer:
            analyse["en_tetes_pieds"].append(footer)

    # Styles
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            analyse["styles"]["paragraphe"].append(analyser_style_complet(style))
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            analyse["styles"]["caractere"].append(analyser_style_complet(style))

    # Paragraphes (limiter pour fichier raisonnable)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Ignorer paragraphes vides
            analyse["paragraphes"].append(analyser_paragraphe(para, i))

    # Tableaux
    for i, table in enumerate(doc.tables):
        analyse["tableaux"].append(analyser_tableau(table, i))

    return analyse


def main():
    parser = argparse.ArgumentParser(description="Analyse COMPLÈTE d'un document Word")
    parser.add_argument("fichier", help="Chemin vers le fichier .docx")
    parser.add_argument("-o", "--output", help="Fichier JSON de sortie")
    parser.add_argument("--text", action="store_true", help="Afficher en texte structuré")
    parser.add_argument("--styles-only", action="store_true", help="Afficher seulement les styles")
    parser.add_argument("--sections-only", action="store_true", help="Afficher seulement les sections")

    args = parser.parse_args()

    chemin = Path(args.fichier)
    if not chemin.exists():
        print(f"Erreur: Fichier introuvable: {chemin}")
        return 1

    print(f"Analyse de: {chemin}")
    print("=" * 60)

    doc = Document(chemin)

    if args.text:
        print(exporter_texte_structure(doc))
        return 0

    if args.styles_only:
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                s = analyser_style_complet(style)
                if s.get("police") or s.get("paragraphe"):
                    print(f"\n{s['nom']}:")
                    if s.get("police"):
                        p = s["police"]
                        attrs = []
                        if p.get("gras"): attrs.append("gras")
                        if p.get("italique"): attrs.append("italique")
                        if p.get("souligne"): attrs.append("souligné")
                        if p.get("majuscules"): attrs.append("MAJUSCULES")
                        if p.get("petites_majuscules"): attrs.append("petites-caps")
                        print(f"  Police: {p.get('nom')} {p.get('taille_pt')}pt {' '.join(attrs)}")
                    if s.get("paragraphe"):
                        pg = s["paragraphe"]
                        print(f"  Alignement: {pg.get('alignement')}")
                        if pg.get("retrait_premiere_ligne_cm"):
                            print(f"  Retrait 1ère ligne: {pg['retrait_premiere_ligne_cm']}cm")
                        if pg.get("espace_avant_pt"):
                            print(f"  Espace avant: {pg['espace_avant_pt']}pt")
                        if pg.get("espace_apres_pt"):
                            print(f"  Espace après: {pg['espace_apres_pt']}pt")
        return 0

    if args.sections_only:
        for i, section in enumerate(doc.sections):
            s = analyser_section(section, i)
            print(f"\nSection {i+1}:")
            print(f"  Page: {s['page']['largeur_cm']} x {s['page']['hauteur_cm']} cm")
            print(f"  Marges: G={s['marges']['gauche_cm']}cm D={s['marges']['droite_cm']}cm H={s['marges']['haut_cm']}cm B={s['marges']['bas_cm']}cm")
            print(f"  Reliure: {s['marges']['reliure_cm']}cm")
            print(f"  Marges miroir: {s['marges'].get('miroir', False)}")
        return 0

    # Analyse complète
    analyse = analyser_document_complet(chemin)

    # Résumé
    print(f"\nRésumé:")
    print(f"  Paragraphes: {analyse['meta']['nb_paragraphes']}")
    print(f"  Tableaux: {analyse['meta']['nb_tableaux']}")
    print(f"  Sections: {analyse['meta']['nb_sections']}")
    print(f"  Styles paragraphe: {len(analyse['styles']['paragraphe'])}")

    # Sauvegarder si demandé
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(analyse, f, indent=2, ensure_ascii=False)
        print(f"\nAnalyse sauvegardée: {args.output}")
    else:
        # Afficher JSON compact
        print("\n" + json.dumps(analyse, indent=2, ensure_ascii=False)[:5000] + "\n...")
        print("\nUtilisez -o fichier.json pour sauvegarder l'analyse complète")

    return 0


if __name__ == "__main__":
    exit(main())
