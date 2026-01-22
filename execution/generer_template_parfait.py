#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_template_parfait.py
============================

Génère un template PARFAIT avec TOUT le contenu du DOCX original
en remplaçant les valeurs spécifiques par des variables Jinja2.
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
import re
from docx import Document
from pathlib import Path

def generer_variable_jinja(texte_original, contexte_ligne):
    """Génère une variable Jinja2 appropriée selon le contexte"""

    # Mapping intelligent selon le contexte
    mappings = [
        # Nombres et références
        (r'\b100583001\b', '{{ reference_acte }}'),
        (r'\bCDI/ALG\b', '{{ initiales_notaire }}'),

        # Dates complètes
        (r'23 décembre 2025', '{{ date_acte.day }} {{ date_acte.month|mois_en_lettres }} {{ date_acte.year }}'),
        (r'DEUX MILLE VINGT CINQ', '{{ date_acte.year|nombre_en_lettres|upper }}'),
        (r'VINGT TROIS DÉCEMBRE', '{{ date_acte.day|jour_en_lettres|upper }} {{ date_acte.month|mois_en_lettres|upper }}'),

        # Noms donateurs
        (r'\bDominique Christian AUVRAY\b', '{{ donateur_1.prenom }} {{ donateur_1.deuxieme_prenom }} {{ donateur_1.nom|upper }}'),
        (r'\bDominique AUVRAY\b', '{{ donateur_1.prenom }} {{ donateur_1.nom|upper }}'),
        (r'\bViviane Blanche Célie ROBIN\b', '{{ donateur_2.prenom }} {{ donateur_2.deuxieme_prenom }} {{ donateur_2.troisieme_prenom }} {{ donateur_2.nom_naissance|upper }}'),
        (r'\bViviane AUVRAY née ROBIN\b', '{{ donateur_2.prenom }} {{ donateur_2.nom|upper }} née {{ donateur_2.nom_naissance|upper }}'),
        (r'\bViviane ROBIN\b', '{{ donateur_2.prenom }} {{ donateur_2.nom_naissance|upper }}'),

        # Noms donataires
        (r'\bAntoine Christian Edmond AUVRAY\b', '{{ donataires[0].prenom }} {{ donataires[0].deuxieme_prenom }} {{ donataires[0].troisieme_prenom }} {{ donataires[0].nom|upper }}'),
        (r'\bAntoine AUVRAY\b', '{{ donataires[0].prenom }} {{ donataires[0].nom|upper }}'),
        (r'\bMathilde Yvonne Nicole AUVRAY\b', '{{ donataires[1].prenom }} {{ donataires[1].deuxieme_prenom }} {{ donataires[1].troisieme_prenom }} {{ donataires[1].nom|upper }}'),
        (r'\bMathilde AUVRAY\b', '{{ donataires[1].prenom }} {{ donataires[1].nom|upper }}'),

        # Lieux donateurs
        (r'\bDARDILLY \(69570\)', '{{ donateur_1.adresse.ville|upper }} ({{ donateur_1.adresse.code_postal }})'),
        (r'29 chemin du [Pp]anorama', '{{ donateur_1.adresse.numero }} {{ donateur_1.adresse.voie }}'),

        # Lieux donataires
        (r'\bLYON 3ÈME ARRONDISSEMENT \(69003\)', '{{ donataires[0].adresse.ville|upper }} ({{ donataires[0].adresse.code_postal }})'),
        (r'16 cours Lafayette', '{{ donataires[0].adresse.numero }} {{ donataires[0].adresse.voie }}'),
        (r'\bPARIS 9ÈME ARRONDISSEMENT \(75009\)', '{{ donataires[1].adresse.ville|upper }} ({{ donataires[1].adresse.code_postal }})'),
        (r'18 rue Milton', '{{ donataires[1].adresse.numero }} {{ donataires[1].adresse.voie }}'),

        # Dates de naissance
        (r'20 février 1961', '{{ donateur_1.naissance.date.day }} {{ donateur_1.naissance.date.month|mois_en_lettres }} {{ donateur_1.naissance.date.year }}'),
        (r'16 avril 1961', '{{ donateur_2.naissance.date.day }} {{ donateur_2.naissance.date.month|mois_en_lettres }} {{ donateur_2.naissance.date.year }}'),
        (r'19 février 1991', '{{ donataires[0].naissance.date.day }} {{ donataires[0].naissance.date.month|mois_en_lettres }} {{ donataires[0].naissance.date.year }}'),
        (r'9 août 1993', '{{ donataires[1].naissance.date.day }} {{ donataires[1].naissance.date.month|mois_en_lettres }} {{ donataires[1].naissance.date.year }}'),

        # Lieux de naissance
        (r'\bPARIS 12ÈME ARRONDISSEMENT \(75012\)', '{{ donateur_1.naissance.ville|upper }} ({{ donateur_1.naissance.code_postal }})'),
        (r'\bTAIN-L\'HERMITAGE \(26600\)', '{{ donateur_2.naissance.ville|upper }} ({{ donateur_2.naissance.code_postal }})'),
        (r'\bLYON 9ÈME ARRONDISSEMENT \(69009\)', 'LYON 9ÈME ARRONDISSEMENT (69009)'),  # Commun aux 2 donataires

        # Notaire
        (r'\bCharlotte DIAZ\b', '{{ notaire.prenom }} {{ notaire.nom|upper }}'),
        (r'\bdes Cèdres\b', '{{ notaire.nom_societe }}'),
        (r'\bTASSIN-LA-DEMI-LUNE \(Rhône\)', '{{ notaire.office.ville|upper }} ({{ notaire.office.departement }})'),
        (r'93 avenue du 11 Novembre 1918', '{{ notaire.office.adresse }}'),
        (r'\b69182\b', '{{ notaire.crpcen }}'),

        # Mariage
        (r'\bCHANTEMERLE-LES-BLES \(26600\)', '{{ mariage.ville|upper }} ({{ mariage.code_postal }})'),
        (r'9 septembre 1989', '{{ mariage.date.day }} {{ mariage.date.month|mois_en_lettres }} {{ mariage.date.year }}'),
        (r'\bPierre COHENDET\b', '{{ mariage.notaire.prenom }} {{ mariage.notaire.nom|upper }}'),
        (r'\bCHAROLLES \(71120\)', '{{ mariage.notaire.ville|upper }} ({{ mariage.notaire.code_postal }})'),
        (r'24 juin 1989', '{{ mariage.contrat.date.day }} {{ mariage.contrat.date.month|mois_en_lettres }} {{ mariage.contrat.date.year }}'),

        # Société
        (r'\bBLOUGE\b', '{{ societe.nom }}'),
        (r'\b985 331 354\b', '{{ societe.siren }}'),
        (r'19 février 2024', '{{ societe.constitution.date.day }} {{ societe.constitution.date.month|mois_en_lettres }} {{ societe.constitution.date.year }}'),

        # Montants
        (r'QUATRE CENT MILLE EUROS', '{{ donation_anterieure.valeur|nombre_en_lettres|upper }} EUROS'),
        (r'400 000,00 EUR', '{{ donation_anterieure.valeur|format_montant }} EUR'),
        (r'MILLE EUROS \(1 000', '{{ societe.capital|nombre_en_lettres|upper }} EUROS ({{ societe.capital|format_montant }}'),

        # Professions
        (r'\bretraité\b', '{{ donateur_1.profession }}'),
        (r'\bretraitée\b', '{{ donateur_2.profession }}'),
        (r'\bingénieur référent agrivoltaïque\b', '{{ donataires[0].profession }}'),
        (r'\bgérante d\'une entreprise de conseil\b', '{{ donataires[1].profession }}'),
    ]

    texte = texte_original
    for pattern, replacement in mappings:
        texte = re.sub(pattern, replacement, texte)

    return texte

def extraire_et_convertir():
    """Extrait le DOCX et crée un template Jinja2 complet"""

    doc = Document('docs_originels/Donation partage (2).docx')

    template_lines = []

    # En-tête page 1 (spécial)
    template_lines.append('{FIRST_PAGE_HEADER_START}')
    template_lines.append('{{ reference_acte }}')
    template_lines.append('{{ initiales_notaire }}/')
    template_lines.append('')
    template_lines.append('Le {{ date_acte.day }} {{ date_acte.month|mois_en_lettres }} {{ date_acte.year }}')
    template_lines.append('{FIRST_PAGE_HEADER_END}')
    template_lines.append('')
    template_lines.append('**DONATION-PARTAGE**')
    template_lines.append('**Par {{ donateur_1.civilite }} {{ donateur_1.prenom }} {{ donateur_1.nom|upper }}**')
    template_lines.append('**Et {{ donateur_2.civilite }} {{ donateur_2.prenom }} {{ donateur_2.nom|upper }} née {{ donateur_2.nom_naissance|upper }}**')
    template_lines.append('**Au profit de leurs {{ nombre_enfants_lettres }} enfants**')
    template_lines.append('***************************************************************')
    template_lines.append('')

    # Convertir tous les paragraphes (en sautant les 3 premiers déjà traités)
    para_a_ignorer = 2  # Les 2 premiers paragraphes (ref et date/titre)
    compteur = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            template_lines.append('')
            continue

        compteur += 1
        if compteur <= para_a_ignorer:
            continue

        # Détecter les encadrés (titres en majuscules)
        is_titre_encadre = False
        if text in ['ELEMENTS PREALABLES', 'TERMINOLOGIE', 'DECLARATIONS DES PARTIES',
                    'DOCUMENTS RELATIFS A LA CAPACITE ET A LA QUALITE DES PARTIES',
                    'EXPOSE PREALABLE']:
            is_titre_encadre = True
            nom_encadre = text.replace(' ', '_').replace('À', 'A')
            template_lines.append(f'{{{nom_encadre}_BOX_START}}')
            template_lines.append(f'## {text}')
            template_lines.append(f'{{{nom_encadre}_BOX_END}}')
            template_lines.append('')
            continue

        # Convertir le texte
        text_converti = generer_variable_jinja(text, i)

        # Détecter les titres (gras + majuscules)
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        is_upper = text.isupper() and len(text) > 5

        if is_bold and is_upper:
            template_lines.append(f'## {text_converti}')
        elif is_bold:
            template_lines.append(f'**{text_converti}**')
        else:
            template_lines.append(text_converti)

        template_lines.append('')

    # Sauvegarder
    output_path = Path('templates/donation_partage_parfait.md')
    output_path.write_text('\n'.join(template_lines), encoding='utf-8')

    print(f'✅ Template parfait créé: {output_path}')
    print(f'   Lignes: {len(template_lines)}')
    print(f'   Paragraphes DOCX: {len(doc.paragraphs)}')

if __name__ == '__main__':
    extraire_et_convertir()
