#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extraire_docx_vers_markdown.py
==============================

Extrait le contenu COMPLET d'un document Word vers Markdown structuré.
Préserve: styles, formatage, tableaux, structure des paragraphes.

Usage:
    python extraire_docx_vers_markdown.py "document.docx" -o "document.md"
    python extraire_docx_vers_markdown.py "document.docx" --variables  # Détecte les variables
"""

import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# =============================================================================
# PATTERNS DE VARIABLES (texte à remplacer)
# =============================================================================

# Patterns typiques de variables dans les actes notariaux
PATTERNS_VARIABLES = [
    # Points de suspension (3 ou plus)
    (r'\.{3,}', '{{VARIABLE}}'),
    # Texte entre crochets
    (r'\[([^\]]+)\]', r'{{VAR_\1}}'),
    # Dates incomplètes
    (r'\b(\d{1,2})/(\d{1,2})/____', '{{DATE}}'),
    (r'\b____/____/____', '{{DATE}}'),
    # Montants avec blancs
    (r'\b____+\s*(euros?|€)', '{{MONTANT}} euros'),
    # Noms en majuscules avec blancs
    (r'M\.\s*____+', 'M. {{NOM}}'),
    (r'Mme\s*____+', 'Mme {{NOM}}'),
    # Références vides
    (r'n°\s*____+', 'n° {{REFERENCE}}'),
]


def detecter_variables(texte: str) -> str:
    """Détecte et marque les variables dans le texte."""
    for pattern, remplacement in PATTERNS_VARIABLES:
        texte = re.sub(pattern, remplacement, texte)
    return texte


# =============================================================================
# CONVERSION FORMATAGE
# =============================================================================

def convertir_run_en_markdown(run) -> str:
    """Convertit un run Word en Markdown avec formatage."""
    texte = run.text
    if not texte:
        return ""

    # Appliquer les formatages
    if run.bold and run.italic:
        texte = f"***{texte}***"
    elif run.bold:
        texte = f"**{texte}**"
    elif run.italic:
        texte = f"*{texte}*"

    # Soulignement (utiliser __texte__ comme convention)
    if run.underline:
        texte = f"__{texte}__"

    return texte


def convertir_paragraphe_en_markdown(para, detect_vars: bool = False) -> tuple:
    """
    Convertit un paragraphe Word en Markdown.
    Retourne: (texte_markdown, style_name, est_titre)
    """
    style_name = para.style.name if para.style else "Normal"

    # Construire le texte avec formatage
    texte_parts = []
    for run in para.runs:
        texte_parts.append(convertir_run_en_markdown(run))

    texte = "".join(texte_parts).strip()

    if not texte:
        return ("", style_name, False)

    # Détecter les variables si demandé
    if detect_vars:
        texte = detecter_variables(texte)

    # Convertir selon le style
    est_titre = False

    if style_name == "Heading 1":
        texte = f"# {texte}"
        est_titre = True
    elif style_name == "Heading 2":
        texte = f"## {texte}"
        est_titre = True
    elif style_name == "Heading 3":
        texte = f"### {texte}"
        est_titre = True
    elif style_name == "Heading 4":
        texte = f"#### {texte}"
        est_titre = True
    elif style_name == "Heading 5":
        texte = f"##### {texte}"
        est_titre = True
    elif style_name in ["Quote", "Citation"]:
        texte = f"> {texte}"
    elif style_name in ["Enumération", "Enumeration", "List Paragraph"]:
        # Garder tel quel, sera traité comme liste
        pass
    elif style_name in ["Chapitre"]:
        texte = f"# {texte}"
        est_titre = True

    return (texte, style_name, est_titre)


def convertir_tableau_en_markdown(table) -> str:
    """Convertit un tableau Word en tableau Markdown."""
    if not table.rows:
        return ""

    lignes_md = []
    nb_cols = len(table.columns)

    for i, row in enumerate(table.rows):
        cellules = []
        for cell in row.cells:
            # Extraire le texte de la cellule (peut avoir plusieurs paragraphes)
            cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cell_text = cell_text.replace("|", "\\|")  # Échapper les pipes
            cellules.append(cell_text)

        # Assurer le bon nombre de colonnes
        while len(cellules) < nb_cols:
            cellules.append("")

        ligne = "| " + " | ".join(cellules[:nb_cols]) + " |"
        lignes_md.append(ligne)

        # Ajouter la ligne de séparation après l'en-tête
        if i == 0:
            sep = "| " + " | ".join(["---"] * nb_cols) + " |"
            lignes_md.append(sep)

    return "\n".join(lignes_md)


# =============================================================================
# EXTRACTION PRINCIPALE
# =============================================================================

def extraire_document(chemin_docx: Path, detect_vars: bool = False) -> str:
    """
    Extrait tout le contenu d'un document Word vers Markdown.
    """
    doc = Document(str(chemin_docx))

    # Collecter tous les éléments (paragraphes et tableaux) dans l'ordre
    # Note: python-docx ne permet pas facilement de mélanger paragraphes et tableaux
    # On va donc traiter séquentiellement

    output_lines = []

    # En-tête du fichier
    output_lines.append(f"<!-- Extrait de: {chemin_docx.name} -->")
    output_lines.append(f"<!-- Paragraphes: {len(doc.paragraphs)}, Tableaux: {len(doc.tables)} -->")
    output_lines.append("")

    # Suivre les tableaux déjà traités
    tableaux_traites = set()

    # Parcourir le document via l'élément XML pour garder l'ordre
    body = doc.element.body

    for element in body:
        # Paragraphe
        if element.tag.endswith('}p'):
            # Trouver le paragraphe correspondant
            for para in doc.paragraphs:
                if para._element is element:
                    texte, style, est_titre = convertir_paragraphe_en_markdown(para, detect_vars)
                    if texte:
                        output_lines.append(texte)
                        # Ajouter une ligne vide après les titres
                        if est_titre:
                            output_lines.append("")
                    break

        # Tableau
        elif element.tag.endswith('}tbl'):
            for i, table in enumerate(doc.tables):
                if table._element is element and i not in tableaux_traites:
                    tableaux_traites.add(i)
                    tableau_md = convertir_tableau_en_markdown(table)
                    if tableau_md:
                        output_lines.append("")
                        output_lines.append(tableau_md)
                        output_lines.append("")
                    break

    # Nettoyer les lignes vides multiples
    resultat = "\n".join(output_lines)
    resultat = re.sub(r'\n{3,}', '\n\n', resultat)

    return resultat


def extraire_avec_sections(chemin_docx: Path, detect_vars: bool = False) -> str:
    """
    Extrait le document en ajoutant des marqueurs de section.
    """
    doc = Document(str(chemin_docx))

    output_lines = []
    output_lines.append(f"<!-- Template extrait de: {chemin_docx.name} -->")
    output_lines.append(f"<!-- Date extraction: AUTO -->")
    output_lines.append("")

    # Sections de document notarial typiques
    section_courante = None
    sections_detectees = []

    body = doc.element.body

    for element in body:
        if element.tag.endswith('}p'):
            for para in doc.paragraphs:
                if para._element is element:
                    texte, style, est_titre = convertir_paragraphe_en_markdown(para, detect_vars)

                    # Détecter les changements de section majeure
                    if style == "Heading 1" and texte:
                        texte_clean = texte.replace("#", "").strip()
                        section_id = re.sub(r'[^a-zA-Z0-9]', '_', texte_clean.lower())
                        output_lines.append(f"\n<!-- SECTION: {section_id} | OBLIGATOIRE -->")
                        sections_detectees.append(section_id)

                    if texte:
                        output_lines.append(texte)
                        if est_titre:
                            output_lines.append("")
                    break

        elif element.tag.endswith('}tbl'):
            for i, table in enumerate(doc.tables):
                if table._element is element:
                    tableau_md = convertir_tableau_en_markdown(table)
                    if tableau_md:
                        output_lines.append("")
                        output_lines.append(tableau_md)
                        output_lines.append("")
                    break

    resultat = "\n".join(output_lines)
    resultat = re.sub(r'\n{3,}', '\n\n', resultat)

    # Ajouter un résumé des sections en début
    if sections_detectees:
        header = "<!-- SECTIONS DETECTEES:\n"
        for s in sections_detectees:
            header += f"  - {s}\n"
        header += "-->\n\n"
        resultat = header + resultat

    return resultat


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Extrait un document Word vers Markdown structuré"
    )
    parser.add_argument("fichier", help="Chemin vers le fichier .docx")
    parser.add_argument("-o", "--output", help="Fichier Markdown de sortie")
    parser.add_argument(
        "--variables", "-v",
        action="store_true",
        help="Détecter et marquer les variables ({{...}})"
    )
    parser.add_argument(
        "--sections", "-s",
        action="store_true",
        help="Ajouter des marqueurs de section"
    )

    args = parser.parse_args()

    chemin = Path(args.fichier)
    if not chemin.exists():
        print(f"Erreur: Fichier introuvable: {chemin}")
        return 1

    print(f"Extraction de: {chemin}")

    # Extraire
    if args.sections:
        contenu = extraire_avec_sections(chemin, args.variables)
    else:
        contenu = extraire_document(chemin, args.variables)

    # Sauvegarder ou afficher
    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(contenu)
        print(f"Markdown sauvegardé: {output_path}")

        # Stats
        nb_lignes = len(contenu.split('\n'))
        nb_sections = contenu.count('<!-- SECTION:')
        nb_variables = contenu.count('{{')
        print(f"  Lignes: {nb_lignes}")
        print(f"  Sections: {nb_sections}")
        print(f"  Variables: {nb_variables}")
    else:
        print(contenu)

    return 0


if __name__ == "__main__":
    exit(main())
