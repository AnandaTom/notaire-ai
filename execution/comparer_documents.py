#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
comparer_documents.py
---------------------
Compare la structure d'un document DOCX généré avec un document original
pour valider la conformité du template.

Analyse:
- Structure des titres et sections
- Styles appliqués
- Nombre de paragraphes
- Tableaux présents
- Éléments de mise en forme

Usage:
    python comparer_documents.py --original docs_originels/trame.docx --genere outputs/acte.docx
    python comparer_documents.py --original docs_originels/trame.docx --genere outputs/acte.docx --rapport .tmp/rapport.json
"""

import argparse
import json
import re
import sys
from collections import Counter
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Erreur: python-docx n'est pas installé. Exécutez: pip install python-docx")
    sys.exit(1)

try:
    from rich.console import Console
    from rich.table import Table
    from rich.panel import Panel
    from rich.progress import track
except ImportError:
    print("Erreur: rich n'est pas installé. Exécutez: pip install rich")
    sys.exit(1)

# Configuration
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent

console = Console()


@dataclass
class StructureDocument:
    """Structure extraite d'un document DOCX."""
    fichier: str
    nb_paragraphes: int
    nb_tableaux: int
    nb_sections: int
    titres: List[Dict[str, Any]]
    styles_utilises: Dict[str, int]
    longueur_totale: int
    tables_info: List[Dict[str, Any]]
    sections_detectees: List[str]


@dataclass
class ResultatComparaison:
    """Résultat de la comparaison entre deux documents."""
    conformite_globale: float  # Score de 0 à 100
    conformite_structure: float
    conformite_titres: float
    conformite_tableaux: float
    differences: List[Dict[str, Any]]
    avertissements: List[str]
    recommandations: List[str]


def extraire_structure(chemin_docx: Path) -> StructureDocument:
    """
    Extrait la structure d'un document DOCX.
    """
    doc = Document(str(chemin_docx))

    titres = []
    styles_utilises = Counter()
    longueur_totale = 0
    sections_detectees = []

    # Analyse des paragraphes
    for i, para in enumerate(doc.paragraphs):
        texte = para.text.strip()
        style_name = para.style.name if para.style else "Normal"

        styles_utilises[style_name] += 1
        longueur_totale += len(texte)

        # Détection des titres
        if style_name.startswith("Heading") or style_name.startswith("Titre"):
            niveau = 1
            if "1" in style_name:
                niveau = 1
            elif "2" in style_name:
                niveau = 2
            elif "3" in style_name:
                niveau = 3

            titres.append({
                "index": i,
                "niveau": niveau,
                "texte": texte[:100],  # Limite à 100 caractères
                "style": style_name
            })

        # Détection des sections par mots-clés
        texte_lower = texte.lower()
        if any(kw in texte_lower for kw in ["chapitre", "partie", "section", "titre"]):
            sections_detectees.append(texte[:80])

        # Détection par formatage (majuscules, gras)
        elif texte and texte.isupper() and len(texte) > 5:
            sections_detectees.append(texte[:80])

    # Analyse des tableaux
    tables_info = []
    for i, table in enumerate(doc.tables):
        nb_lignes = len(table.rows)
        nb_colonnes = len(table.columns)

        # Extrait les en-têtes si possible
        headers = []
        if nb_lignes > 0:
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip()[:30])

        tables_info.append({
            "index": i,
            "lignes": nb_lignes,
            "colonnes": nb_colonnes,
            "headers": headers
        })

    # Calcul du nombre de sections (basé sur les titres Heading 1)
    nb_sections = len([t for t in titres if t["niveau"] == 1])

    return StructureDocument(
        fichier=str(chemin_docx.name),
        nb_paragraphes=len(doc.paragraphs),
        nb_tableaux=len(doc.tables),
        nb_sections=nb_sections,
        titres=titres,
        styles_utilises=dict(styles_utilises),
        longueur_totale=longueur_totale,
        tables_info=tables_info,
        sections_detectees=sections_detectees
    )


def normaliser_texte(texte: str) -> str:
    """Normalise un texte pour comparaison."""
    # Supprime les espaces multiples, la ponctuation, met en minuscules
    texte = texte.lower()
    texte = re.sub(r'\s+', ' ', texte)
    texte = re.sub(r'[^\w\s]', '', texte)
    return texte.strip()


def calculer_similarite_titres(titres1: List[Dict], titres2: List[Dict]) -> Tuple[float, List[Dict]]:
    """
    Compare les titres de deux documents et retourne un score de similarité.
    """
    if not titres1 and not titres2:
        return 100.0, []

    if not titres1 or not titres2:
        return 0.0, [{"type": "missing_titles", "message": "Un des documents n'a pas de titres détectés"}]

    differences = []
    matches = 0
    total = max(len(titres1), len(titres2))

    # Normalise les titres
    titres1_norm = [normaliser_texte(t["texte"]) for t in titres1]
    titres2_norm = [normaliser_texte(t["texte"]) for t in titres2]

    # Cherche les correspondances
    matched_indices = set()
    for i, t1 in enumerate(titres1_norm):
        for j, t2 in enumerate(titres2_norm):
            if j in matched_indices:
                continue
            # Similarité basique: texte identique ou contenu l'un dans l'autre
            if t1 == t2 or t1 in t2 or t2 in t1:
                matches += 1
                matched_indices.add(j)
                break
            # Similarité par mots communs
            mots1 = set(t1.split())
            mots2 = set(t2.split())
            if mots1 and mots2:
                intersection = len(mots1 & mots2)
                union = len(mots1 | mots2)
                if union > 0 and intersection / union > 0.5:
                    matches += 0.7  # Match partiel
                    matched_indices.add(j)
                    break

    # Titres manquants dans le document généré
    for i, t1 in enumerate(titres1):
        found = False
        t1_norm = normaliser_texte(t1["texte"])
        for t2_norm in titres2_norm:
            if t1_norm == t2_norm or t1_norm in t2_norm or t2_norm in t1_norm:
                found = True
                break
        if not found and len(t1["texte"]) > 3:
            differences.append({
                "type": "titre_manquant_genere",
                "titre_original": t1["texte"],
                "niveau": t1["niveau"]
            })

    # Titres en trop dans le document généré
    for j, t2 in enumerate(titres2):
        if j not in matched_indices:
            differences.append({
                "type": "titre_supplementaire",
                "titre_genere": t2["texte"],
                "niveau": t2["niveau"]
            })

    score = (matches / total) * 100 if total > 0 else 100.0
    return min(score, 100.0), differences


def calculer_similarite_tableaux(tables1: List[Dict], tables2: List[Dict]) -> Tuple[float, List[Dict]]:
    """
    Compare les tableaux de deux documents.
    """
    if not tables1 and not tables2:
        return 100.0, []

    differences = []
    nb1 = len(tables1)
    nb2 = len(tables2)

    if nb1 != nb2:
        differences.append({
            "type": "nombre_tableaux_different",
            "original": nb1,
            "genere": nb2
        })

    # Compare les tableaux existants
    for i in range(min(nb1, nb2)):
        t1 = tables1[i]
        t2 = tables2[i]

        if t1["colonnes"] != t2["colonnes"]:
            differences.append({
                "type": "colonnes_differentes",
                "tableau": i + 1,
                "original": t1["colonnes"],
                "genere": t2["colonnes"]
            })

        # Compare les headers
        h1_norm = [normaliser_texte(h) for h in t1["headers"]]
        h2_norm = [normaliser_texte(h) for h in t2["headers"]]

        if h1_norm != h2_norm:
            differences.append({
                "type": "headers_differents",
                "tableau": i + 1,
                "original": t1["headers"],
                "genere": t2["headers"]
            })

    # Score basé sur le nombre de tableaux correspondants
    if max(nb1, nb2) == 0:
        score = 100.0
    else:
        score = (min(nb1, nb2) / max(nb1, nb2)) * 100

    # Pénalité pour les différences de structure
    score -= len(differences) * 5
    score = max(0, score)

    return score, differences


def comparer_documents(original: Path, genere: Path) -> ResultatComparaison:
    """
    Compare deux documents DOCX et produit un rapport de conformité.
    """
    console.print(f"[cyan]Extraction de la structure de l'original...[/cyan]")
    struct_original = extraire_structure(original)

    console.print(f"[cyan]Extraction de la structure du document généré...[/cyan]")
    struct_genere = extraire_structure(genere)

    differences = []
    avertissements = []
    recommandations = []

    # Comparaison de la structure globale
    ratio_paragraphes = min(struct_genere.nb_paragraphes, struct_original.nb_paragraphes) / \
                        max(struct_genere.nb_paragraphes, struct_original.nb_paragraphes) \
                        if max(struct_genere.nb_paragraphes, struct_original.nb_paragraphes) > 0 else 1.0

    if abs(struct_genere.nb_paragraphes - struct_original.nb_paragraphes) > 20:
        differences.append({
            "type": "paragraphes",
            "original": struct_original.nb_paragraphes,
            "genere": struct_genere.nb_paragraphes,
            "ecart": abs(struct_genere.nb_paragraphes - struct_original.nb_paragraphes)
        })

    # Comparaison de la longueur
    ratio_longueur = min(struct_genere.longueur_totale, struct_original.longueur_totale) / \
                     max(struct_genere.longueur_totale, struct_original.longueur_totale) \
                     if max(struct_genere.longueur_totale, struct_original.longueur_totale) > 0 else 1.0

    if ratio_longueur < 0.7:
        avertissements.append(
            f"Le document généré est significativement {'plus court' if struct_genere.longueur_totale < struct_original.longueur_totale else 'plus long'} "
            f"({struct_genere.longueur_totale} vs {struct_original.longueur_totale} caractères)"
        )

    # Comparaison des titres
    score_titres, diff_titres = calculer_similarite_titres(
        struct_original.titres,
        struct_genere.titres
    )
    differences.extend(diff_titres)

    # Comparaison des tableaux
    score_tableaux, diff_tableaux = calculer_similarite_tableaux(
        struct_original.tables_info,
        struct_genere.tables_info
    )
    differences.extend(diff_tableaux)

    # Comparaison des styles utilisés
    styles_manquants = set(struct_original.styles_utilises.keys()) - set(struct_genere.styles_utilises.keys())
    if styles_manquants:
        for style in styles_manquants:
            if struct_original.styles_utilises[style] > 2:  # Ignore les styles peu utilisés
                avertissements.append(f"Style '{style}' présent dans l'original mais absent du document généré")

    # Score de structure
    score_structure = (ratio_paragraphes * 0.5 + ratio_longueur * 0.5) * 100

    # Score global
    score_global = (
        score_structure * 0.3 +
        score_titres * 0.4 +
        score_tableaux * 0.3
    )

    # Génération des recommandations
    if score_titres < 80:
        recommandations.append("Vérifier que tous les titres de section sont présents dans le template")

    if score_tableaux < 80:
        recommandations.append("Vérifier la structure des tableaux (colonnes, en-têtes)")

    if ratio_longueur < 0.8:
        recommandations.append("Le document généré semble incomplet. Vérifier que toutes les sections sont remplies.")

    for style in ["Heading 1", "Heading 2", "Titre 1", "Titre 2"]:
        if style in struct_original.styles_utilises and style not in struct_genere.styles_utilises:
            recommandations.append(f"Ajouter le style '{style}' au template pour une meilleure structure")

    return ResultatComparaison(
        conformite_globale=round(score_global, 1),
        conformite_structure=round(score_structure, 1),
        conformite_titres=round(score_titres, 1),
        conformite_tableaux=round(score_tableaux, 1),
        differences=differences,
        avertissements=avertissements,
        recommandations=recommandations
    )


def afficher_rapport(resultat: ResultatComparaison, struct_original: StructureDocument, struct_genere: StructureDocument) -> None:
    """Affiche un rapport de comparaison formaté."""

    # Score global
    couleur_score = "green" if resultat.conformite_globale >= 80 else "yellow" if resultat.conformite_globale >= 60 else "red"

    console.print()
    console.print(Panel(
        f"[bold {couleur_score}]Score de conformité: {resultat.conformite_globale}%[/bold {couleur_score}]",
        title="RÉSULTAT",
        border_style=couleur_score
    ))

    # Détail des scores
    table = Table(title="Détail des scores", show_header=True)
    table.add_column("Critère", style="cyan")
    table.add_column("Score", justify="right")
    table.add_column("Statut", justify="center")

    for critere, score in [
        ("Structure", resultat.conformite_structure),
        ("Titres", resultat.conformite_titres),
        ("Tableaux", resultat.conformite_tableaux)
    ]:
        if score >= 80:
            statut = "[green]✓[/green]"
        elif score >= 60:
            statut = "[yellow]~[/yellow]"
        else:
            statut = "[red]✗[/red]"
        table.add_row(critere, f"{score}%", statut)

    console.print(table)

    # Statistiques comparatives
    console.print()
    console.print("[bold]Statistiques comparatives:[/bold]")

    stats_table = Table(show_header=True)
    stats_table.add_column("Métrique", style="cyan")
    stats_table.add_column("Original", justify="right")
    stats_table.add_column("Généré", justify="right")
    stats_table.add_column("Écart", justify="right")

    stats_table.add_row(
        "Paragraphes",
        str(struct_original.nb_paragraphes),
        str(struct_genere.nb_paragraphes),
        str(struct_genere.nb_paragraphes - struct_original.nb_paragraphes)
    )
    stats_table.add_row(
        "Tableaux",
        str(struct_original.nb_tableaux),
        str(struct_genere.nb_tableaux),
        str(struct_genere.nb_tableaux - struct_original.nb_tableaux)
    )
    stats_table.add_row(
        "Titres détectés",
        str(len(struct_original.titres)),
        str(len(struct_genere.titres)),
        str(len(struct_genere.titres) - len(struct_original.titres))
    )
    stats_table.add_row(
        "Longueur (car.)",
        str(struct_original.longueur_totale),
        str(struct_genere.longueur_totale),
        str(struct_genere.longueur_totale - struct_original.longueur_totale)
    )

    console.print(stats_table)

    # Différences
    if resultat.differences:
        console.print()
        console.print("[bold yellow]Différences détectées:[/bold yellow]")
        for diff in resultat.differences[:10]:  # Limite à 10
            if diff["type"] == "titre_manquant_genere":
                console.print(f"  [red]- Titre manquant:[/red] {diff['titre_original'][:60]}...")
            elif diff["type"] == "titre_supplementaire":
                console.print(f"  [yellow]+ Titre en plus:[/yellow] {diff['titre_genere'][:60]}...")
            elif diff["type"] == "nombre_tableaux_different":
                console.print(f"  [yellow]⚠ Tableaux:[/yellow] {diff['original']} original vs {diff['genere']} généré")
            else:
                console.print(f"  [dim]{diff}[/dim]")

        if len(resultat.differences) > 10:
            console.print(f"  [dim]... et {len(resultat.differences) - 10} autres différences[/dim]")

    # Avertissements
    if resultat.avertissements:
        console.print()
        console.print("[bold yellow]Avertissements:[/bold yellow]")
        for avert in resultat.avertissements:
            console.print(f"  [yellow]⚠[/yellow] {avert}")

    # Recommandations
    if resultat.recommandations:
        console.print()
        console.print("[bold blue]Recommandations:[/bold blue]")
        for reco in resultat.recommandations:
            console.print(f"  [blue]→[/blue] {reco}")


def main():
    """Point d'entrée principal."""
    parser = argparse.ArgumentParser(
        description="Compare la structure d'un document généré avec un original",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples:
  python comparer_documents.py --original docs_originels/trame.docx --genere outputs/acte.docx
  python comparer_documents.py -o docs_originels/trame.docx -g outputs/acte.docx --rapport .tmp/rapport.json
        """
    )

    parser.add_argument(
        "--original", "-o",
        required=True,
        type=str,
        help="Chemin vers le document original (trame)"
    )

    parser.add_argument(
        "--genere", "-g",
        required=True,
        type=str,
        help="Chemin vers le document généré"
    )

    parser.add_argument(
        "--rapport", "-r",
        type=str,
        help="Fichier de sortie pour le rapport JSON (optionnel)"
    )

    parser.add_argument(
        "--seuil",
        type=float,
        default=80.0,
        help="Seuil de conformité minimum (défaut: 80%%)"
    )

    args = parser.parse_args()

    # Vérifie les fichiers
    original_path = Path(args.original)
    genere_path = Path(args.genere)

    if not original_path.exists():
        console.print(f"[red]Fichier original non trouvé: {original_path}[/red]")
        sys.exit(1)

    if not genere_path.exists():
        console.print(f"[red]Fichier généré non trouvé: {genere_path}[/red]")
        sys.exit(1)

    # Affiche l'en-tête
    console.print()
    console.print(Panel(
        f"[bold white]COMPARAISON DE DOCUMENTS[/bold white]\n\n"
        f"[cyan]Original:[/cyan] {original_path.name}\n"
        f"[cyan]Généré:[/cyan] {genere_path.name}",
        border_style="blue"
    ))

    # Extraction des structures
    struct_original = extraire_structure(original_path)
    struct_genere = extraire_structure(genere_path)

    # Comparaison
    resultat = comparer_documents(original_path, genere_path)

    # Affichage du rapport
    afficher_rapport(resultat, struct_original, struct_genere)

    # Export JSON si demandé
    if args.rapport:
        rapport_path = Path(args.rapport)
        rapport_path.parent.mkdir(parents=True, exist_ok=True)

        rapport_data = {
            "original": asdict(struct_original),
            "genere": asdict(struct_genere),
            "resultat": asdict(resultat)
        }

        with open(rapport_path, "w", encoding="utf-8") as f:
            json.dump(rapport_data, f, ensure_ascii=False, indent=2)

        console.print(f"\n[green]Rapport exporté: {rapport_path}[/green]")

    # Code de retour basé sur le seuil
    if resultat.conformite_globale >= args.seuil:
        console.print(f"\n[green]✓ Conformité atteinte ({resultat.conformite_globale}% >= {args.seuil}%)[/green]")
        sys.exit(0)
    else:
        console.print(f"\n[red]✗ Conformité insuffisante ({resultat.conformite_globale}% < {args.seuil}%)[/red]")
        sys.exit(1)


if __name__ == "__main__":
    main()
