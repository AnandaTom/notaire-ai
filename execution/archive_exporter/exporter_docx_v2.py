#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
exporter_docx.py
================

Export d'actes notariaux HTML/Markdown vers DOCX (Word).
Version corrigee avec: tableaux, espacement reduit, pagination simple, hyphenation.

Usage:
    python exporter_docx.py --input <acte.md> --output <acte.docx>
    python exporter_docx.py --input <acte.md> --output <acte.docx> --zones-grisees

Options:
    --zones-grisees : Conserver les zones grisees sur les variables remplies

Base sur l'analyse des originaux RTF:
- Police: Times New Roman 11pt
- Interligne: simple
- Espacement: 3pt (normal), 6pt (sous-titres), 12pt (titres)
- Marges miroir: interieur 60mm, exterieur 15mm
- Marge haut/bas: 25mm
"""

import argparse
import re
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# CONFIGURATION GLOBALE
# =============================================================================

# Option pour conserver les zones grisees sur les variables remplies
ZONES_GRISEES_ACTIVES = False

# Marqueurs pour identifier les variables remplies (inseres par assembler_acte.py)
# Format: <<<VAR_START>>>valeur<<<VAR_END>>>
MARQUEUR_VAR_START = "<<<VAR_START>>>"
MARQUEUR_VAR_END = "<<<VAR_END>>>"


def appliquer_fond_gris(run):
    """
    Applique un fond grise (shading) a un run Word.
    Couleur: gris clair (#D9D9D9) - identique aux zones grisees des trames originales.
    """
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')  # Gris clair
    rPr.append(shd)


# =============================================================================
# NETTOYAGE XML
# =============================================================================

def nettoyer_texte_xml(texte: str) -> str:
    """
    Supprime les caracteres de controle invalides pour XML.
    XML 1.0 accepte: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
    """
    if not texte:
        return texte
    # Supprimer tous les caracteres de controle sauf tab, newline, carriage return
    return ''.join(
        c for c in texte
        if c in '\t\n\r' or (ord(c) >= 0x20 and ord(c) <= 0xD7FF) or
        (ord(c) >= 0xE000 and ord(c) <= 0xFFFD)
    )


# =============================================================================
# TABLEAUX MARKDOWN
# =============================================================================

def detecter_tableau_markdown(lignes: list, index: int) -> tuple:
    """
    Detecte un tableau Markdown et retourne (est_tableau, fin_index, donnees).
    Un tableau commence par | et a une ligne de separation |---|---|.
    """
    if index >= len(lignes):
        return False, index, None

    ligne = lignes[index].strip()
    if not ligne.startswith('|'):
        return False, index, None

    # Verifier la ligne suivante pour le separateur
    if index + 1 >= len(lignes):
        return False, index, None

    sep_ligne = lignes[index + 1].strip()
    if not re.match(r'^\|[\s\-:]+(\|[\s\-:]+)+\|?$', sep_ligne):
        return False, index, None

    # C'est un tableau - extraire toutes les lignes
    tableau = []
    entete = [cell.strip() for cell in ligne.split('|') if cell.strip() or cell == '']
    # Nettoyer les cellules vides aux extremites
    entete = [c for c in ligne.split('|')[1:-1]]
    entete = [c.strip() for c in entete]
    tableau.append(entete)

    # Extraire alignements depuis le separateur
    alignements = []
    for cell in sep_ligne.split('|')[1:-1]:
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            alignements.append('center')
        elif cell.endswith(':'):
            alignements.append('right')
        else:
            alignements.append('left')

    # Lignes de donnees
    fin_index = index + 2
    while fin_index < len(lignes):
        data_ligne = lignes[fin_index].strip()
        if not data_ligne.startswith('|'):
            break
        cells = [c.strip() for c in data_ligne.split('|')[1:-1]]
        tableau.append(cells)
        fin_index += 1

    return True, fin_index, {'lignes': tableau, 'alignements': alignements}


# =============================================================================
# TABLEAUX "APLATIS" (convertis depuis DOC sans delimiteurs pipe)
# =============================================================================

# Patterns d'en-tetes de tableaux notariaux courants (chaque element = 1 colonne)
ENTETES_TABLEAUX_CONNUS = [
    # ==========================================================================
    # 1. DESIGNATION CADASTRALE (tableaux de biens immobiliers)
    # ==========================================================================
    ['Section', 'N°', 'Lieudit', 'Surface'],
    ['Section', 'Numéro', 'Lieudit', 'Surface'],
    ['Section', 'N°', 'Lieu-dit', 'Contenance'],
    ['Section', 'N°', 'Lieu-dit', 'Surface'],
    ['Section', 'Numéro', 'Lieu-dit', 'Contenance'],
    ['Section', 'N°', 'Nature', 'Lieudit', 'Surface'],  # 5 colonnes
    ['Section', 'N°', 'Nature', 'Lieu-dit', 'Contenance'],
    ['Commune', 'Section', 'N°', 'Lieudit', 'Surface'],  # Avec commune
    ['Commune', 'Section', 'Numéro', 'Lieu-dit', 'Contenance'],
    ['Parcelle', 'Section', 'N°', 'Surface'],
    ['Référence cadastrale', 'Surface'],  # 2 colonnes simplifié

    # ==========================================================================
    # 2. REPARTITION DES CHARGES
    # ==========================================================================
    ['Type', 'Montant', 'Quote-part'],
    ['Type de charge', 'Montant', 'Quote-part'],
    ['Nature', 'Montant', 'Quote-part'],
    ['Charge', 'Montant annuel', 'Quote-part'],
    ['Désignation', 'Montant', 'Répartition'],
    ['Nature de la charge', 'Montant (€)', 'Quote-part (%)'],
    ['Charges', 'Montant', 'Part vendeur', 'Part acquéreur'],  # 4 colonnes
    ['Type de charge', 'Montant', 'Part vendeur', 'Part acquéreur'],

    # ==========================================================================
    # 3. ETAT HYPOTHECAIRE / INSCRIPTIONS
    # ==========================================================================
    ['Créancier', 'Montant', 'Date'],
    ['Créancier', 'Capital', 'Date inscription'],
    ['Créancier', 'Montant', 'Date', 'Mainlevée'],  # 4 colonnes
    ['Nature inscription', 'Créancier', 'Montant', 'Date'],
    ['Inscription', 'Créancier', 'Capital', 'Date'],
    ['Type', 'Créancier', 'Montant initial', 'Solde'],
    ['Privilège/Hypothèque', 'Créancier', 'Montant', 'Date'],
    ['Nature', 'Bénéficiaire', 'Montant', 'Date inscription'],

    # ==========================================================================
    # 4. FRAIS D'ACTE / EMOLUMENTS
    # ==========================================================================
    ['Nature', 'Base', 'Taux', 'Montant'],
    ['Désignation', 'Base', 'Taux', 'Montant'],
    ['Nature des frais', 'Assiette', 'Taux', 'Montant (€)'],
    ['Frais', 'Base', 'Taux', 'Total'],
    ['Émoluments', 'Base', 'Taux', 'Montant'],
    ['Nature', 'Assiette (€)', 'Taux (%)', 'Montant (€)'],
    ['Poste', 'Base taxable', 'Taux', 'Montant HT'],
    ['Nature', 'Base', 'Taux', 'Montant HT', 'TVA', 'Montant TTC'],  # 6 colonnes

    # ==========================================================================
    # 5. REPARTITION DU PRIX
    # ==========================================================================
    ['Désignation', 'Prix'],
    ['Désignation', 'Montant'],
    ['Bien', 'Prix'],
    ['Lot', 'Désignation', 'Prix'],  # 3 colonnes
    ['N° lot', 'Désignation', 'Prix'],
    ['Désignation du bien', 'Prix (€)'],
    ['Nature', 'Désignation', 'Prix'],
    ['Élément', 'Montant'],
    ['Répartition', 'Montant (€)'],

    # ==========================================================================
    # 6. IMPOTS FONCIERS / TAXES
    # ==========================================================================
    ['Taxe', 'Base', 'Taux', 'Montant'],
    ['Impôt', 'Base imposable', 'Taux', 'Montant'],
    ['Nature taxe', 'Assiette', 'Taux', 'Montant dû'],
    ['Taxe foncière', 'Base', 'Taux', 'Montant'],
    ['Contribution', 'Base', 'Taux', 'Montant'],
    ['Type de contribution', 'Assiette (€)', 'Taux', 'Montant (€)'],
    ['Type de contribution', 'Assiette', 'Taux', 'Montant'],
    ['Impôt', 'Année', 'Montant'],  # 3 colonnes
    ['Taxe', 'Période', 'Montant'],

    # ==========================================================================
    # 7. DIAGNOSTICS IMMOBILIERS
    # ==========================================================================
    ['Type', 'Date', 'Validité'],
    ['Diagnostic', 'Date', 'Validité'],
    ['Nature diagnostic', 'Date réalisation', 'Date fin validité'],
    ['Diagnostic', 'Date', 'Échéance'],
    ['Type diagnostic', 'Diagnostiqueur', 'Date', 'Validité'],  # 4 colonnes
    ['Diagnostic', 'N° rapport', 'Date', 'Validité'],
    ['Nature', 'Date établissement', 'Durée validité'],
    ['DDT', 'Date', 'Validité'],
    ['Contrôle', 'Date', 'Résultat', 'Validité'],

    # ==========================================================================
    # 8. SERVITUDES
    # ==========================================================================
    ['Nature', 'Bénéficiaire'],
    ['Servitude', 'Bénéficiaire'],
    ['Type servitude', 'Fonds dominant', 'Fonds servant'],  # 3 colonnes
    ['Nature', 'Bénéficiaire', 'Origine'],
    ['Servitude', 'Nature', 'Bénéficiaire'],
    ['Type', 'Description', 'Bénéficiaire'],
    ['Nature servitude', 'Assiette', 'Bénéficiaire'],
    ['Servitude', 'Parcelle grevée', 'Parcelle bénéficiaire'],

    # ==========================================================================
    # 9. ASSURANCES
    # ==========================================================================
    ['Compagnie', 'N° Police', 'Échéance'],
    ['Assureur', 'N° contrat', 'Échéance'],
    ['Compagnie', 'Numéro police', 'Date échéance'],
    ['Assurance', 'N° Police', 'Montant', 'Échéance'],  # 4 colonnes
    ['Type assurance', 'Compagnie', 'N° Police', 'Échéance'],
    ['Risque', 'Assureur', 'Police', 'Prime'],
    ['Nature', 'Compagnie', 'Contrat', 'Validité'],
    ['Assurance', 'Compagnie', 'N°', 'Montant garantie'],

    # ==========================================================================
    # 10. FINANCEMENT / PRETS
    # ==========================================================================
    ['Prêteur', 'Montant', 'Taux', 'Durée'],
    ['Banque', 'Capital', 'Taux', 'Durée'],
    ['Organisme', 'Montant prêt', 'Taux nominal', 'Durée mois'],
    ['Prêteur', 'Capital emprunté', 'Taux', 'Durée', 'Mensualité'],  # 5 colonnes
    ['Type prêt', 'Organisme', 'Montant', 'Taux', 'Durée'],
    ['Financement', 'Montant', 'Taux', 'Durée'],
    ['Source', 'Montant', 'Conditions'],  # 3 colonnes
    ['Nature', 'Prêteur', 'Montant', 'Taux'],
    ['Prêt', 'Banque', 'Capital', 'Intérêts'],

    # ==========================================================================
    # 11. TABLEAUX RECAPITULATIFS / TOTAUX
    # ==========================================================================
    ['Poste', 'Montant'],
    ['Désignation', 'Total'],
    ['Rubrique', 'Montant (€)'],
    ['Élément', 'Sous-total', 'Total'],  # 3 colonnes
    ['Description', 'Montant HT', 'TVA', 'Montant TTC'],  # 4 colonnes

    # ==========================================================================
    # 12. TABLEAUX DE DROITS FISCAUX
    # ==========================================================================
    ['Mt à payer'],
    ['Droit', 'Base', 'Taux', 'Montant'],
    ['Nature droit', 'Assiette', 'Taux applicable', 'Montant dû'],
    ['Droit d\'enregistrement', 'Base', 'Taux', 'Total'],

    # ==========================================================================
    # 13. LOTS DE COPROPRIETE
    # ==========================================================================
    ['N° lot', 'Désignation', 'Quote-part', 'Tantièmes'],
    ['Lot', 'Description', 'Tantièmes'],
    ['N°', 'Nature', 'Étage', 'Tantièmes'],  # 4 colonnes
    ['Lot', 'Consistance', 'Millièmes'],
    ['N° lot', 'Bâtiment', 'Étage', 'Nature', 'Tantièmes'],  # 5 colonnes

    # ==========================================================================
    # 14. URBANISME
    # ==========================================================================
    ['Document', 'Date', 'Zone'],
    ['PLU', 'Zone', 'Règlement'],
    ['Zonage', 'Destination', 'COS'],
    ['Document urbanisme', 'Date approbation', 'Zone'],

    # ==========================================================================
    # 15. PARTIES / INTERVENANTS
    # ==========================================================================
    ['Qualité', 'Nom', 'Adresse'],
    ['Partie', 'Qualité', 'Quote-part'],
    ['Intervenant', 'Qualité', 'Rôle'],
    ['Nom', 'Qualité', 'Domicile'],

    # ==========================================================================
    # 16. ORIGINE DE PROPRIETE
    # ==========================================================================
    ['Nature acte', 'Date', 'Notaire', 'Publication'],  # 4 colonnes
    ['Acte', 'Date', 'Officier public'],
    ['Titre', 'Date', 'Origine'],
    ['Acquisition', 'Date', 'Vendeur', 'Prix'],

    # ==========================================================================
    # 17. MOBILIER
    # ==========================================================================
    ['Désignation', 'Quantité', 'Valeur'],
    ['Article', 'Description', 'Prix'],
    ['Mobilier', 'Estimation'],
    ['Élément mobilier', 'Valeur (€)'],

    # ==========================================================================
    # 18. PROVISIONS / REGULARISATIONS
    # ==========================================================================
    ['Poste', 'Provision', 'Réel', 'Régularisation'],
    ['Charge', 'Appelé', 'Consommé', 'Solde'],
    ['Nature', 'Montant provisionné', 'Montant réel', 'Différence'],

    # ==========================================================================
    # 19. SEQUESTRE / CONSIGNATION
    # ==========================================================================
    ['Motif', 'Montant', 'Bénéficiaire'],
    ['Consignation', 'Montant', 'Durée'],
    ['Séquestre', 'Montant', 'Conditions libération'],

    # ==========================================================================
    # 20. BORNAGE
    # ==========================================================================
    ['Point', 'Coordonnées', 'Nature'],
    ['Borne', 'X', 'Y', 'Z'],  # 4 colonnes
    ['Repère', 'Coordonnées Lambert'],
]

# Patterns speciaux pour tables avec structure asymetrique
# Ces tables ont moins d'en-tetes que de colonnes de donnees
ENTETES_ASYMETRIQUES = {
    # Tableau recapitulatif financement: en-tete 3 cols -> table 4 cols
    # En-tete: M…., Mme…, TOTAL
    # Donnees: Label, val1, val2, val3
    'financement': {
        'pattern': ['M….', 'Mme…', 'TOTAL'],
        'nb_colonnes': 4,  # Ajoute une colonne implicite pour les labels de ligne
    }
}

# Mots-cles qui indiquent potentiellement un debut d'en-tete de tableau
# Ces mots-cles doivent etre des EN-TETES de colonnes, pas des donnees
MOTS_CLES_ENTETE_TABLEAU = [
    # Designation cadastrale
    'Section', 'Commune', 'Parcelle', 'Référence cadastrale',
    # Charges et repartitions
    'Type de charge', 'Nature de la charge', 'Charges',
    # Etat hypothecaire
    'Créancier', 'Nature inscription', 'Inscription', 'Privilège/Hypothèque',
    # Frais d'acte
    'Nature des frais', 'Frais', 'Émoluments',
    # Repartition prix
    'Désignation du bien', 'Répartition',
    # Impots fonciers
    'Type de contribution', 'Taxe foncière', 'Contribution', 'Impôt', 'Taxe',
    # Diagnostics
    'Diagnostic', 'Nature diagnostic', 'Type diagnostic', 'DDT', 'Contrôle',
    # Servitudes
    'Servitude', 'Type servitude', 'Nature servitude',
    # Assurances
    'Compagnie', 'Assureur', 'Assurance', 'Type assurance', 'Risque',
    # Financement
    'Prêteur', 'Banque', 'Organisme', 'Type prêt', 'Financement', 'Prêt',
    # Recapitulatifs
    'Poste', 'Rubrique', 'Élément',
    # Droits fiscaux
    'Mt à payer', 'Droit', 'Nature droit', "Droit d'enregistrement",
    # Lots copropriete
    'N° lot', 'Lot',
    # Urbanisme
    'Document', 'PLU', 'Zonage', 'Document urbanisme',
    # Parties
    'Qualité', 'Partie', 'Intervenant',
    # Origine propriete
    'Nature acte', 'Acte', 'Titre', 'Acquisition',
    # Mobilier
    'Article', 'Mobilier', 'Élément mobilier',
    # Provisions
    'Charge',
    # Sequestre
    'Motif', 'Consignation', 'Séquestre',
    # Bornage
    'Point', 'Borne', 'Repère',
    # Patterns asymetriques (financement recapitulatif)
    'M….', 'M....',
]


def detecter_tableau_aplati(lignes: list, index: int) -> tuple:
    """
    Detecte un tableau "aplati" (chaque cellule sur une ligne separee, separees par lignes vides).

    Pattern typique d'un tableau cadastral aplati:
        Section          <- en-tete col 1
        (ligne vide)
        N°               <- en-tete col 2
        (ligne vide)
        Lieudit          <- en-tete col 3
        (ligne vide)
        Surface          <- en-tete col 4
        (ligne vide)
        AH               <- donnee col 1
        (ligne vide)
        211              <- donnee col 2
        ...

    Retourne: (est_tableau, fin_index, donnees) ou (False, index, None)
    """
    if index >= len(lignes):
        return False, index, None

    ligne = lignes[index].strip()

    # Verifier si c'est un mot-cle d'en-tete connu
    est_entete_potentiel = any(
        ligne.lower() == kw.lower() or ligne.lower().startswith(kw.lower())
        for kw in MOTS_CLES_ENTETE_TABLEAU
    )

    if not est_entete_potentiel:
        return False, index, None

    # Collecter les lignes non-vides consecutives (avec lignes vides entre elles)
    cellules = []
    i = index
    lignes_vides_consecutives = 0
    max_lignes_vides = 2  # Tolerer jusqu'a 2 lignes vides entre cellules

    while i < len(lignes) and lignes_vides_consecutives <= max_lignes_vides:
        ligne_courante = lignes[i].strip()

        if not ligne_courante:
            lignes_vides_consecutives += 1
            i += 1
            continue

        # Arreter si on rencontre un titre markdown
        if ligne_courante.startswith('#'):
            break

        # Arreter si la ligne est trop longue (probablement un paragraphe)
        if len(ligne_courante) > 80:
            break

        # Arreter si on tombe sur une phrase complete (contient un point suivi d'espace ou fin)
        if re.search(r'\.\s+[A-Z]', ligne_courante) or (ligne_courante.endswith('.') and len(ligne_courante) > 50):
            break

        lignes_vides_consecutives = 0
        cellules.append(ligne_courante)
        i += 1

    # Il faut au moins 6 cellules pour un tableau minimal (ex: 4 en-tetes + 4 donnees pour 4 colonnes min 2 lignes)
    if len(cellules) < 6:
        return False, index, None

    # Essayer de determiner le nombre de colonnes avec les en-tetes connus
    nb_colonnes, est_asymetrique = detecter_colonnes_tableau_aplati_v2(cellules)

    if nb_colonnes < 2:
        return False, index, None

    # Pour les tables asymetriques, on doit ajuster les cellules
    # Ex: pour financement, on a 3 en-tetes mais 4 colonnes, donc on insere une cellule vide
    if est_asymetrique:
        # Trouver le nombre d'en-tetes du pattern asymetrique
        for nom, config in ENTETES_ASYMETRIQUES.items():
            pattern = config['pattern']
            if len(cellules) >= len(pattern) and all(
                cellules[i].lower().strip() == pattern[i].lower().strip()
                for i in range(len(pattern))
            ):
                # Inserer une cellule vide au debut de l'en-tete
                cellules.insert(0, '')
                break

    # Verifier que le nombre de cellules est coherent
    if len(cellules) % nb_colonnes != 0:
        # Ajuster en tronquant les cellules excedentaires
        nb_lignes_completes = len(cellules) // nb_colonnes
        if nb_lignes_completes < 2:
            return False, index, None
        cellules = cellules[:nb_lignes_completes * nb_colonnes]

    # Construire le tableau ligne par ligne
    tableau = []
    for row_start in range(0, len(cellules), nb_colonnes):
        row = cellules[row_start:row_start + nb_colonnes]
        if len(row) == nb_colonnes:
            tableau.append(row)

    if len(tableau) < 2:  # Au moins en-tete + 1 ligne de donnees
        return False, index, None

    # Calculer les alignements par defaut
    alignements = ['left'] * nb_colonnes
    # Aligner a droite les colonnes qui semblent numeriques
    for col_idx in range(nb_colonnes):
        valeurs_col = [tableau[row_idx][col_idx] for row_idx in range(1, len(tableau))]
        if all(est_valeur_numerique(v) for v in valeurs_col if v.strip()):
            alignements[col_idx] = 'right'

    return True, i, {'lignes': tableau, 'alignements': alignements}


def detecter_colonnes_tableau_aplati_v2(cellules: list) -> tuple:
    """
    Determine le nombre de colonnes d'un tableau aplati en comparant
    avec les en-tetes connus et en utilisant des heuristiques avancees.

    Retourne: (nb_colonnes, est_asymetrique)
    - nb_colonnes: nombre de colonnes detectees (0 si pas de tableau)
    - est_asymetrique: True si l'en-tete a moins de colonnes que les donnees
    """
    if not cellules:
        return 0, False

    premiere_cellule = cellules[0].lower().strip()

    # ==========================================================================
    # ETAPE 1: Verifier les patterns asymetriques (en-tete != nb colonnes donnees)
    # ==========================================================================
    for nom, config in ENTETES_ASYMETRIQUES.items():
        pattern = config['pattern']
        nb_cols = config['nb_colonnes']
        if len(cellules) >= nb_cols * 2:  # Au moins en-tete + 1 ligne de donnees
            match = all(
                cellules[i].lower().strip() == pattern[i].lower().strip()
                for i in range(len(pattern))
            )
            if match:
                return nb_cols, True  # Asymetrique

    # ==========================================================================
    # ETAPE 2: Verifier correspondance exacte avec patterns d'en-tetes connus
    # ==========================================================================
    for entete_connu in ENTETES_TABLEAUX_CONNUS:
        nb_cols = len(entete_connu)
        if len(cellules) >= nb_cols * 2:  # Au moins en-tete + 1 ligne de donnees
            # Verifier si les premieres cellules correspondent aux en-tetes
            match = all(
                cellules[i].lower().strip() == entete_connu[i].lower().strip()
                for i in range(nb_cols)
            )
            if match:
                return nb_cols, False  # Symetrique

    # ==========================================================================
    # ETAPE 3: Heuristiques specifiques par type de tableau notarial
    # ==========================================================================

    # --- 3.1 Tableaux cadastraux (Section en premiere position) ---
    if premiere_cellule == 'section':
        # Verifier si c'est un tableau 4 ou 5 colonnes selon presence de "Nature"
        if len(cellules) >= 10 and cellules[2].lower().strip() == 'nature':
            return 5, False  # Section, N°, Nature, Lieudit, Surface
        elif len(cellules) >= 8:
            return 4, False  # Section, N°, Lieudit, Surface

    # --- 3.2 Tableaux cadastraux (Commune en premiere position) ---
    if premiere_cellule == 'commune':
        if len(cellules) >= 10:
            return 5, False  # Commune, Section, N°, Lieudit, Surface

    # --- 3.3 Tableaux de droits fiscaux avec formules ---
    if premiere_cellule == 'mt à payer':
        # Verifier si on a des "x" et "=" dans les cellules (indicateurs de formule)
        has_x = any(c.strip().lower() == 'x' for c in cellules[1:])
        has_eq = any(c.strip() == '=' for c in cellules[1:])
        if has_x and has_eq:
            return 6, False  # Desc, Montant, x, %, =, Resultat

    # --- 3.4 Tableaux de diagnostics ---
    if premiere_cellule in ['diagnostic', 'type', 'ddt', 'contrôle']:
        # Chercher "date" et "validité" dans les cellules suivantes
        has_date = any('date' in c.lower() for c in cellules[1:5])
        has_validite = any('validité' in c.lower() or 'échéance' in c.lower() for c in cellules[1:5])
        if has_date and has_validite:
            # Determiner si 3 ou 4 colonnes selon presence diagnostiqueur/resultat
            if any('diagnostiqueur' in c.lower() or 'résultat' in c.lower() for c in cellules[1:5]):
                return 4, False
            return 3, False

    # --- 3.5 Tableaux de financement/prets ---
    if premiere_cellule in ['prêteur', 'banque', 'organisme', 'financement', 'prêt']:
        # Chercher "montant", "taux", "durée"
        has_montant = any('montant' in c.lower() or 'capital' in c.lower() for c in cellules[1:6])
        has_taux = any('taux' in c.lower() for c in cellules[1:6])
        has_duree = any('durée' in c.lower() or 'mois' in c.lower() for c in cellules[1:6])
        if has_montant and has_taux:
            if has_duree:
                # Verifier si mensualité present (5 colonnes)
                if any('mensualité' in c.lower() for c in cellules[1:6]):
                    return 5, False
                return 4, False  # Prêteur, Montant, Taux, Durée
            return 3, False

    # --- 3.6 Tableaux d'assurances ---
    if premiere_cellule in ['compagnie', 'assureur', 'assurance']:
        has_police = any('police' in c.lower() or 'contrat' in c.lower() for c in cellules[1:5])
        has_echeance = any('échéance' in c.lower() or 'validité' in c.lower() for c in cellules[1:5])
        if has_police:
            if any('montant' in c.lower() or 'prime' in c.lower() for c in cellules[1:5]):
                return 4, False  # Compagnie, Police, Montant, Échéance
            return 3, False

    # --- 3.7 Tableaux de servitudes ---
    if premiere_cellule in ['servitude', 'nature servitude', 'type servitude']:
        has_beneficiaire = any('bénéficiaire' in c.lower() for c in cellules[1:4])
        if has_beneficiaire:
            if any('origine' in c.lower() or 'fonds' in c.lower() for c in cellules[1:4]):
                return 3, False
            return 2, False

    # --- 3.8 Tableaux de charges ---
    if premiere_cellule in ['type de charge', 'nature de la charge', 'charge', 'charges']:
        has_montant = any('montant' in c.lower() for c in cellules[1:5])
        has_quotepart = any('quote-part' in c.lower() or 'part' in c.lower() for c in cellules[1:5])
        if has_montant:
            if any('vendeur' in c.lower() and 'acquéreur' in c.lower() for c in cellules[1:5]):
                return 4, False  # Charge, Montant, Part vendeur, Part acquéreur
            if has_quotepart:
                return 3, False
            return 2, False

    # --- 3.9 Tableaux de frais d'acte ---
    if premiere_cellule in ['nature des frais', 'frais', 'émoluments', 'nature']:
        has_base = any('base' in c.lower() or 'assiette' in c.lower() for c in cellules[1:5])
        has_taux = any('taux' in c.lower() for c in cellules[1:5])
        has_montant = any('montant' in c.lower() or 'total' in c.lower() for c in cellules[1:5])
        if has_base and has_taux and has_montant:
            # Verifier si TVA present (6 colonnes)
            if any('tva' in c.lower() or 'ttc' in c.lower() for c in cellules[1:7]):
                return 6, False
            return 4, False

    # --- 3.10 Tableaux etat hypothecaire ---
    if premiere_cellule in ['créancier', 'inscription', 'nature inscription']:
        has_montant = any('montant' in c.lower() or 'capital' in c.lower() for c in cellules[1:5])
        has_date = any('date' in c.lower() for c in cellules[1:5])
        if has_montant and has_date:
            if any('mainlevée' in c.lower() or 'solde' in c.lower() for c in cellules[1:5]):
                return 4, False
            return 3, False

    # --- 3.11 Tableaux de lots de copropriete ---
    if premiere_cellule in ['n° lot', 'lot']:
        has_designation = any('désignation' in c.lower() or 'description' in c.lower() or 'nature' in c.lower() for c in cellules[1:5])
        has_tantiemes = any('tantièmes' in c.lower() or 'millièmes' in c.lower() for c in cellules[1:5])
        if has_tantiemes:
            if any('étage' in c.lower() or 'bâtiment' in c.lower() for c in cellules[1:6]):
                return 5, False  # Lot, Bâtiment, Étage, Nature, Tantièmes
            if has_designation and any('quote-part' in c.lower() for c in cellules[1:5]):
                return 4, False
            return 3, False

    # --- 3.12 Tableaux de contributions/impots ---
    if premiere_cellule in ['type de contribution', 'taxe', 'impôt', 'contribution', 'taxe foncière']:
        has_base = any('base' in c.lower() or 'assiette' in c.lower() for c in cellules[1:5])
        has_taux = any('taux' in c.lower() for c in cellules[1:5])
        has_montant = any('montant' in c.lower() for c in cellules[1:5])
        if has_base and has_taux and has_montant:
            return 4, False
        elif has_montant:
            return 3, False

    # --- 3.13 Tableaux origine de propriete ---
    if premiere_cellule in ['nature acte', 'acte', 'titre', 'acquisition']:
        has_date = any('date' in c.lower() for c in cellules[1:5])
        has_notaire = any('notaire' in c.lower() or 'officier' in c.lower() for c in cellules[1:5])
        if has_date:
            if has_notaire and any('publication' in c.lower() for c in cellules[1:5]):
                return 4, False
            return 3, False

    # ==========================================================================
    # ETAPE 4: Heuristiques generales basees sur le contenu
    # ==========================================================================

    # --- 4.1 Detection par presence de valeurs monetaires et pourcentages ---
    nb_euros = sum(1 for c in cellules if '€' in c or re.match(r'^[\d\s,\.]+\s*euros?$', c.strip(), re.I))
    nb_pourcent = sum(1 for c in cellules if '%' in c)
    nb_numerique = sum(1 for c in cellules if est_valeur_numerique(c))

    if nb_euros > 0 or nb_pourcent > 0:
        # Tableau financier probable - essayer differentes tailles
        for cols in [6, 5, 4, 3, 2]:
            if len(cellules) % cols == 0 and len(cellules) >= cols * 2:
                # Verifier que les colonnes numeriques sont bien reparties
                nb_lignes = len(cellules) // cols
                # Au moins 2 lignes dont 1 d'en-tete
                if nb_lignes >= 2:
                    return cols, False

    # --- 4.2 Detection par presence de dates ---
    nb_dates = sum(1 for c in cellules if re.match(r'^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$', c.strip()))
    if nb_dates > 0:
        # Tableau avec dates - generalement 3 ou 4 colonnes
        for cols in [4, 3]:
            if len(cellules) % cols == 0 and len(cellules) >= cols * 2:
                return cols, False

    # --- 4.3 Detection par surfaces cadastrales ---
    nb_surfaces = sum(1 for c in cellules if re.match(r'^\d+\s*ha\s*\d+\s*a\s*\d+\s*ca$', c.strip(), re.I))
    if nb_surfaces > 0:
        # Tableau cadastral - 4 ou 5 colonnes
        for cols in [5, 4]:
            if len(cellules) % cols == 0 and len(cellules) >= cols * 2:
                return cols, False

    # ==========================================================================
    # ETAPE 5: Fallback - essayer les tailles standard si divisible
    # ==========================================================================
    for cols in [6, 5, 4, 3, 2]:
        if len(cellules) % cols == 0 and len(cellules) >= cols * 2:
            # Verifier coherence minimale: pas trop de cellules vides
            nb_vides = sum(1 for c in cellules if not c.strip())
            if nb_vides < len(cellules) * 0.3:  # Max 30% vides
                return cols, False

    return 0, False  # Pas de tableau detecte


def est_valeur_numerique(texte: str) -> bool:
    """Verifie si un texte est une valeur numerique (montant, pourcentage, surface, etc.)."""
    texte = texte.strip()
    if not texte:
        return False

    # Montants avec devise
    if re.match(r'^[\d\s,\.]+\s*(€|euros?|EUR)?$', texte, re.IGNORECASE):
        return True

    # Pourcentages
    if re.match(r'^[\d\s,\.]+\s*%$', texte):
        return True

    # Surface cadastrale (ex: 00 ha 28 a 21 ca)
    if re.match(r'^\d+\s*ha\s*\d+\s*a\s*\d+\s*ca$', texte, re.IGNORECASE):
        return True

    # Nombres simples
    if re.match(r'^[\d\s,\.]+$', texte):
        return True

    # Operateurs (x, =)
    if texte in ['x', 'X', '=']:
        return True

    return False


def ajouter_tableau_word(doc: Document, donnees: dict):
    """Ajoute un tableau Word depuis les donnees Markdown."""
    lignes = donnees['lignes']
    alignements = donnees['alignements']

    if not lignes or len(lignes) < 1:
        return

    nb_cols = max(len(row) for row in lignes)
    nb_lignes = len(lignes)

    table = doc.add_table(rows=nb_lignes, cols=nb_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Appliquer bordures
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Remplir le tableau
    for i, row_data in enumerate(lignes):
        row = table.rows[i]
        for j in range(nb_cols):
            cell = row.cells[j]
            cell_text = row_data[j] if j < len(row_data) else ''

            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)

            # Alignement
            if j < len(alignements):
                if alignements[j] == 'center':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignements[j] == 'right':
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Texte
            ajouter_texte_formate(para, cell_text)

            # En-tete en gras
            if i == 0:
                for run in para.runs:
                    run.bold = True

    # Espacement apres tableau
    para_after = doc.add_paragraph()
    para_after.paragraph_format.space_after = Pt(18)
    para_after.paragraph_format.space_before = Pt(0)


# =============================================================================
# PARSER HTML
# =============================================================================

class NotarialHTMLParser(HTMLParser):
    """Parser HTML specialise pour les actes notariaux."""

    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_div_class = None
        self.text_buffer = ""
        self.format_stack = []
        self.in_list = False

    def get_current_format(self):
        fmt = {'bold': False, 'italic': False, 'underline': False}
        for f in self.format_stack:
            if f == 'strong' or f == 'b':
                fmt['bold'] = True
            elif f == 'em' or f == 'i':
                fmt['italic'] = True
            # NOTE: On ignore 'u' ici - le soulignement n'est appliqué que via __text__ en Markdown
        return fmt

    def flush_text(self):
        if self.text_buffer and self.current_paragraph is not None:
            fmt = self.get_current_format()
            run = self.current_paragraph.add_run(self.text_buffer)
            run.bold = fmt['bold']
            run.italic = fmt['italic']
            run.underline = fmt['underline']
            appliquer_police(run)
        self.text_buffer = ""

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag == 'div':
            self.flush_text()
            css_class = attrs_dict.get('class', '')
            self.current_div_class = css_class

            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.paragraph_format.first_line_indent = Pt(0)

            if css_class == 'header-ref':
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self.current_paragraph.paragraph_format.space_after = Pt(8)
            elif css_class == 'header-titre':
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.current_paragraph.paragraph_format.space_after = Pt(18)
                self.current_paragraph.paragraph_format.space_before = Pt(10)
            elif css_class == 'header-notaire':
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self.current_paragraph.paragraph_format.space_after = Pt(8)
            elif css_class == 'personne':
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self.current_paragraph.paragraph_format.space_after = Pt(8)
            else:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self.current_paragraph.paragraph_format.space_after = Pt(8)

        elif tag == 'br':
            if self.current_paragraph is not None:
                self.flush_text()
                run = self.current_paragraph.add_run()
                run.add_break()

        elif tag in ['strong', 'b']:
            self.flush_text()
            self.format_stack.append('strong')

        elif tag in ['em', 'i']:
            self.flush_text()
            self.format_stack.append('em')

        # NOTE: On ignore les balises <u> - le soulignement n'est appliqué que via __text__ en Markdown

        elif tag == 'li':
            self.flush_text()
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.current_paragraph.paragraph_format.left_indent = Mm(6)
            self.current_paragraph.paragraph_format.first_line_indent = Mm(-3)
            self.current_paragraph.paragraph_format.space_after = Pt(2)
            self.text_buffer = "- "

    def handle_endtag(self, tag):
        if tag == 'div':
            self.flush_text()
            self.current_div_class = None
            self.current_paragraph = None

        elif tag in ['strong', 'b']:
            self.flush_text()
            if 'strong' in self.format_stack:
                self.format_stack.remove('strong')

        elif tag in ['em', 'i']:
            self.flush_text()
            if 'em' in self.format_stack:
                self.format_stack.remove('em')

        # NOTE: On ignore les balises </u> - le soulignement n'est appliqué que via __text__ en Markdown

        elif tag == 'li':
            self.flush_text()
            self.current_paragraph = None

    def handle_data(self, data):
        if data.strip() or self.text_buffer:
            self.text_buffer += data


# =============================================================================
# CONFIGURATION DOCUMENT
# =============================================================================

def configurer_styles(doc: Document):
    """
    Configure les styles du document EXACTEMENT comme l'original DOCX.
    Basé sur l'analyse de docs_originels/Trame vente lots de copropriété.docx:

    - Normal: 11pt, justified, first line indent 1.251cm
    - Heading 1: bold, ALL CAPS, underline, CENTERED, space after 12pt
    - Heading 2: bold, small caps, underline, CENTERED, space after 12pt
    - Heading 3: bold, underline, CENTERED, space after 12pt
    - Heading 4: bold only, space BEFORE 6pt (not after)
    - Heading 5: bold, underline
    """
    # Style Normal - base pour tous les autres
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style_normal.paragraph_format.space_after = Pt(0)  # Original: null (pas d'espace)
    style_normal.paragraph_format.space_before = Pt(0)
    style_normal.paragraph_format.first_line_indent = Mm(12.51)  # Original: 1.251cm

    # Forcer Times New Roman partout
    rPr = style_normal._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

    # Heading 1: bold, ALL CAPS, underline, CENTERED, espace avant/après
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(11)
    style_h1.font.bold = True
    style_h1.font.all_caps = True
    style_h1.font.underline = True
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h1.paragraph_format.space_before = Pt(24)  # Espace avant pour séparer les sections
    style_h1.paragraph_format.space_after = Pt(18)   # Espace après
    style_h1.paragraph_format.first_line_indent = Pt(0)
    style_h1.paragraph_format.keep_with_next = True
    style_h1.paragraph_format.keep_together = True

    # Heading 2: bold, small caps, underline, CENTERED, espace avant/après
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(11)
    style_h2.font.bold = True
    style_h2.font.small_caps = True
    style_h2.font.underline = True
    style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # CENTERED!
    style_h2.paragraph_format.space_before = Pt(24)  # Espace avant pour séparer
    style_h2.paragraph_format.space_after = Pt(18)   # Espace après
    style_h2.paragraph_format.first_line_indent = Pt(0)
    style_h2.paragraph_format.keep_with_next = True
    style_h2.paragraph_format.keep_together = True

    # Heading 3: bold, underline, CENTERED, espace avant/après
    style_h3 = doc.styles['Heading 3']
    style_h3.font.name = 'Times New Roman'
    style_h3.font.size = Pt(11)
    style_h3.font.bold = True
    style_h3.font.underline = True
    style_h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # CENTERED!
    style_h3.paragraph_format.space_before = Pt(10)   # Petit espace avant
    style_h3.paragraph_format.space_after = Pt(18)   # Espace après
    style_h3.paragraph_format.first_line_indent = Pt(0)
    style_h3.paragraph_format.keep_with_next = True
    style_h3.paragraph_format.keep_together = True

    # Heading 4: bold only, space BEFORE 6pt (NOT after)
    style_h4 = doc.styles['Heading 4']
    style_h4.font.name = 'Times New Roman'
    style_h4.font.size = Pt(11)
    style_h4.font.bold = True
    style_h4.font.italic = False
    style_h4.font.underline = False
    style_h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Hérite de Normal
    style_h4.paragraph_format.space_before = Pt(10)  # Original: 6pt AVANT
    style_h4.paragraph_format.space_after = Pt(0)   # Pas d'espace après
    style_h4.paragraph_format.first_line_indent = Pt(0)
    style_h4.paragraph_format.keep_with_next = True
    style_h4.paragraph_format.keep_together = True

    # Heading 5: bold, underline (utilisé rarement)
    style_h5 = doc.styles['Heading 5']
    style_h5.font.name = 'Times New Roman'
    style_h5.font.size = Pt(11)
    style_h5.font.bold = True
    style_h5.font.underline = True
    style_h5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_h5.paragraph_format.first_line_indent = Pt(0)

    # ===== STYLES PERSONNALISES =====
    # Ces styles sont définis dans l'original et utilisés fréquemment

    # Quote (Citation): italique - 40 occurrences dans l'original
    try:
        style_quote = doc.styles['Quote']
        style_quote.font.name = 'Times New Roman'
        style_quote.font.size = Pt(11)
        style_quote.font.italic = True
        style_quote.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_quote.paragraph_format.first_line_indent = Mm(12.51)
    except KeyError:
        pass  # Style n'existe pas dans ce template

    # Title: 12pt, bold, underline
    style_title = doc.styles['Title']
    style_title.font.name = 'Times New Roman'
    style_title.font.size = Pt(12)
    style_title.font.bold = True
    style_title.font.underline = True


def configurer_marges(doc: Document):
    """
    Configure les marges du document EXACTEMENT comme l'original DOCX.
    Basé sur l'analyse: G=6.0cm D=1.5cm H=2.5cm B=2.5cm
    AVEC marges miroir pour documents notariaux (pages paires/impaires inversées)
    """
    for section in doc.sections:
        section.left_margin = Mm(60)   # 6.0 cm (page impaire/droite)
        section.right_margin = Mm(15)  # 1.5 cm (page impaire/droite)
        section.top_margin = Mm(25)    # 2.5 cm
        section.bottom_margin = Mm(25) # 2.5 cm
        section.page_width = Mm(210)   # A4
        section.page_height = Mm(297)  # A4
        section.gutter = Mm(0)         # Pas de reliure supplémentaire

        # Distance en-tête/pied de page (original: 1.27cm)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        # En-tête différent sur première page (original: true)
        section.different_first_page_header_footer = True

        # ACTIVER les marges miroir via XML (pages paires inversent G/D)
        sectPr = section._sectPr
        mirrorMargins = sectPr.find(qn('w:mirrorMargins'))
        if mirrorMargins is None:
            mirrorMargins = OxmlElement('w:mirrorMargins')
            sectPr.append(mirrorMargins)


def configurer_compatibilite(doc: Document):
    """Configure les options de compatibilite Word pour meilleure justification."""
    settings = doc.settings._element

    compat = settings.find(qn('w:compat'))
    if compat is None:
        compat = OxmlElement('w:compat')
        settings.append(compat)

    # Ne pas etendre les espaces sur lignes courtes
    do_not_expand = OxmlElement('w:doNotExpandShiftReturn')
    compat.append(do_not_expand)


def ajouter_pagination(doc: Document):
    """Ajoute la pagination en haut a droite (numero seul)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_after = Pt(0)

        # Champ PAGE uniquement
        run = para.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)

        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char_begin)

        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'PAGE'
        run._r.append(instr_text)

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_end)


# =============================================================================
# FORMATAGE MARKDOWN
# =============================================================================

def traiter_formatage_markdown(texte: str):
    """
    Parse le formatage Markdown et retourne une liste de tuples (texte, format).
    Gere aussi les marqueurs de variables pour les zones grisees.
    Format: {'bold': bool, 'italic': bool, 'underline': bool, 'zone_grisee': bool}
    """
    segments = []
    i = 0
    current_text = ""
    bold = False
    italic = False
    underline = False
    zone_grisee = False

    while i < len(texte):
        # Detecter debut de zone variable (marqueur)
        if texte[i:i+len(MARQUEUR_VAR_START)] == MARQUEUR_VAR_START:
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            zone_grisee = True
            i += len(MARQUEUR_VAR_START)
            continue

        # Detecter fin de zone variable (marqueur)
        if texte[i:i+len(MARQUEUR_VAR_END)] == MARQUEUR_VAR_END:
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            zone_grisee = False
            i += len(MARQUEUR_VAR_END)
            continue

        if texte[i:i+3] == '***':
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            bold = not bold
            italic = not italic
            i += 3
            continue

        if texte[i:i+2] == '**':
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            bold = not bold
            i += 2
            continue

        if texte[i] == '*' and (i+1 >= len(texte) or texte[i+1] != '*'):
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            italic = not italic
            i += 1
            continue

        if texte[i:i+2] == '__':
            if current_text:
                segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))
                current_text = ""
            underline = not underline
            i += 2
            continue

        current_text += texte[i]
        i += 1

    if current_text:
        segments.append((current_text, {'bold': bold, 'italic': italic, 'underline': underline, 'zone_grisee': zone_grisee}))

    return segments


def ajouter_texte_formate(paragraph, texte: str):
    """
    Ajoute du texte avec formatage Markdown a un paragraphe.
    Applique un fond gris aux variables si ZONES_GRISEES_ACTIVES est True.
    """
    segments = traiter_formatage_markdown(texte)
    for text, fmt in segments:
        if text:
            # Nettoyer les caracteres invalides pour XML
            text = nettoyer_texte_xml(text)
            if text:  # Verifier apres nettoyage
                run = paragraph.add_run(text)
                run.bold = fmt['bold']
                run.italic = fmt['italic']
                run.underline = fmt['underline']
                appliquer_police(run)
                # Appliquer fond gris si c'est une zone variable et l'option est activee
                if ZONES_GRISEES_ACTIVES and fmt.get('zone_grisee', False):
                    appliquer_fond_gris(run)


def appliquer_police(run):
    """Applique la police Times New Roman a un run."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)


# =============================================================================
# DETECTION TITRES
# =============================================================================

def detecter_titre_markdown(ligne: str):
    """Detecte si une ligne est un titre markdown et retourne (niveau, texte)."""
    ligne_strip = ligne.strip()
    if ligne_strip.startswith('### '):
        return (3, ligne_strip[4:])
    elif ligne_strip.startswith('## '):
        return (2, ligne_strip[3:])
    elif ligne_strip.startswith('# '):
        return (1, ligne_strip[2:])
    return (0, ligne_strip)


def est_titre_notarial(texte: str) -> bool:
    """Verifie si le texte est un titre notarial principal."""
    texte_upper = texte.upper().strip()

    # Titres exacts ou au debut de ligne (pour les titres avec ponctuation comme "CECI EXPOSE,")
    titres_debut = [
        'PARTIE NORMALISEE', 'PARTIE NORMALISÉE',
        'PARTIE DEVELOPPEE', 'PARTIE DÉVELOPPÉE',
        'IDENTIFICATION DES PARTIES',
        'DESIGNATION', 'DÉSIGNATION',
        'ORIGINE DE PROPRIETE', 'ORIGINE DE PROPRIÉTÉ',
        'CHARGES ET CONDITIONS',
        'PRIX ET PAIEMENT',
        'GARANTIES',
        'TERMINOLOGIE',
        'CECI EXPOSE',
        'NATURE ET QUOTITÉ', 'NATURE ET QUOTITE',
        'FIN DE PARTIE',
        'FIXATION DE LA PROPORTION',
    ]
    for t in titres_debut:
        if texte_upper.startswith(t):
            return True
    return False


def est_sous_titre_notarial(texte: str) -> bool:
    """Verifie si le texte est un sous-titre notarial."""
    texte_upper = texte.upper().strip()
    sous_titres = [
        'VENDEUR', 'ACQUEREUR', 'ACQUÉREUR',
        'QUOTITÉS', 'QUOTITES',
        'PRESENCE', 'PRÉSENCE',
        'REPRESENTATION', 'REPRÉSENTATION',
        'DECLARATIONS', 'DÉCLARATIONS',
        'DOCUMENTS RELATIFS',
        'DONT QUITTANCE',
        'FINANCEMENT',
        'TOTAL',
        'CONCERNANT',
    ]
    for s in sous_titres:
        if texte_upper.startswith(s):
            return True
    return False


# =============================================================================
# CONVERSION PRINCIPALE
# =============================================================================

def appliquer_bordures_tableau(table):
    """
    Applique des bordures noires simples à un tableau Word.
    Utilisé pour créer des encadrés (boxes).
    """
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def convertir_contenu_vers_docx(contenu: str, doc: Document):
    """Convertit le contenu HTML/Markdown vers Word."""

    # Nettoyer les caracteres de controle invalides pour XML
    contenu = nettoyer_texte_xml(contenu)

    # CORRECTION 3: Supprimer tous les commentaires HTML
    contenu = re.sub(r'<!--.*?-->', '', contenu, flags=re.DOTALL)
    contenu = re.sub(r'\n\s*\n\s*\n', '\n\n', contenu)

    # Supprimer les balises <u> et </u> du HTML source
    # Le soulignement n'est appliqué que via __text__ en Markdown
    contenu = re.sub(r'</?u>', '', contenu)

    has_html = '<div' in contenu or '<strong>' in contenu or '<br' in contenu

    lignes = contenu.split('\n')
    i = 0
    parser = NotarialHTMLParser(doc) if has_html else None
    html_buffer = ""
    in_html_block = False
    in_first_page_header = False
    first_page_header_started = False
    in_box = False
    box_table = None

    while i < len(lignes):
        ligne = lignes[i]
        ligne_strip = ligne.strip()

        # Detecter debut/fin de box (encadré)
        if 'BOX_START}' in ligne_strip:
            in_box = True
            # Créer un tableau avec une seule cellule pour faire l'encadré
            box_table = doc.add_table(rows=1, cols=1)
            box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # Appliquer les bordures
            appliquer_bordures_tableau(box_table)
            i += 1
            continue
        if 'BOX_END}' in ligne_strip:
            in_box = False
            box_table = None
            i += 1
            continue

        # Detecter debut/fin de header de premiere page
        if '{FIRST_PAGE_HEADER_START}' in ligne_strip:
            in_first_page_header = True
            first_page_header_started = True
            # Ajouter un grand espace en haut de la premiere page (environ 10cm)
            para_espace = doc.add_paragraph()
            para_espace.paragraph_format.space_before = Mm(100)  # 10 cm d'espace avant
            para_espace.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        if '{FIRST_PAGE_HEADER_END}' in ligne_strip:
            in_first_page_header = False
            # Ajouter un espace après le header (environ 1cm)
            para_espace_apres = doc.add_paragraph()
            para_espace_apres.paragraph_format.space_after = Mm(10)
            i += 1
            continue

        # Traiter ligne header de premiere page (reference, initiales, date)
        if in_first_page_header and ligne_strip:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centré comme dans l'original
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.first_line_indent = Mm(0)  # Pas de retrait
            para.paragraph_format.left_indent = Mm(0)  # Pas d'indentation
            run = para.add_run(ligne_strip)
            appliquer_police(run)
            i += 1
            continue

        # Ignorer lignes vides
        if not ligne_strip:
            i += 1
            continue

        # CORRECTION 2: Detecter tableaux Markdown (format standard avec |)
        est_tableau, fin_idx, donnees = detecter_tableau_markdown(lignes, i)
        if est_tableau:
            ajouter_tableau_word(doc, donnees)
            i = fin_idx
            continue

        # CORRECTION 6: Detecter tableaux "aplatis" (convertis depuis DOC sans delimiteurs)
        est_tableau_aplati, fin_idx_aplati, donnees_aplati = detecter_tableau_aplati(lignes, i)
        if est_tableau_aplati:
            ajouter_tableau_word(doc, donnees_aplati)
            i = fin_idx_aplati
            continue

        # Traiter blocs HTML
        if has_html and '<div' in ligne_strip:
            in_html_block = True
            html_buffer = ligne + "\n"
            i += 1
            continue

        if in_html_block:
            html_buffer += ligne + "\n"
            if '</div>' in ligne_strip:
                parser.feed(html_buffer)
                html_buffer = ""
                in_html_block = False
            i += 1
            continue

        # Ligne markdown simple
        # Si on est dans une box, ajouter le contenu dans la cellule du tableau
        if in_box and box_table:
            cell = box_table.rows[0].cells[0]
            traiter_ligne_markdown_dans_conteneur(ligne_strip, cell)
        else:
            traiter_ligne_markdown(ligne_strip, doc)
        i += 1


def traiter_ligne_markdown_dans_conteneur(ligne: str, cell):
    """
    Traite une ligne de Markdown et l'ajoute dans une cellule de tableau (pour les encadrés).
    Utilise les mêmes styles que traiter_ligne_markdown mais adapté pour un conteneur.
    """
    # Ignorer separateurs
    if ligne in ['---', '***', '___']:
        return

    # Nettoyer echappements
    ligne = ligne.replace('\\-', '-')
    ligne = ligne.replace('\\*', '*')

    # Titres markdown avec # → ajouter avec formatage approprié
    niveau, texte = detecter_titre_markdown(ligne)
    if niveau > 0:
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(texte)
        run.bold = True
        run.underline = True
        if niveau == 1:
            run.font.all_caps = True
        elif niveau == 2:
            run.font.small_caps = True
        appliquer_police(run)
        return

    # Paragraphe normal dans la cellule
    para = cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.first_line_indent = Mm(12.51)
    ajouter_texte_formate(para, ligne)


def traiter_ligne_markdown(ligne: str, doc: Document):
    """
    Traite une ligne de Markdown et l'ajoute au document.
    Applique les styles selon l'analyse du RTF original:
    - Heading 1 (titres majeurs): bold, ALL CAPS, underline, centered
    - Heading 2 (sous-sections): bold, small caps, underline
    - Heading 3 (sous-sous-sections): bold, underline
    - Heading 4 (petits titres): bold only
    """

    # Ignorer separateurs
    if ligne in ['---', '***', '___']:
        return

    # Nettoyer echappements
    ligne = ligne.replace('\\-', '-')
    ligne = ligne.replace('\\*', '*')

    # Titres markdown avec # → utiliser styles Heading
    niveau, texte = detecter_titre_markdown(ligne)
    if niveau > 0:
        if niveau == 1:
            para = doc.add_paragraph(style='Heading 1')
        elif niveau == 2:
            para = doc.add_paragraph(style='Heading 2')
        else:
            para = doc.add_paragraph(style='Heading 3')
        run = para.add_run(texte)
        appliquer_police(run)
        return

    # Listes a puces
    if ligne.startswith('* ') or ligne.startswith('- '):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Mm(6)
        para.paragraph_format.first_line_indent = Mm(-3)
        para.paragraph_format.space_after = Pt(2)
        texte_liste = ligne[2:]
        run = para.add_run('- ')
        appliquer_police(run)
        ajouter_texte_formate(para, texte_liste)
        return

    # Titres notariaux (texte en gras)
    # Detecter soit avec **texte** soit le texte seul sur une ligne
    texte_clean = ligne.replace('**', '').strip()
    est_entoure_bold = ligne.startswith('**') and ligne.endswith('**')

    # Les titres doivent etre relativement courts (moins de 80 caracteres)
    # sauf s'ils sont explicitement marques en bold avec **
    if est_titre_notarial(texte_clean) and (est_entoure_bold or len(texte_clean) < 80):
        # Titre principal → Heading 1 (bold, ALL CAPS, underline, centered)
        para = doc.add_paragraph(style='Heading 1')
        run = para.add_run(texte_clean)
        appliquer_police(run)
        return

    # Sous-titres notariaux - detecter avec ou sans ** marqueurs
    # Verifier que c'est une ligne courte et autonome (pas une phrase complete)
    if est_sous_titre_notarial(texte_clean) and (est_entoure_bold or len(texte_clean) < 50):
        # Sous-titre → Heading 2 (bold, small caps, underline)
        para = doc.add_paragraph(style='Heading 2')
        run = para.add_run(texte_clean)
        appliquer_police(run)
        return

    # Paragraphe normal - EXACTEMENT comme l'original
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)  # Original: pas d'espace après
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.first_line_indent = Mm(12.51)  # Original: 1.251cm
    ajouter_texte_formate(para, ligne)


# =============================================================================
# EXPORT PRINCIPAL
# =============================================================================

def exporter_docx(chemin_entree: Path, chemin_sortie: Path, zones_grisees: bool = False) -> bool:
    """
    Exporte un fichier HTML/Markdown vers DOCX.

    Args:
        chemin_entree: Fichier source (Markdown/HTML)
        chemin_sortie: Fichier DOCX de sortie
        zones_grisees: Si True, conserve les zones grisees sur les variables remplies
    """
    global ZONES_GRISEES_ACTIVES
    ZONES_GRISEES_ACTIVES = zones_grisees

    with open(chemin_entree, 'r', encoding='utf-8') as f:
        contenu = f.read()

    doc = Document()
    configurer_styles(doc)
    configurer_marges(doc)
    configurer_compatibilite(doc)  # CORRECTION 5
    ajouter_pagination(doc)
    convertir_contenu_vers_docx(contenu, doc)

    chemin_sortie.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(chemin_sortie))

    return True


def main():
    parser = argparse.ArgumentParser(description='Exporter un acte HTML/Markdown vers DOCX')
    parser.add_argument('--input', '-i', type=Path, required=True, help='Fichier source')
    parser.add_argument('--output', '-o', type=Path, required=True, help='Fichier DOCX de sortie')
    parser.add_argument('--zones-grisees', '-z', action='store_true',
                        help='Conserver les zones grisees sur les variables remplies')

    args = parser.parse_args()

    if not args.input.exists():
        print(f'[ERREUR] Fichier non trouve: {args.input}')
        return 1

    try:
        if exporter_docx(args.input, args.output, zones_grisees=args.zones_grisees):
            taille = args.output.stat().st_size / 1024
            option_gris = " (avec zones grisees)" if args.zones_grisees else ""
            print(f'[OK] DOCX genere: {args.output} ({taille:.1f} Ko){option_gris}')
            return 0
    except Exception as e:
        print(f'[ERREUR] {e}')
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    exit(main())
