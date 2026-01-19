#!/usr/bin/env python3
"""
Analyse le formatage détaillé d'un document Word (.docx ou .doc converti).
Extrait toutes les informations de mise en forme pour reproduction fidèle.
"""

import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def twips_to_cm(twips):
    """Convertit des twips en centimètres."""
    if twips is None:
        return None
    return round(twips / 567, 2)


def pt_to_float(pt_value):
    """Convertit une valeur Pt en float."""
    if pt_value is None:
        return None
    return float(pt_value.pt) if hasattr(pt_value, 'pt') else float(pt_value)


def analyser_sections(doc):
    """Analyse les sections et leurs marges."""
    sections = []
    for i, section in enumerate(doc.sections):
        section_info = {
            "numero": i + 1,
            "largeur_page_cm": round(section.page_width.cm, 2) if section.page_width else None,
            "hauteur_page_cm": round(section.page_height.cm, 2) if section.page_height else None,
            "marge_gauche_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "marge_droite_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
            "marge_haut_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "marge_bas_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "marge_reliure_cm": round(section.gutter.cm, 2) if section.gutter else None,
            "en_tete_distance_cm": round(section.header_distance.cm, 2) if section.header_distance else None,
            "pied_page_distance_cm": round(section.footer_distance.cm, 2) if section.footer_distance else None,
        }

        # Vérifier marges miroir via XML
        sectPr = section._sectPr
        if sectPr is not None:
            mirrorMargins = sectPr.find(qn('w:mirrorMargins'))
            section_info["marges_miroir"] = mirrorMargins is not None

        sections.append(section_info)
    return sections


def analyser_styles(doc):
    """Analyse tous les styles définis."""
    styles = []
    for style in doc.styles:
        if style.type in [WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER]:
            style_info = {
                "nom": style.name,
                "type": "paragraphe" if style.type == WD_STYLE_TYPE.PARAGRAPH else "caractere",
                "base_sur": style.base_style.name if style.base_style else None,
            }

            # Formatage de police
            if style.font:
                font = style.font
                style_info["police"] = {
                    "nom": font.name,
                    "taille_pt": pt_to_float(font.size),
                    "gras": font.bold,
                    "italique": font.italic,
                    "souligne": font.underline,
                    "majuscules": font.all_caps,
                    "petites_majuscules": font.small_caps,
                }

            # Formatage de paragraphe (seulement pour styles paragraphe)
            if style.type == WD_STYLE_TYPE.PARAGRAPH and style.paragraph_format:
                pf = style.paragraph_format
                alignment_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "gauche",
                    WD_ALIGN_PARAGRAPH.CENTER: "centre",
                    WD_ALIGN_PARAGRAPH.RIGHT: "droite",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justifie",
                }
                style_info["paragraphe"] = {
                    "alignement": alignment_map.get(pf.alignment, str(pf.alignment)),
                    "retrait_premiere_ligne_cm": round(pf.first_line_indent.cm, 2) if pf.first_line_indent else None,
                    "retrait_gauche_cm": round(pf.left_indent.cm, 2) if pf.left_indent else None,
                    "retrait_droite_cm": round(pf.right_indent.cm, 2) if pf.right_indent else None,
                    "espace_avant_pt": pt_to_float(pf.space_before),
                    "espace_apres_pt": pt_to_float(pf.space_after),
                    "interligne": pf.line_spacing,
                    "conserver_avec_suivant": pf.keep_with_next,
                    "conserver_ensemble": pf.keep_together,
                }

            styles.append(style_info)

    return styles


def analyser_paragraphes(doc, limite=50):
    """Analyse les premiers paragraphes pour comprendre le formatage appliqué."""
    paragraphes = []
    for i, para in enumerate(doc.paragraphs[:limite]):
        if not para.text.strip():
            continue

        para_info = {
            "numero": i + 1,
            "texte_apercu": para.text[:100] + "..." if len(para.text) > 100 else para.text,
            "style": para.style.name if para.style else None,
        }

        # Formatage direct du paragraphe
        pf = para.paragraph_format
        if pf:
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "gauche",
                WD_ALIGN_PARAGRAPH.CENTER: "centre",
                WD_ALIGN_PARAGRAPH.RIGHT: "droite",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justifie",
            }
            para_info["format_direct"] = {
                "alignement": alignment_map.get(pf.alignment, str(pf.alignment)) if pf.alignment else None,
                "retrait_premiere_ligne_cm": round(pf.first_line_indent.cm, 2) if pf.first_line_indent else None,
                "espace_avant_pt": pt_to_float(pf.space_before) if pf.space_before else None,
                "espace_apres_pt": pt_to_float(pf.space_after) if pf.space_after else None,
            }

        # Analyser les runs pour voir le formatage de caractères
        runs_info = []
        for run in para.runs[:5]:  # Limiter à 5 runs par paragraphe
            if run.text.strip():
                run_info = {
                    "texte": run.text[:50] + "..." if len(run.text) > 50 else run.text,
                    "gras": run.bold,
                    "italique": run.italic,
                    "souligne": run.underline,
                }
                if run.font:
                    run_info["police"] = run.font.name
                    run_info["taille_pt"] = pt_to_float(run.font.size)
                runs_info.append(run_info)

        if runs_info:
            para_info["runs"] = runs_info

        paragraphes.append(para_info)

    return paragraphes


def analyser_tableaux(doc):
    """Analyse les tableaux du document."""
    tableaux = []
    for i, table in enumerate(doc.tables):
        table_info = {
            "numero": i + 1,
            "nb_lignes": len(table.rows),
            "nb_colonnes": len(table.columns),
            "largeurs_colonnes_cm": [],
        }

        # Largeurs des colonnes
        for col in table.columns:
            if col.width:
                table_info["largeurs_colonnes_cm"].append(round(col.width.cm, 2))

        # Aperçu du contenu
        apercu = []
        for row in table.rows[:3]:  # Limiter à 3 lignes
            row_text = [cell.text[:30] for cell in row.cells[:4]]
            apercu.append(row_text)
        table_info["apercu"] = apercu

        tableaux.append(table_info)

    return tableaux


def analyser_document(chemin_docx):
    """Analyse complète d'un document Word."""
    doc = Document(chemin_docx)

    analyse = {
        "fichier": str(chemin_docx),
        "sections": analyser_sections(doc),
        "styles_principaux": [s for s in analyser_styles(doc)
                             if s["nom"] in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4",
                                            "Title", "Titre", "Titre 1", "Titre 2", "Titre 3", "Titre 4"]],
        "tous_styles": analyser_styles(doc),
        "paragraphes_echantillon": analyser_paragraphes(doc, limite=30),
        "tableaux": analyser_tableaux(doc),
        "statistiques": {
            "nb_paragraphes": len(doc.paragraphs),
            "nb_tableaux": len(doc.tables),
            "nb_sections": len(doc.sections),
        }
    }

    return analyse


def main():
    parser = argparse.ArgumentParser(description="Analyse le formatage d'un document Word")
    parser.add_argument("fichier", help="Chemin vers le fichier .docx à analyser")
    parser.add_argument("-o", "--output", help="Fichier JSON de sortie (optionnel)")
    parser.add_argument("--compact", action="store_true", help="Affichage compact (styles principaux uniquement)")

    args = parser.parse_args()

    chemin = Path(args.fichier)
    if not chemin.exists():
        print(f"Erreur: Le fichier '{chemin}' n'existe pas.")
        return 1

    print(f"Analyse de: {chemin}")
    print("-" * 60)

    analyse = analyser_document(chemin)

    if args.compact:
        # Affichage compact
        print("\n=== SECTIONS ET MARGES ===")
        for section in analyse["sections"]:
            print(f"\nSection {section['numero']}:")
            print(f"  Page: {section['largeur_page_cm']} x {section['hauteur_page_cm']} cm")
            print(f"  Marges: G={section['marge_gauche_cm']}cm, D={section['marge_droite_cm']}cm, H={section['marge_haut_cm']}cm, B={section['marge_bas_cm']}cm")
            if section.get('marge_reliure_cm'):
                print(f"  Reliure: {section['marge_reliure_cm']}cm")
            print(f"  Marges miroir: {section.get('marges_miroir', False)}")

        print("\n=== STYLES PRINCIPAUX ===")
        for style in analyse["styles_principaux"]:
            print(f"\n{style['nom']}:")
            if style.get("police"):
                p = style["police"]
                attrs = []
                if p.get("gras"): attrs.append("gras")
                if p.get("italique"): attrs.append("italique")
                if p.get("souligne"): attrs.append("souligné")
                if p.get("majuscules"): attrs.append("MAJUSCULES")
                if p.get("petites_majuscules"): attrs.append("PETITES CAPS")
                print(f"  Police: {p.get('nom')} {p.get('taille_pt')}pt {' '.join(attrs)}")
            if style.get("paragraphe"):
                pg = style["paragraphe"]
                print(f"  Alignement: {pg.get('alignement')}")
                if pg.get('retrait_premiere_ligne_cm'):
                    print(f"  Retrait 1ère ligne: {pg['retrait_premiere_ligne_cm']}cm")
                if pg.get('espace_apres_pt'):
                    print(f"  Espace après: {pg['espace_apres_pt']}pt")

        print(f"\n=== STATISTIQUES ===")
        stats = analyse["statistiques"]
        print(f"Paragraphes: {stats['nb_paragraphes']}")
        print(f"Tableaux: {stats['nb_tableaux']}")
        print(f"Sections: {stats['nb_sections']}")

    else:
        # Affichage complet JSON
        print(json.dumps(analyse, indent=2, ensure_ascii=False))

    # Sauvegarder si demandé
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(analyse, f, indent=2, ensure_ascii=False)
        print(f"\nAnalyse sauvegardée dans: {args.output}")

    return 0


if __name__ == "__main__":
    exit(main())
