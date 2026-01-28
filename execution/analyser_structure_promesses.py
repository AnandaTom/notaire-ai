#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analyser_structure_promesses.py
===============================

Analyse la structure des trames de promesse de vente pour identifier:
- Les TITRES/HEADINGS (texte en gras, souligné, majuscules)
- Les grandes SECTIONS principales
- Les DIFFÉRENCES de structure entre les trames

Usage:
    python analyser_structure_promesses.py
"""

import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt
from collections import OrderedDict
import json
import re

# Forcer UTF-8 pour Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')


def analyser_style_paragraphe(para) -> dict:
    """Analyse le style d'un paragraphe pour identifier s'il s'agit d'un titre."""
    info = {
        'texte': para.text.strip(),
        'style_name': para.style.name if para.style else None,
        'is_bold': False,
        'is_underline': False,
        'is_caps': False,
        'is_small_caps': False,
        'is_centered': False,
        'niveau_heading': None,
        'font_size': None
    }

    if not para.text.strip():
        return info

    # Vérifier l'alignement
    if para.alignment:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        info['is_centered'] = para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    # Analyser les runs du paragraphe
    for run in para.runs:
        if run.bold:
            info['is_bold'] = True
        if run.underline:
            info['is_underline'] = True
        if run.font.all_caps:
            info['is_caps'] = True
        if run.font.small_caps:
            info['is_small_caps'] = True
        if run.font.size:
            info['font_size'] = run.font.size.pt

    # Vérifier si le texte est en majuscules (même si pas marqué comme tel)
    if para.text.strip() and para.text.strip().upper() == para.text.strip():
        # Vérifier que ce n'est pas juste des chiffres ou symboles
        if any(c.isalpha() for c in para.text):
            info['is_caps'] = True

    # Déterminer le niveau de heading
    style_name = (para.style.name or '').lower() if para.style else ''
    if 'heading 1' in style_name or 'titre 1' in style_name:
        info['niveau_heading'] = 1
    elif 'heading 2' in style_name or 'titre 2' in style_name:
        info['niveau_heading'] = 2
    elif 'heading 3' in style_name or 'titre 3' in style_name:
        info['niveau_heading'] = 3
    elif 'heading 4' in style_name or 'titre 4' in style_name:
        info['niveau_heading'] = 4

    return info


def est_titre_ou_heading(info: dict) -> tuple:
    """
    Détermine si le paragraphe est un titre/heading et son niveau.
    Retourne (est_titre, niveau, type_titre)
    """
    texte = info['texte']

    if not texte or len(texte) < 2:
        return (False, None, None)

    # Ignorer les lignes trop longues (probablement du contenu, pas un titre)
    if len(texte) > 200:
        return (False, None, None)

    # Niveau 1: MAJUSCULES + GRAS + SOULIGNE (ou combinaison)
    if info['is_caps'] and info['is_bold'] and info['is_underline']:
        return (True, 1, 'H1_CAPS_BOLD_UNDERLINE')

    # Niveau 1 alternatif: TOUT EN MAJUSCULES et centré
    if info['is_caps'] and info['is_centered']:
        return (True, 1, 'H1_CAPS_CENTERED')

    # Niveau 2: MAJUSCULES + GRAS
    if info['is_caps'] and info['is_bold']:
        return (True, 2, 'H2_CAPS_BOLD')

    # Niveau 2: MAJUSCULES seules (mais pas trop court)
    if info['is_caps'] and len(texte) > 5:
        return (True, 2, 'H2_CAPS')

    # Niveau 3: GRAS + SOULIGNÉ
    if info['is_bold'] and info['is_underline']:
        return (True, 3, 'H3_BOLD_UNDERLINE')

    # Niveau 3 alternatif: Small caps + gras
    if info['is_small_caps'] and info['is_bold']:
        return (True, 3, 'H3_SMALLCAPS_BOLD')

    # Niveau 4: GRAS seul (si pas trop long et semble être un titre)
    if info['is_bold'] and len(texte) < 100:
        # Vérifier si ça ressemble à un titre (commence par majuscule, pas de . à la fin)
        if texte[0].isupper() and not texte.endswith('.'):
            return (True, 4, 'H4_BOLD')

    # Vérifier les patterns de numérotation de section
    patterns_section = [
        r'^TITRE\s+\w+',  # TITRE PREMIER, TITRE I
        r'^CHAPITRE\s+\w+',  # CHAPITRE I
        r'^ARTICLE\s+\d+',  # ARTICLE 1
        r'^SECTION\s+\w+',  # SECTION I
        r'^[IVX]+[\s\.\-]+',  # I. II. III. etc.
        r'^\d+[\.\)\-]\s*[A-Z]',  # 1. Titre, 1) Titre
        r'^[A-Z][\.\)\-]\s*[A-Z]',  # A. Titre, A) Titre
    ]

    for pattern in patterns_section:
        if re.match(pattern, texte):
            return (True, 2, 'SECTION_NUMEROTEE')

    return (False, None, None)


def extraire_structure_document(docx_path: Path) -> dict:
    """Extrait la structure complète d'un document DOCX."""
    doc = Document(docx_path)

    structure = {
        'fichier': str(docx_path.name),
        'titres': [],
        'sections_principales': [],
        'total_paragraphes': len(doc.paragraphs)
    }

    titres_trouves = []

    for idx, para in enumerate(doc.paragraphs):
        info = analyser_style_paragraphe(para)
        est_titre, niveau, type_titre = est_titre_ou_heading(info)

        if est_titre and info['texte']:
            titre_info = {
                'index': idx,
                'texte': info['texte'],
                'niveau': niveau,
                'type': type_titre,
                'style_word': info['style_name'],
                'formatage': {
                    'bold': info['is_bold'],
                    'underline': info['is_underline'],
                    'caps': info['is_caps'],
                    'small_caps': info['is_small_caps'],
                    'centered': info['is_centered']
                }
            }
            titres_trouves.append(titre_info)

    structure['titres'] = titres_trouves

    # Identifier les sections principales (niveau 1 et 2)
    sections_principales = [t for t in titres_trouves if t['niveau'] in [1, 2]]
    structure['sections_principales'] = [t['texte'] for t in sections_principales]

    return structure


def normaliser_titre(texte: str) -> str:
    """Normalise un titre pour comparaison (sans accents, lowercase, sans ponctuation)."""
    import unicodedata
    # Supprimer accents
    texte = unicodedata.normalize('NFKD', texte).encode('ASCII', 'ignore').decode('ASCII')
    # Lowercase
    texte = texte.lower()
    # Supprimer ponctuation et espaces multiples
    texte = re.sub(r'[^\w\s]', ' ', texte)
    texte = re.sub(r'\s+', ' ', texte).strip()
    return texte


def comparer_structures(structures: list) -> dict:
    """Compare les structures de plusieurs documents."""
    # Collecter tous les titres normalisés par document
    titres_par_doc = {}
    for s in structures:
        nom = s['fichier']
        titres_par_doc[nom] = {
            'originaux': [t['texte'] for t in s['titres']],
            'normalises': [normaliser_titre(t['texte']) for t in s['titres']],
            'niveaux': {normaliser_titre(t['texte']): t['niveau'] for t in s['titres']}
        }

    # Trouver les titres communs (présents dans les 3)
    tous_titres_norm = set()
    for doc_data in titres_par_doc.values():
        tous_titres_norm.update(doc_data['normalises'])

    titres_communs = []
    titres_specifiques = {nom: [] for nom in titres_par_doc.keys()}

    for titre_norm in tous_titres_norm:
        docs_avec_titre = []
        for nom, doc_data in titres_par_doc.items():
            if titre_norm in doc_data['normalises']:
                # Retrouver le titre original
                idx = doc_data['normalises'].index(titre_norm)
                docs_avec_titre.append({
                    'doc': nom,
                    'titre_original': doc_data['originaux'][idx],
                    'niveau': doc_data['niveaux'].get(titre_norm)
                })

        if len(docs_avec_titre) == len(structures):
            # Titre commun aux 3 documents
            titres_communs.append({
                'titre_normalise': titre_norm,
                'variantes': docs_avec_titre
            })
        else:
            # Titre spécifique
            for doc_info in docs_avec_titre:
                titres_specifiques[doc_info['doc']].append({
                    'titre': doc_info['titre_original'],
                    'niveau': doc_info['niveau'],
                    'present_dans': [d['doc'] for d in docs_avec_titre]
                })

    return {
        'titres_communs': titres_communs,
        'titres_specifiques': titres_specifiques
    }


def generer_rapport(structures: list, comparaison: dict) -> str:
    """Génère un rapport textuel structuré."""
    lignes = []
    lignes.append("=" * 80)
    lignes.append("RAPPORT D'ANALYSE DES TRAMES DE PROMESSE DE VENTE")
    lignes.append("=" * 80)
    lignes.append("")

    # Section 1: Analyse individuelle de chaque document
    lignes.append("1. ANALYSE INDIVIDUELLE PAR DOCUMENT")
    lignes.append("-" * 50)

    for s in structures:
        lignes.append(f"\n### {s['fichier']}")
        lignes.append(f"    Paragraphes totaux: {s['total_paragraphes']}")
        lignes.append(f"    Titres identifies: {len(s['titres'])}")
        lignes.append(f"    Sections principales: {len(s['sections_principales'])}")
        lignes.append("")
        lignes.append("    TITRES/HEADINGS:")

        for t in s['titres']:
            indent = "    " * t['niveau']
            niveau_str = f"[H{t['niveau']}]"
            texte_tronque = t['texte'][:70] + "..." if len(t['texte']) > 70 else t['texte']
            lignes.append(f"    {indent}{niveau_str} {texte_tronque}")

        lignes.append("")

    # Section 2: Sections communes
    lignes.append("\n" + "=" * 80)
    lignes.append("2. SECTIONS COMMUNES (obligatoires)")
    lignes.append("-" * 50)

    # Trier par niveau puis ordre d'apparition
    communs = sorted(comparaison['titres_communs'],
                    key=lambda x: (x['variantes'][0]['niveau'] or 99, x['variantes'][0]['titre_original']))

    for item in communs:
        niveau = item['variantes'][0]['niveau']
        titre = item['variantes'][0]['titre_original'][:60]
        lignes.append(f"  [H{niveau}] {titre}")

    lignes.append(f"\n  TOTAL: {len(communs)} sections communes")

    # Section 3: Sections spécifiques
    lignes.append("\n" + "=" * 80)
    lignes.append("3. SECTIONS SPECIFIQUES (conditionnelles)")
    lignes.append("-" * 50)

    for doc_nom, titres in comparaison['titres_specifiques'].items():
        if titres:
            lignes.append(f"\n  ### Specifiques a {doc_nom}:")
            for t in titres:
                niveau = t['niveau']
                autres_docs = [d for d in t['present_dans'] if d != doc_nom]
                presence = f" (aussi dans: {', '.join(autres_docs)})" if autres_docs else " (UNIQUE)"
                lignes.append(f"    [H{niveau}] {t['titre'][:50]}{presence}")

    # Section 4: Ordre des sections
    lignes.append("\n" + "=" * 80)
    lignes.append("4. ORDRE DES SECTIONS (par document)")
    lignes.append("-" * 50)

    for s in structures:
        lignes.append(f"\n### {s['fichier']} - Ordre des sections H1/H2:")
        sections_h12 = [t for t in s['titres'] if t['niveau'] in [1, 2]]
        for idx, t in enumerate(sections_h12, 1):
            lignes.append(f"  {idx:2}. {t['texte'][:60]}")

    return "\n".join(lignes)


def main():
    # Chemins des fichiers
    base_dir = Path(__file__).parent.parent / "docs_original"
    fichiers = [
        base_dir / "Trame_promesse_A.docx",
        base_dir / "Trame_promesse_B.docx",
        base_dir / "Trame_promesse_C.docx"
    ]

    # Vérifier existence
    for f in fichiers:
        if not f.exists():
            print(f"[ERREUR] Fichier non trouve: {f}")
            return 1

    print("Analyse des trames de promesse de vente...")
    print("-" * 50)

    # Extraire la structure de chaque document
    structures = []
    for f in fichiers:
        print(f"Analyse de {f.name}...")
        s = extraire_structure_document(f)
        structures.append(s)
        print(f"  -> {len(s['titres'])} titres, {len(s['sections_principales'])} sections principales")

    # Comparer les structures
    print("\nComparaison des structures...")
    comparaison = comparer_structures(structures)

    # Générer le rapport
    rapport = generer_rapport(structures, comparaison)
    print("\n" + rapport)

    # Sauvegarder le rapport
    output_dir = Path(__file__).parent.parent / ".tmp"
    output_dir.mkdir(exist_ok=True)

    # Rapport texte
    rapport_path = output_dir / "rapport_analyse_promesses.txt"
    with open(rapport_path, 'w', encoding='utf-8') as f:
        f.write(rapport)
    print(f"\nRapport sauvegarde: {rapport_path}")

    # Données JSON pour traitement ultérieur
    json_path = output_dir / "structures_promesses.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            'structures': structures,
            'comparaison': comparaison
        }, f, ensure_ascii=False, indent=2)
    print(f"Donnees JSON: {json_path}")

    return 0


if __name__ == '__main__':
    exit(main())
