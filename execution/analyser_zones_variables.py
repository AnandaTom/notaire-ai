#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
analyser_zones_variables.py
===========================

Analyse un document DOCX pour identifier les zones grisées (variables)
et leurs répétitions dans le texte.

Les zones grisées dans les documents notariaux sont généralement:
- Du texte avec surlignage (highlight)
- Du texte avec ombrage (shading)
- Des champs de formulaire
- Du texte entre crochets ou accolades

Usage:
    python analyser_zones_variables.py --input <document.docx> --output <rapport.json>
"""

import argparse
import json
import re
from pathlib import Path
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX


def extraire_zones_grisees(doc_path: Path) -> dict:
    """
    Extrait toutes les zones grisées (variables) d'un document DOCX.

    Returns:
        Dictionnaire avec les zones trouvées et leur contexte
    """
    doc = Document(str(doc_path))

    zones_grisees = []
    zones_par_type = defaultdict(list)
    texte_complet = []

    # Parcourir tous les paragraphes
    for para_idx, para in enumerate(doc.paragraphs):
        texte_para = para.text
        texte_complet.append(texte_para)

        # Parcourir les runs (segments de texte avec même formatage)
        for run in para.runs:
            texte_run = run.text.strip()
            if not texte_run:
                continue

            est_grise = False
            type_gris = None

            # Vérifier le surlignage (highlight)
            if run.font.highlight_color is not None:
                est_grise = True
                type_gris = f"highlight_{run.font.highlight_color}"

            # Vérifier l'ombrage (shading) via XML
            rPr = run._element.rPr
            if rPr is not None:
                shd = rPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill.lower() not in ['auto', 'ffffff', None]:
                        est_grise = True
                        type_gris = f"shading_{fill}"

            # Vérifier la couleur de fond via XML direct
            if rPr is not None:
                highlight = rPr.find(qn('w:highlight'))
                if highlight is not None:
                    val = highlight.get(qn('w:val'))
                    if val and val.lower() not in ['none', 'white']:
                        est_grise = True
                        type_gris = f"highlight_xml_{val}"

            if est_grise:
                zone = {
                    'texte': texte_run,
                    'paragraphe_idx': para_idx,
                    'contexte': texte_para[:100] + '...' if len(texte_para) > 100 else texte_para,
                    'type': type_gris
                }
                zones_grisees.append(zone)
                zones_par_type[type_gris].append(texte_run)

    # Analyser les tableaux
    zones_tableaux = []
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        texte_run = run.text.strip()
                        if not texte_run:
                            continue

                        est_grise = False
                        type_gris = None

                        # Mêmes vérifications que pour les paragraphes
                        if run.font.highlight_color is not None:
                            est_grise = True
                            type_gris = f"highlight_{run.font.highlight_color}"

                        rPr = run._element.rPr
                        if rPr is not None:
                            shd = rPr.find(qn('w:shd'))
                            if shd is not None:
                                fill = shd.get(qn('w:fill'))
                                if fill and fill.lower() not in ['auto', 'ffffff', None]:
                                    est_grise = True
                                    type_gris = f"shading_{fill}"

                        if est_grise:
                            zone = {
                                'texte': texte_run,
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'type': type_gris,
                                'contexte': para.text[:100]
                            }
                            zones_tableaux.append(zone)
                            zones_par_type[type_gris].append(texte_run)

    return {
        'zones_paragraphes': zones_grisees,
        'zones_tableaux': zones_tableaux,
        'par_type': dict(zones_par_type),
        'texte_complet': '\n'.join(texte_complet)
    }


def detecter_patterns_variables(texte: str) -> list:
    """
    Détecte les patterns de variables dans le texte:
    - [texte entre crochets]
    - {texte entre accolades}
    - __texte avec underscores__
    - TEXTE EN MAJUSCULES isolé
    - Dates au format XX/XX/XXXX
    - Montants (chiffres avec espaces ou virgules)
    """
    patterns = []

    # Crochets
    for match in re.finditer(r'\[([^\]]+)\]', texte):
        patterns.append({
            'type': 'crochets',
            'texte': match.group(1),
            'position': match.start()
        })

    # Accolades
    for match in re.finditer(r'\{([^}]+)\}', texte):
        patterns.append({
            'type': 'accolades',
            'texte': match.group(1),
            'position': match.start()
        })

    # Points de suspension ou tirets répétés (placeholder)
    for match in re.finditer(r'\.{3,}|_{3,}|-{3,}', texte):
        patterns.append({
            'type': 'placeholder',
            'texte': match.group(0),
            'position': match.start()
        })

    return patterns


def identifier_variables_semantiques(zones: list) -> dict:
    """
    Identifie et catégorise les variables par leur sens sémantique.
    """
    categories = {
        'identite_vendeur': [],
        'identite_acquereur': [],
        'adresse': [],
        'dates': [],
        'montants': [],
        'cadastre': [],
        'copropriete': [],
        'notaire': [],
        'prets': [],
        'diagnostics': [],
        'autres': []
    }

    # Patterns de reconnaissance
    patterns_identite = [
        r'(monsieur|madame|m\.|mme)',
        r'(nom|prénom|prénoms)',
        r'(né|née|naissance)',
        r'(profession)',
        r'(nationalité)',
        r'(célibataire|marié|pacsé|divorcé|veuf)',
    ]

    patterns_adresse = [
        r'(rue|avenue|boulevard|place|allée)',
        r'(code postal|\d{5})',
        r'(demeurant|domicilié)',
    ]

    patterns_dates = [
        r'\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}',
        r'(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)',
    ]

    patterns_montants = [
        r'\d+[\s\u00a0]?\d*[\s\u00a0]?(€|euros?)',
        r'(prix|montant|somme)',
    ]

    patterns_cadastre = [
        r'(section|parcelle|lieudit|cadastr)',
        r'(tantièmes|millièmes)',
        r'(lot|cave|parking|appartement)',
    ]

    for zone in zones:
        texte_lower = zone.get('texte', '').lower()
        contexte_lower = zone.get('contexte', '').lower()

        categorise = False

        # Vérifier chaque catégorie
        if any(re.search(p, texte_lower) or re.search(p, contexte_lower) for p in patterns_identite):
            if 'vendeur' in contexte_lower or 'promettant' in contexte_lower:
                categories['identite_vendeur'].append(zone)
            elif 'acquéreur' in contexte_lower or 'bénéficiaire' in contexte_lower:
                categories['identite_acquereur'].append(zone)
            else:
                categories['identite_vendeur'].append(zone)  # Default
            categorise = True

        if any(re.search(p, texte_lower) for p in patterns_adresse):
            categories['adresse'].append(zone)
            categorise = True

        if any(re.search(p, texte_lower) for p in patterns_dates):
            categories['dates'].append(zone)
            categorise = True

        if any(re.search(p, texte_lower) for p in patterns_montants):
            categories['montants'].append(zone)
            categorise = True

        if any(re.search(p, texte_lower) or re.search(p, contexte_lower) for p in patterns_cadastre):
            categories['cadastre'].append(zone)
            categorise = True

        if not categorise:
            categories['autres'].append(zone)

    return categories


def trouver_repetitions(zones_grisees: list, texte_complet: str) -> list:
    """
    Trouve les répétitions des variables grisées dans le texte non grisé.
    """
    repetitions = []

    for zone in zones_grisees:
        texte_var = zone.get('texte', '').strip()
        if len(texte_var) < 3:  # Ignorer les textes trop courts
            continue

        # Compter les occurrences dans le texte complet
        count = texte_complet.count(texte_var)

        if count > 1:  # Plus d'une occurrence = répétition
            # Trouver les positions
            positions = []
            start = 0
            while True:
                pos = texte_complet.find(texte_var, start)
                if pos == -1:
                    break
                positions.append(pos)
                start = pos + 1

            repetitions.append({
                'variable': texte_var,
                'occurrences': count,
                'positions': positions,
                'zone_origine': zone
            })

    return repetitions


def generer_rapport(doc_path: Path, output_path: Path = None):
    """
    Génère un rapport complet d'analyse des variables.
    """
    print(f"Analyse de: {doc_path}")

    # Extraire les zones grisées
    resultats = extraire_zones_grisees(doc_path)

    # Détecter les patterns dans le texte complet
    patterns = detecter_patterns_variables(resultats['texte_complet'])

    # Catégoriser les variables
    toutes_zones = resultats['zones_paragraphes'] + resultats['zones_tableaux']
    categories = identifier_variables_semantiques(toutes_zones)

    # Trouver les répétitions
    repetitions = trouver_repetitions(toutes_zones, resultats['texte_complet'])

    rapport = {
        'fichier_analyse': str(doc_path),
        'statistiques': {
            'zones_grisees_paragraphes': len(resultats['zones_paragraphes']),
            'zones_grisees_tableaux': len(resultats['zones_tableaux']),
            'patterns_detectes': len(patterns),
            'repetitions_trouvees': len(repetitions)
        },
        'zones_par_type': resultats['par_type'],
        'zones_paragraphes': resultats['zones_paragraphes'],
        'zones_tableaux': resultats['zones_tableaux'],
        'patterns': patterns,
        'categories': {k: len(v) for k, v in categories.items()},
        'categories_detail': {k: v for k, v in categories.items() if v},
        'repetitions': repetitions
    }

    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(rapport, f, ensure_ascii=False, indent=2)
        print(f"Rapport sauvegardé: {output_path}")

    # Afficher un résumé
    print("\n" + "="*60)
    print("RÉSUMÉ DE L'ANALYSE")
    print("="*60)
    print(f"Zones grisées dans paragraphes: {rapport['statistiques']['zones_grisees_paragraphes']}")
    print(f"Zones grisées dans tableaux: {rapport['statistiques']['zones_grisees_tableaux']}")
    print(f"Patterns détectés: {rapport['statistiques']['patterns_detectes']}")
    print(f"Variables avec répétitions: {rapport['statistiques']['repetitions_trouvees']}")

    print("\nTypes de zones trouvées:")
    for type_zone, zones in resultats['par_type'].items():
        print(f"  - {type_zone}: {len(zones)} occurrences")

    print("\nCatégories de variables:")
    for cat, count in rapport['categories'].items():
        if count > 0:
            print(f"  - {cat}: {count}")

    return rapport


def main():
    parser = argparse.ArgumentParser(
        description="Analyse les zones grisées (variables) d'un document DOCX"
    )
    parser.add_argument(
        '--input', '-i',
        type=Path,
        required=True,
        help="Document DOCX à analyser"
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help="Fichier JSON de sortie pour le rapport"
    )

    args = parser.parse_args()

    if not args.input.exists():
        print(f"[ERREUR] Fichier non trouvé: {args.input}")
        return 1

    rapport = generer_rapport(args.input, args.output)
    return 0


if __name__ == '__main__':
    exit(main())
