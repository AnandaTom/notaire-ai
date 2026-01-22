#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
exporter_docx_officiel.py
=========================

Exporte un acte Markdown vers DOCX avec les VRAIES marges et tailles de police du document original.

Basé sur l'analyse du DOCX original :
- Times New Roman 11pt (texte principal)
- Times New Roman 9pt (certaines sections)
- Marges standards : G=25.4mm, D=10mm, H=6.14mm, B=15.56mm
- Alternance pages paires/impaires avec marges variables
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import comtypes.client
import pythoncom

def set_cell_border(cell, **kwargs):
    """
    Définit les bordures d'une cellule de tableau.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Liste des bordures possibles
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))
            tcPr.append(element)


def convertir_markdown_vers_docx(input_md, output_docx):
    """
    Convertit un fichier Markdown en DOCX avec formatage fidèle à l'original.

    Marges du document original :
    - Gauche : 25.4 mm (pages impaires), 12-15 mm (pages paires)
    - Droite : 10 mm
    - Haut : 6.14 mm
    - Bas : 15.56 mm (varie entre 12-25 mm)
    """

    # Lire le Markdown
    md_path = Path(input_md)
    content = md_path.read_text(encoding='utf-8')

    # Créer le document Word
    doc = Document()

    # Configuration de la première section
    section = doc.sections[0]
    section.page_height = Mm(297)  # A4
    section.page_width = Mm(210)
    section.left_margin = Mm(25.4)   # Marge gauche standard
    section.right_margin = Mm(10)     # Marge droite réduite
    section.top_margin = Mm(6.14)     # Marge haute réduite
    section.bottom_margin = Mm(15.56) # Marge basse standard

    # Style Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_normal.paragraph_format.space_after = Pt(0.8)  # Ajusté pour atteindre exactement 28 pages
    style_normal.paragraph_format.line_spacing = 1.0
    style_normal.paragraph_format.first_line_indent = Mm(12.51)

    # Style Heading 1 (## Titres majeurs)
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(11)
    style_h1.font.bold = True
    style_h1.font.all_caps = True
    style_h1.font.underline = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h1.paragraph_format.space_before = Pt(6)  # Réduit de 12 à 6
    style_h1.paragraph_format.space_after = Pt(3)   # Réduit de 6 à 3
    style_h1.paragraph_format.first_line_indent = Mm(0)

    # Style Heading 2 (### Sous-titres)
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(11)
    style_h2.font.bold = True
    style_h2.font.small_caps = True
    style_h2.font.underline = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    style_h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h2.paragraph_format.space_before = Pt(6)  # Réduit de 12 à 6
    style_h2.paragraph_format.space_after = Pt(3)   # Réduit de 6 à 3
    style_h2.paragraph_format.first_line_indent = Mm(0)

    # Style Heading 3 (#### Titres de section)
    style_h3 = doc.styles['Heading 3']
    style_h3.font.name = 'Times New Roman'
    style_h3.font.size = Pt(11)
    style_h3.font.bold = True
    style_h3.font.underline = True
    style_h3.font.color.rgb = RGBColor(0, 0, 0)
    style_h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h3.paragraph_format.space_before = Pt(3)  # Réduit de 6 à 3
    style_h3.paragraph_format.space_after = Pt(3)   # Réduit de 6 à 3
    style_h3.paragraph_format.first_line_indent = Mm(0)

    # Style Heading 4
    style_h4 = doc.styles['Heading 4']
    style_h4.font.name = 'Times New Roman'
    style_h4.font.size = Pt(11)
    style_h4.font.bold = True
    style_h4.font.color.rgb = RGBColor(0, 0, 0)
    style_h4.paragraph_format.space_before = Pt(3)  # Réduit de 6 à 3
    style_h4.paragraph_format.space_after = Pt(0)   # Réduit de 3 à 0
    style_h4.paragraph_format.first_line_indent = Mm(12.51)

    # Traiter ligne par ligne
    lines = content.split('\n')
    i = 0
    in_box = False
    box_content = []
    current_box_name = None
    first_page_header_lines = []
    in_first_page_header = False

    while i < len(lines):
        line = lines[i]

        # Gestion des marqueurs spéciaux
        if '{FIRST_PAGE_HEADER_START}' in line:
            in_first_page_header = True
            i += 1
            continue

        if '{FIRST_PAGE_HEADER_END}' in line:
            in_first_page_header = False
            # Créer le header de la première page
            if first_page_header_lines:
                for header_line in first_page_header_lines:
                    if header_line.strip():
                        para = doc.add_paragraph(header_line.strip())
                        para.style = style_normal
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.first_line_indent = Mm(0)
                        for run in para.runs:
                            run.font.size = Pt(9)  # Taille réduite pour header
                first_page_header_lines = []
            i += 1
            continue

        if in_first_page_header:
            first_page_header_lines.append(line)
            i += 1
            continue

        # Gestion des encadrés
        if '_BOX_START}' in line:
            in_box = True
            current_box_name = line.strip('{} ')
            box_content = []
            i += 1
            continue

        if '_BOX_END}' in line:
            in_box = False
            # Créer le tableau pour l'encadré
            if box_content:
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                table.allow_autofit = False

                cell = table.cell(0, 0)
                cell.width = Mm(160)

                # Ajouter le contenu
                for box_line in box_content:
                    if box_line.startswith('## '):
                        p = cell.paragraphs[0] if len(cell.paragraphs) == 1 else cell.add_paragraph()
                        p.text = box_line[3:].strip()
                        p.style = style_h1
                    elif box_line.startswith('### '):
                        p = cell.add_paragraph()
                        p.text = box_line[4:].strip()
                        p.style = style_h2
                    else:
                        p = cell.add_paragraph()
                        p.text = box_line.strip()

                # Bordures
                set_cell_border(
                    cell,
                    top={"val": "single", "sz": "12", "color": "000000"},
                    bottom={"val": "single", "sz": "12", "color": "000000"},
                    left={"val": "single", "sz": "12", "color": "000000"},
                    right={"val": "single", "sz": "12", "color": "000000"}
                )

                # Espacement après
                doc.add_paragraph()

            box_content = []
            current_box_name = None
            i += 1
            continue

        if in_box:
            box_content.append(line)
            i += 1
            continue

        # Ligne vide
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Ligne d'astérisques (séparateur)
        if re.match(r'^\*{5,}$', line.strip()):
            para = doc.add_paragraph(line.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Mm(0)
            i += 1
            continue

        # Titres
        if line.startswith('#### '):
            para = doc.add_paragraph(line[5:].strip())
            para.style = style_h4
        elif line.startswith('### '):
            para = doc.add_paragraph(line[4:].strip())
            para.style = style_h3
        elif line.startswith('## '):
            para = doc.add_paragraph(line[3:].strip())
            # Détecter si c'est un titre majeur ou sous-titre
            if line[3:].strip().isupper():
                para.style = style_h1
            else:
                para.style = style_h2
        # Texte en gras
        elif line.startswith('**') and line.endswith('**'):
            para = doc.add_paragraph()
            run = para.add_run(line.strip('*').strip())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            para.paragraph_format.first_line_indent = Mm(0)
            para.paragraph_format.space_after = Pt(0)  # Réduit de 3 à 0
        # Texte normal
        else:
            # Détecter le gras inline
            if '**' in line:
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part.strip('*'))
                        run.bold = True
                    else:
                        run = para.add_run(part)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
            else:
                para = doc.add_paragraph(line.strip())
                para.style = style_normal

        i += 1

    # Sauvegarder
    doc.save(output_docx)
    print(f'✅ DOCX créé: {output_docx}')
    return output_docx


def convertir_docx_vers_pdf(docx_path, pdf_path):
    """
    Convertit un fichier DOCX en PDF via Microsoft Word COM.
    """
    try:
        pythoncom.CoInitialize()

        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        docx_abs = str(Path(docx_path).resolve())
        pdf_abs = str(Path(pdf_path).resolve())

        doc = word.Documents.Open(docx_abs)
        doc.SaveAs(pdf_abs, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close()
        word.Quit()

        print(f'✅ PDF créé: {pdf_path}')

    except Exception as e:
        print(f'❌ Erreur conversion PDF: {e}')
        raise
    finally:
        pythoncom.CoUninitialize()


if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(description='Exporte Markdown vers DOCX/PDF (marges originales)')
    parser.add_argument('--input', required=True, help='Fichier Markdown source')
    parser.add_argument('--output', required=True, help='Fichier DOCX de sortie')
    parser.add_argument('--pdf', action='store_true', help='Générer aussi le PDF')

    args = parser.parse_args()

    # Générer DOCX
    docx_path = convertir_markdown_vers_docx(args.input, args.output)

    # Générer PDF si demandé
    if args.pdf:
        pdf_path = str(Path(args.output).with_suffix('.pdf'))
        convertir_docx_vers_pdf(docx_path, pdf_path)
