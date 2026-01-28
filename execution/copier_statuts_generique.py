#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
copier_statuts_generique.py
============================

Copie le document Statuts (1).docx en remplaçant toutes les valeurs spécifiques
par des variables génériques, tout en préservant 100% du formatage.
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from pathlib import Path
import comtypes.client
import pythoncom


def remplacer_valeurs_generiques(doc):
    """Remplace toutes les valeurs spécifiques par des variables génériques."""

    # Dictionnaire complet des remplacements
    remplacements = {
        # Société
        'BLOUGE': 'NOM_SOCIETE',
        'Blouge': 'Nom_Societe',
        '985 331 354': '000 000 000',
        '985 331 354 R.C.S. LYON': '000 000 000 R.C.S. METROPOLE',

        # Personnes - Noms de famille
        'AUVRAY': 'NOM_FAMILLE',
        'Auvray': 'Nom_Famille',
        'ROBIN': 'NOM_NAISSANCE',
        'Robin': 'Nom_Naissance',

        # Notaires et professionnels
        'COHENDET': 'NOM_NOTAIRE',
        'Cohendet': 'Nom_Notaire',
        'Pierre': 'Prénom_Notaire',
        'GIEN': 'NOM_AVOCAT',
        'Gien': 'Nom_Avocat',

        # Prénoms
        'Dominique': 'Prénom1',
        'Christian': 'Prénom3',
        'Viviane': 'Prénom2',
        'Blanche': 'Prénom6',
        'Célie': 'Prénom7',

        # Adresses - Villes
        'DARDILLY': 'VILLE',
        'Dardilly': 'Ville',
        'PARIS': 'METROPOLE',
        'Paris': 'Metropole',
        "TAIN L'HERMITAGE": 'VILLE_NAISSANCE',
        "Tain l'Hermitage": 'Ville_Naissance',
        'CHANTEMERLE-LES-BLES': 'VILLE_MARIAGE',
        'Chantemerle-les-Bles': 'Ville_Mariage',
        'Charolles': 'Ville_Notaire',
        'LYON': 'METROPOLE',
        'Lyon': 'Metropole',

        # Adresses - Rues
        '29 Chemin du Panorama': '00 rue Exemple',
        '29 chemin du Panorama': '00 rue Exemple',

        # Codes postaux
        '69570': '00000',
        '75012': '00001',
        '26600': '00002',
        '71120': '00003',

        # Dates
        '20 février 1961': 'JJ mois AAAA',
        '16 avril 1961': 'JJ mois AAAA',
        '09 septembre 1989': 'JJ mois AAAA',
        '24 juin 1989': 'JJ mois AAAA',

        # Organismes
        'FIDUCIAL BANQUE': 'BANQUE_EXEMPLE',
        'Fiducial Banque': 'Banque_Exemple',
        'FIDUCIAL SOFIRAL': 'CABINET_EXEMPLE',
        'Fiducial Sofiral': 'Cabinet_Exemple',
        'FIDUCIAL': 'CABINET',
        'Fiducial': 'Cabinet',

        # Emails
        'lyon.avocat@fiducial.fr': 'email@exemple.fr',
    }

    count_paras = 0
    count_tables = 0

    # Remplacer dans les paragraphes
    for para in doc.paragraphs:
        for run in para.runs:
            texte_original = run.text
            texte_modifie = texte_original

            for ancien, nouveau in remplacements.items():
                texte_modifie = texte_modifie.replace(ancien, nouveau)

            if texte_modifie != texte_original:
                run.text = texte_modifie
                count_paras += 1

    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        texte_original = run.text
                        texte_modifie = texte_original

                        for ancien, nouveau in remplacements.items():
                            texte_modifie = texte_modifie.replace(ancien, nouveau)

                        if texte_modifie != texte_original:
                            run.text = texte_modifie
                            count_tables += 1

    return count_paras, count_tables


def generer_statuts_officiel():
    """Génère le document Statuts Officiel à partir de l'original."""

    print('=' * 80)
    print('GÉNÉRATION STATUTS OFFICIEL')
    print('=' * 80)

    # Charger le document original
    print('\n1. Chargement du document original...')
    doc_original = Document('docs_original/Statuts (1).docx')
    print(f'   ✅ Chargé: {len(doc_original.paragraphs)} paragraphes, {len(doc_original.tables)} tableaux')

    # Remplacer les valeurs
    print('\n2. Remplacement des valeurs spécifiques...')
    count_paras, count_tables = remplacer_valeurs_generiques(doc_original)
    print(f'   ✅ Paragraphes modifiés: {count_paras}')
    print(f'   ✅ Cellules de tableaux modifiées: {count_tables}')

    # Sauvegarder DOCX
    print('\n3. Sauvegarde du document...')
    output_dir = Path('outputs')
    output_dir.mkdir(exist_ok=True)

    docx_path = output_dir / 'Statuts Officiel.docx'
    doc_original.save(str(docx_path))
    print(f'   ✅ DOCX sauvegardé: {docx_path}')

    # Convertir en PDF
    print('\n4. Conversion en PDF...')
    pdf_path = output_dir / 'Statuts Officiel.pdf'

    pythoncom.CoInitialize()
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False

        docx_abs = str(docx_path.resolve())
        pdf_abs = str(pdf_path.resolve())

        doc = word.Documents.Open(docx_abs)
        doc.SaveAs(pdf_abs, FileFormat=17)
        doc.Close()
        word.Quit()

        print(f'   ✅ PDF créé: {pdf_path}')
    except Exception as e:
        print(f'   ❌ Erreur conversion PDF: {e}')
        raise
    finally:
        pythoncom.CoUninitialize()

    print('\n' + '=' * 80)
    print('✅ GÉNÉRATION TERMINÉE')
    print('=' * 80)
    print(f'\nFichiers créés:')
    print(f'  - {docx_path}')
    print(f'  - {pdf_path}')

    return docx_path, pdf_path


if __name__ == '__main__':
    generer_statuts_officiel()
