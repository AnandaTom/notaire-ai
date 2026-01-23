#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generer_statuts_propre.py
==========================

Génère les statuts avec:
1. Variables génériques
2. Fond bleu sur tableau STATUTS (table 1)
3. Suppression paragraphes vides qui créent page blanche
4. Réduction taille tableau STATUTS pour qu'il soit moins gros
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm
import comtypes.client
import pythoncom


def set_cell_background(cell, color):
    """Définit la couleur de fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def generer_statuts():
    """Génère le document Statuts."""

    print('=' * 80)
    print('GENERATION STATUTS PROPRE')
    print('=' * 80)

    # Charger le document original
    print('\n1. Chargement document original...')
    doc = Document('docs_originels/Statuts (1).docx')
    print(f'   OK - {len(doc.paragraphs)} paragraphes')

    # Dictionnaire de remplacements
    remplacements = {
        'FIDUCIAL SOFIRAL': 'CABINET_EXEMPLE',
        'Fiducial Sofiral': 'Cabinet_Exemple',
        'FIDUCIAL BANQUE': 'BANQUE_EXEMPLE',
        'Fiducial Banque': 'Banque_Exemple',
        'FIDUCIAL': 'CABINET',
        'Fiducial': 'Cabinet',
        '29 Chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 chemin du Panorama - 69570 DARDILLY': '00 rue Exemple - 00000 VILLE',
        '29 Chemin du Panorama': '00 rue Exemple',
        '29 chemin du Panorama': '00 rue Exemple',
        "TAIN L'HERMITAGE": 'VILLE_NAISSANCE',
        "Tain l'Hermitage": 'Ville_Naissance',
        'CHANTEMERLE-LES-BLES': 'VILLE_MARIAGE',
        'Chantemerle-les-Bles': 'Ville_Mariage',
        '985 331 354 R.C.S. LYON': '000 000 000 R.C.S. METROPOLE',
        '985 331 354': '000 000 000',
        'BLOUGE': 'NOM_SOCIETE',
        'Blouge': 'Nom_Societe',
        'AUVRAY': 'NOM_FAMILLE',
        'Auvray': 'Nom_Famille',
        'ROBIN': 'NOM_NAISSANCE',
        'Robin': 'Nom_Naissance',
        'COHENDET': 'NOM_NOTAIRE',
        'Cohendet': 'Nom_Notaire',
        'GIEN': 'NOM_AVOCAT',
        'Gien': 'Nom_Avocat',
        'Dominique': 'Prenom1',
        'Christian': 'Prenom3',
        'Viviane': 'Prenom2',
        'Blanche': 'Prenom6',
        'Célie': 'Prenom7',
        'Pierre': 'Prenom_Notaire',
        'DARDILLY': 'VILLE',
        'Dardilly': 'Ville',
        'PARIS': 'METROPOLE',
        'Paris': 'Metropole',
        'LYON': 'METROPOLE',
        'Lyon': 'Metropole',
        'Charolles': 'Ville_Notaire',
        '69570': '00000',
        '75012': '00001',
        '26600': '00002',
        '71120': '00003',
        '20 février 1961': 'JJ mois AAAA',
        '16 avril 1961': 'JJ mois AAAA',
        '09 septembre 1989': 'JJ mois AAAA',
        '24 juin 1989': 'JJ mois AAAA',
        'lyon.avocat@fiducial.fr': 'email@exemple.fr',
    }

    # Remplacer dans paragraphes et tableaux
    print('\n2. Remplacement variables...')
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for ancien, nouveau in remplacements.items():
                if ancien in run.text:
                    run.text = run.text.replace(ancien, nouveau)
                    count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for ancien, nouveau in remplacements.items():
                            if ancien in run.text:
                                run.text = run.text.replace(ancien, nouveau)
                                count += 1

    print(f'   {count} modifications')

    # Ajouter fond bleu au tableau STATUTS
    print('\n3. Formatage tableau STATUTS...')
    table_statuts = doc.tables[1]
    cell_statuts = table_statuts.cell(0, 0)
    set_cell_background(cell_statuts, '001f5f')

    # Réduire largeur ET hauteur du tableau
    table_statuts.autofit = False
    table_statuts.allow_autofit = False

    # Réduire hauteur en réduisant le padding et la police
    from docx.shared import Pt
    for para in cell_statuts.paragraphs:
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        for run in para.runs:
            run.font.size = Pt(10)  # Réduire de 11 à 10pt

    # Largeur réduite
    for row in table_statuts.rows:
        for cell in row.cells:
            cell.width = Mm(100)  # Plus petit

    print('   Fond bleu + taille réduite')

    # Supprimer TOUS les paragraphes vides 2-6
    print('\n4. Suppression paragraphes vides...')
    for i in range(6, 1, -1):  # 6, 5, 4, 3, 2
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)

    print('   Paragraphes vides supprimés (page blanche enlevée)')

    # Sauvegarder
    print('\n5. Sauvegarde...')
    output_dir = Path('outputs')
    output_dir.mkdir(exist_ok=True)

    docx_path = output_dir / 'officiel statuts.docx'
    doc.save(str(docx_path))
    print(f'   {docx_path}')

    # Convertir en PDF
    print('\n6. Conversion PDF...')
    pdf_path = output_dir / 'officiel statuts.pdf'

    pythoncom.CoInitialize()
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc_word = word.Documents.Open(str(docx_path.resolve()))
        doc_word.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        doc_word.Close()
        word.Quit()
        print(f'   {pdf_path}')
    finally:
        pythoncom.CoUninitialize()

    print('\n' + '=' * 80)
    print('TERMINE')
    print('=' * 80)


if __name__ == '__main__':
    generer_statuts()
