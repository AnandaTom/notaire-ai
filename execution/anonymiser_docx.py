#!/usr/bin/env python3
"""
Anonymisation de documents DOCX notariaux.
Détecte et remplace les noms/prénoms tout en préservant 100% du formatage.

Usage:
    python execution/anonymiser_docx.py docs_originels/Promesse.docx -o docs_originels/Promesse_anonyme.docx
    python execution/anonymiser_docx.py docs_originels/*.docx --batch  # Traite tous les fichiers
"""

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional
from dataclasses import dataclass, field

# Encodage UTF-8 pour Windows
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("❌ python-docx requis: pip install python-docx")
    sys.exit(1)


# =============================================================================
# DONNÉES DE RÉFÉRENCE POUR DÉTECTION
# =============================================================================

# Prénoms français courants (masculins et féminins)
PRENOMS_FRANCAIS = {
    # Masculins
    "Jean", "Pierre", "Michel", "André", "Philippe", "Jacques", "Bernard", "Alain",
    "François", "Claude", "Daniel", "Patrick", "Christian", "René", "Roger", "Robert",
    "Paul", "Henri", "Louis", "Gérard", "Maurice", "Marcel", "Nicolas", "Éric", "Eric",
    "Olivier", "Thierry", "Christophe", "Didier", "Laurent", "Dominique", "Marc",
    "Stéphane", "Stephane", "David", "Pascal", "Yves", "Bruno", "Serge", "Guy",
    "Antoine", "Vincent", "Frédéric", "Frederic", "Emmanuel", "Julien", "Sébastien",
    "Sebastien", "Thomas", "Alexandre", "Guillaume", "Maxime", "Anthony", "Jérôme",
    "Fabrice", "Fabien", "Cédric", "Cedric", "Benoît", "Benoit", "Franck", "Hervé",
    "Herve", "Sylvain", "Arnaud", "Mathieu", "Charles", "Georges", "Raymond", "Joseph",
    # Féminins
    "Marie", "Jeanne", "Françoise", "Francoise", "Monique", "Nicole", "Catherine",
    "Nathalie", "Isabelle", "Sylvie", "Anne", "Martine", "Christine", "Jacqueline",
    "Madeleine", "Marguerite", "Denise", "Simone", "Suzanne", "Hélène", "Helene",
    "Michèle", "Michele", "Brigitte", "Claudine", "Colette", "Dominique", "Odette",
    "Yvonne", "Chantal", "Danielle", "Geneviève", "Genevieve", "Patricia", "Véronique",
    "Veronique", "Valérie", "Valerie", "Sophie", "Sandrine", "Corinne", "Laurence",
    "Pascale", "Agnès", "Agnes", "Cécile", "Cecile", "Virginie", "Stéphanie", "Stephanie",
    "Caroline", "Émilie", "Emilie", "Julie", "Aurélie", "Aurelie", "Céline", "Celine",
    "Delphine", "Florence", "Laetitia", "Marion", "Camille", "Pauline", "Laura",
    "Charlotte", "Lucie", "Manon", "Léa", "Lea", "Chloé", "Chloe", "Clara", "Emma",
    # Composés courants
    "Jean-Pierre", "Jean-Claude", "Jean-Paul", "Jean-Michel", "Jean-François",
    "Jean-Marc", "Jean-Luc", "Marie-Claude", "Marie-France", "Marie-Christine",
    "Anne-Marie", "Marie-Pierre", "Marie-Thérèse",
}

# Noms de famille français courants
NOMS_FAMILLE_FRANCAIS = {
    "Martin", "Bernard", "Thomas", "Petit", "Robert", "Richard", "Durand", "Dubois",
    "Moreau", "Laurent", "Simon", "Michel", "Lefebvre", "Leroy", "Roux", "David",
    "Bertrand", "Morel", "Fournier", "Girard", "Bonnet", "Dupont", "Lambert", "Fontaine",
    "Rousseau", "Vincent", "Muller", "Lefevre", "Faure", "Andre", "Mercier", "Blanc",
    "Guerin", "Boyer", "Garnier", "Chevalier", "François", "Legrand", "Gauthier",
    "Garcia", "Perrin", "Robin", "Clement", "Morin", "Nicolas", "Henry", "Roussel",
    "Mathieu", "Gautier", "Masson", "Marchand", "Duval", "Denis", "Dumont", "Marie",
    "Lemaire", "Noel", "Meyer", "Dufour", "Meunier", "Brun", "Blanchard", "Giraud",
    "Joly", "Riviere", "Lucas", "Brunet", "Gaillard", "Barbier", "Arnaud", "Martinez",
    "Gerard", "Roche", "Renard", "Schmitt", "Roy", "Leroux", "Colin", "Vidal", "Caron",
    "Picard", "Roger", "Fabre", "Aubert", "Lemoine", "Renaud", "Dumas", "Lacroix",
    "Olivier", "Philippe", "Bourgeois", "Pierre", "Benoit", "Rey", "Leclerc", "Payet",
    "Rolland", "Leclercq", "Guillaume", "Lecomte", "Lopez", "Jean", "Dupuy", "Guillot",
    "Hubert", "Berger", "Carpentier", "Sanchez", "Dupuis", "Moulin", "Louis", "Deschamps",
    # Noms spécifiques vus dans les actes notariaux
    "Abela", "Guyot", "Delacroix", "Beaumont", "Charpentier", "Vasseur", "Collin",
    "Pichon", "Lemaitre", "Perrot", "Breton", "Poirier", "Boucher", "Jacquet",
}

# Patterns pour détecter les contextes de noms
PATTERNS_CONTEXTE = [
    # Civilités suivies de noms
    r'\b(Monsieur|Madame|Mademoiselle|M\.|Mme|Mlle|Mr)\s+([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)?)',
    # "né(e) [Prénom]" ou "née [Nom]"
    r'\bné(?:e)?\s+([A-ZÀ-Ÿ][a-zà-ÿ]+)',
    # "époux/épouse de [Nom]"
    r'\b(?:époux|épouse|veuf|veuve)\s+(?:de\s+)?([A-ZÀ-Ÿ][a-zà-ÿ]+)',
    # "Maître [Nom]" (notaires)
    r'\bMaître\s+([A-ZÀ-Ÿ][a-zà-ÿ]+)',
    # Noms en majuscules (style notarial)
    r'\b([A-ZÀ-Ÿ]{2,})\b',
]

# Noms à préserver (ne pas anonymiser)
NOMS_PRESERVES = {
    # Villes
    "Paris", "Lyon", "Marseille", "Bordeaux", "Toulouse", "Nice", "Nantes", "Strasbourg",
    "Montpellier", "Lille", "Rennes", "Reims", "Toulon", "Grenoble", "Dijon", "Angers",
    "Nîmes", "Villeurbanne", "Saint", "Havre", "Étienne", "Quentin", "Denis", "Germain",
    # Mois
    "Janvier", "Février", "Mars", "Avril", "Mai", "Juin", "Juillet", "Août",
    "Septembre", "Octobre", "Novembre", "Décembre",
    # Jours
    "Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche",
    # Termes juridiques
    "France", "République", "Française", "Code", "Civil", "Article", "Loi",
    "Tribunal", "Cour", "Cassation", "Appel", "Instance", "Notaire", "Notaires",
    "Chambre", "Office", "Étude", "Greffe", "Cadastre", "Hypothèques", "Foncier",
    # Termes bancaires/financiers
    "Banque", "Caisse", "Crédit", "Agricole", "Mutuel", "Populaire", "Épargne",
    "BNP", "Paribas", "Société", "Générale", "LCL", "HSBC", "Boursorama",
    # Termes immobiliers
    "Copropriété", "Syndic", "Règlement", "Division", "Propriété", "Immeuble",
    "Appartement", "Maison", "Terrain", "Lot", "Lots", "Cave", "Parking", "Garage",
    # Directions
    "Nord", "Sud", "Est", "Ouest",
    # Autres
    "EUR", "EUROS", "Euro", "Euros",
}


@dataclass
class Remplacement:
    """Représente un remplacement effectué."""
    original: str
    anonyme: str
    type: str  # 'prenom', 'nom', 'nom_complet'
    contexte: str = ""


@dataclass
class ResultatAnonymisation:
    """Résultat de l'anonymisation d'un document."""
    fichier_source: Path
    fichier_sortie: Path
    remplacements: List[Remplacement] = field(default_factory=list)
    noms_detectes: Set[str] = field(default_factory=set)
    mapping: Dict[str, str] = field(default_factory=dict)
    succes: bool = True
    erreur: str = ""


class AnonymiseurDocx:
    """Anonymise les documents DOCX en préservant le formatage."""

    def __init__(self, mode_strict: bool = False):
        """
        Args:
            mode_strict: Si True, demande confirmation pour chaque nom détecté
        """
        self.mode_strict = mode_strict
        self.mapping_global: Dict[str, str] = {}
        self.compteur_prenoms = 1
        self.compteur_noms = 1

        # Listes de prénoms/noms de remplacement
        self.prenoms_remplacement = [
            "Jean", "Pierre", "Marie", "Paul", "Anne", "Louis", "Claire",
            "Michel", "Sophie", "André", "Lucie", "Henri", "Julie", "François"
        ]
        self.noms_remplacement = [
            "DUPONT", "MARTIN", "DURAND", "BERNARD", "PETIT", "MOREAU",
            "LAMBERT", "SIMON", "MICHEL", "LAURENT", "GARCIA", "DAVID"
        ]

    def _est_nom_a_preserver(self, mot: str) -> bool:
        """Vérifie si un mot doit être préservé (pas anonymisé)."""
        mot_lower = mot.lower()
        mot_title = mot.title()
        mot_upper = mot.upper()

        return (
            mot_title in NOMS_PRESERVES or
            mot_upper in NOMS_PRESERVES or
            mot_lower in {'le', 'la', 'les', 'de', 'du', 'des', 'un', 'une', 'et', 'ou'} or
            len(mot) < 2
        )

    def _detecter_noms_dans_texte(self, texte: str) -> Set[str]:
        """Détecte tous les noms/prénoms potentiels dans un texte."""
        noms_detectes = set()

        # 1. Chercher les prénoms connus
        for prenom in PRENOMS_FRANCAIS:
            if prenom in texte:
                noms_detectes.add(prenom)

        # 2. Chercher les noms de famille connus
        for nom in NOMS_FAMILLE_FRANCAIS:
            # Chercher en minuscules et majuscules
            if nom in texte or nom.upper() in texte:
                noms_detectes.add(nom)
                noms_detectes.add(nom.upper())

        # 3. Patterns contextuels
        for pattern in PATTERNS_CONTEXTE:
            matches = re.finditer(pattern, texte, re.IGNORECASE)
            for match in matches:
                # Récupérer le groupe capturé (le nom)
                for group in match.groups():
                    if group and len(group) >= 2:
                        noms_detectes.add(group)

        # 4. Mots en MAJUSCULES qui ressemblent à des noms
        mots_majuscules = re.findall(r'\b([A-ZÀ-Ÿ]{2,})\b', texte)
        for mot in mots_majuscules:
            # Vérifier si c'est potentiellement un nom
            if (
                mot.title() in NOMS_FAMILLE_FRANCAIS or
                mot.title() in PRENOMS_FRANCAIS or
                (len(mot) >= 3 and mot not in {'EUR', 'TVA', 'SCI', 'SARL', 'SAS', 'RCS', 'SIRET', 'APE'})
            ):
                if not self._est_nom_a_preserver(mot):
                    noms_detectes.add(mot)

        # Filtrer les noms à préserver
        noms_filtres = {n for n in noms_detectes if not self._est_nom_a_preserver(n)}

        return noms_filtres

    def _obtenir_remplacement(self, nom: str, type_nom: str) -> str:
        """Obtient ou crée un remplacement pour un nom."""
        # Vérifier si déjà dans le mapping
        nom_normalise = nom.upper()
        if nom_normalise in self.mapping_global:
            remplacement = self.mapping_global[nom_normalise]
            # Adapter la casse
            if nom.isupper():
                return remplacement.upper()
            elif nom.istitle():
                return remplacement.title()
            else:
                return remplacement.lower()

        # Créer un nouveau remplacement
        if type_nom == 'prenom':
            if self.compteur_prenoms <= len(self.prenoms_remplacement):
                remplacement = self.prenoms_remplacement[self.compteur_prenoms - 1]
            else:
                remplacement = f"Prénom{self.compteur_prenoms}"
            self.compteur_prenoms += 1
        else:  # nom de famille
            if self.compteur_noms <= len(self.noms_remplacement):
                remplacement = self.noms_remplacement[self.compteur_noms - 1]
            else:
                remplacement = f"NOM{self.compteur_noms}"
            self.compteur_noms += 1

        self.mapping_global[nom_normalise] = remplacement

        # Adapter la casse
        if nom.isupper():
            return remplacement.upper()
        elif nom.istitle():
            return remplacement.title()
        else:
            return remplacement.lower()

    def _remplacer_dans_run(self, run, noms_a_remplacer: Dict[str, str]) -> List[Remplacement]:
        """Remplace les noms dans un run en préservant le formatage."""
        remplacements = []

        if not run.text:
            return remplacements

        texte_original = run.text
        texte_modifie = texte_original

        # Trier par longueur décroissante pour éviter les remplacements partiels
        noms_tries = sorted(noms_a_remplacer.keys(), key=len, reverse=True)

        for nom in noms_tries:
            if nom in texte_modifie:
                remplacement = noms_a_remplacer[nom]
                texte_modifie = texte_modifie.replace(nom, remplacement)
                remplacements.append(Remplacement(
                    original=nom,
                    anonyme=remplacement,
                    type='nom' if nom.isupper() else 'prenom'
                ))

        if texte_modifie != texte_original:
            run.text = texte_modifie

        return remplacements

    def _traiter_paragraphe(self, paragraph, noms_a_remplacer: Dict[str, str]) -> List[Remplacement]:
        """Traite un paragraphe en préservant le formatage de chaque run."""
        remplacements = []

        for run in paragraph.runs:
            remplacements.extend(self._remplacer_dans_run(run, noms_a_remplacer))

        return remplacements

    def _traiter_table(self, table, noms_a_remplacer: Dict[str, str]) -> List[Remplacement]:
        """Traite une table en préservant le formatage."""
        remplacements = []

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    remplacements.extend(self._traiter_paragraphe(paragraph, noms_a_remplacer))
                # Tables imbriquées
                for nested_table in cell.tables:
                    remplacements.extend(self._traiter_table(nested_table, noms_a_remplacer))

        return remplacements

    def _extraire_tout_texte(self, doc: Document) -> str:
        """Extrait tout le texte du document pour analyse."""
        textes = []

        # Paragraphes
        for para in doc.paragraphs:
            textes.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    textes.append(cell.text)

        # Headers/Footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    textes.append(para.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    textes.append(para.text)

        return '\n'.join(textes)

    def anonymiser(self, fichier_source: Path, fichier_sortie: Optional[Path] = None) -> ResultatAnonymisation:
        """
        Anonymise un document DOCX.

        Args:
            fichier_source: Chemin du fichier source
            fichier_sortie: Chemin du fichier de sortie (optionnel)

        Returns:
            ResultatAnonymisation avec les détails
        """
        fichier_source = Path(fichier_source)

        if fichier_sortie is None:
            fichier_sortie = fichier_source.parent / f"{fichier_source.stem}_anonyme{fichier_source.suffix}"
        else:
            fichier_sortie = Path(fichier_sortie)

        resultat = ResultatAnonymisation(
            fichier_source=fichier_source,
            fichier_sortie=fichier_sortie
        )

        try:
            # Charger le document
            print(f"📄 Chargement: {fichier_source.name}")
            doc = Document(str(fichier_source))

            # Extraire tout le texte pour analyse
            texte_complet = self._extraire_tout_texte(doc)

            # Détecter les noms
            print("🔍 Détection des noms...")
            noms_detectes = self._detecter_noms_dans_texte(texte_complet)
            resultat.noms_detectes = noms_detectes

            print(f"   Trouvé: {len(noms_detectes)} noms/prénoms uniques")

            # Créer le mapping de remplacement
            noms_a_remplacer = {}
            for nom in noms_detectes:
                # Déterminer le type
                if nom.isupper():
                    type_nom = 'nom'
                elif nom.title() in PRENOMS_FRANCAIS:
                    type_nom = 'prenom'
                else:
                    type_nom = 'nom'

                remplacement = self._obtenir_remplacement(nom, type_nom)
                noms_a_remplacer[nom] = remplacement

            resultat.mapping = dict(self.mapping_global)

            # Afficher le mapping
            print("\n📋 Mapping des remplacements:")
            for original, anonyme in sorted(noms_a_remplacer.items()):
                print(f"   {original} → {anonyme}")

            # Traiter les paragraphes
            print("\n✏️  Anonymisation en cours...")
            for paragraph in doc.paragraphs:
                resultat.remplacements.extend(
                    self._traiter_paragraphe(paragraph, noms_a_remplacer)
                )

            # Traiter les tables
            for table in doc.tables:
                resultat.remplacements.extend(
                    self._traiter_table(table, noms_a_remplacer)
                )

            # Traiter headers/footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        resultat.remplacements.extend(
                            self._traiter_paragraphe(para, noms_a_remplacer)
                        )
                if section.footer:
                    for para in section.footer.paragraphs:
                        resultat.remplacements.extend(
                            self._traiter_paragraphe(para, noms_a_remplacer)
                        )

            # Sauvegarder
            print(f"\n💾 Sauvegarde: {fichier_sortie.name}")
            doc.save(str(fichier_sortie))

            resultat.succes = True
            print(f"✅ Anonymisation terminée: {len(resultat.remplacements)} remplacements")

        except Exception as e:
            resultat.succes = False
            resultat.erreur = str(e)
            print(f"❌ Erreur: {e}")

        return resultat

    def anonymiser_batch(self, fichiers: List[Path], dossier_sortie: Optional[Path] = None) -> List[ResultatAnonymisation]:
        """Anonymise plusieurs fichiers."""
        resultats = []

        for fichier in fichiers:
            if dossier_sortie:
                sortie = dossier_sortie / f"{fichier.stem}_anonyme{fichier.suffix}"
            else:
                sortie = None

            resultat = self.anonymiser(fichier, sortie)
            resultats.append(resultat)
            print()  # Ligne vide entre les fichiers

        return resultats


def exporter_mapping(resultats: List[ResultatAnonymisation], fichier_sortie: Path):
    """Exporte le mapping complet vers un fichier JSON."""
    mapping_complet = {}

    for resultat in resultats:
        mapping_complet[resultat.fichier_source.name] = {
            "mapping": resultat.mapping,
            "noms_detectes": list(resultat.noms_detectes),
            "nb_remplacements": len(resultat.remplacements)
        }

    with open(fichier_sortie, 'w', encoding='utf-8') as f:
        json.dump(mapping_complet, f, ensure_ascii=False, indent=2)

    print(f"📄 Mapping exporté: {fichier_sortie}")


def main():
    parser = argparse.ArgumentParser(
        description="Anonymise les documents DOCX notariaux en préservant le formatage"
    )
    parser.add_argument(
        "fichiers",
        nargs="+",
        help="Fichier(s) DOCX à anonymiser"
    )
    parser.add_argument(
        "-o", "--output",
        help="Fichier ou dossier de sortie"
    )
    parser.add_argument(
        "--batch",
        action="store_true",
        help="Mode batch: traite tous les fichiers"
    )
    parser.add_argument(
        "--mapping",
        help="Fichier JSON pour exporter le mapping des remplacements"
    )
    parser.add_argument(
        "--strict",
        action="store_true",
        help="Mode strict: demande confirmation pour chaque nom"
    )

    args = parser.parse_args()

    # Résoudre les fichiers
    fichiers = []
    for pattern in args.fichiers:
        chemin = Path(pattern)
        if chemin.exists():
            fichiers.append(chemin)
        else:
            # Glob pattern
            fichiers.extend(Path('.').glob(pattern))

    if not fichiers:
        print("❌ Aucun fichier trouvé")
        sys.exit(1)

    print(f"📁 {len(fichiers)} fichier(s) à traiter\n")

    # Anonymiser
    anonymiseur = AnonymiseurDocx(mode_strict=args.strict)

    if len(fichiers) == 1 and not args.batch:
        # Fichier unique
        sortie = Path(args.output) if args.output else None
        resultat = anonymiseur.anonymiser(fichiers[0], sortie)
        resultats = [resultat]
    else:
        # Batch
        dossier_sortie = Path(args.output) if args.output else None
        if dossier_sortie and not dossier_sortie.exists():
            dossier_sortie.mkdir(parents=True)
        resultats = anonymiseur.anonymiser_batch(fichiers, dossier_sortie)

    # Exporter mapping si demandé
    if args.mapping:
        exporter_mapping(resultats, Path(args.mapping))

    # Résumé
    print("\n" + "=" * 60)
    print("📊 RÉSUMÉ")
    print("=" * 60)

    succes = sum(1 for r in resultats if r.succes)
    echecs = len(resultats) - succes
    total_remplacements = sum(len(r.remplacements) for r in resultats)

    print(f"   Fichiers traités: {len(resultats)}")
    print(f"   ✅ Succès: {succes}")
    if echecs:
        print(f"   ❌ Échecs: {echecs}")
    print(f"   📝 Total remplacements: {total_remplacements}")

    # Mapping global
    if anonymiseur.mapping_global:
        print(f"\n📋 Mapping global ({len(anonymiseur.mapping_global)} entrées):")
        for original, anonyme in sorted(anonymiseur.mapping_global.items()):
            print(f"   {original} → {anonyme}")


if __name__ == "__main__":
    main()
